"""
专精特新中小企业申请书生成器
模板: 浙江省专精特新中小企业申请书.docx
结构: 1 张大表，116行 x 21列，10 个章节
"""
import os, shutil
from datetime import datetime
from docx import Document
from docx.shared import Pt

from modules.docgen.document_headers import add_company_header


TEMPLATE_PATH = os.path.expanduser("~/Desktop/申报通知/专精特新/浙江省专精特新中小企业申请书.docx")

# 列位置 — 选择每组中间列避免合并边界干扰
COL_YEARS = {"2023": 5, "2024": 10, "2025": 16}


def _safe_set(table, row, col, text):
    try:
        cell = table.rows[row].cells[col]
        for p in cell.paragraphs:
            for r in p.runs:
                r.text = ""
        if cell.paragraphs:
            cell.paragraphs[0].text = str(text) if text else ""
    except IndexError:
        pass


def _fill_financial_row(table, row_idx, data, field_name, unit=""):
    """填充三年财务数据行"""
    for yr, col in COL_YEARS.items():
        key = f"fin_{yr}_{field_name}"
        val = data.get(key, "")
        if val:
            _safe_set(table, row_idx, col, str(val))


def generate_zhuanjing(data: dict, output_path: str = None) -> str:
    if not os.path.exists(TEMPLATE_PATH):
        raise FileNotFoundError(f"模板不存在: {TEMPLATE_PATH}")

    if output_path is None:
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        project_root = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
        output_path = os.path.join(project_root, "outputs",
            f"专精特新中小企业申请书_{data.get('company_name','unknown')}_{ts}.docx")

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    shutil.copy(TEMPLATE_PATH, output_path)
    doc = Document(output_path)
    add_company_header(
        doc,
        data.get("company_name"),
        data.get("company_english_name"),
    )
    table = doc.tables[0]

    # === 一、企业基本情况 (Rows 1-14) ===
    _safe_set(table, 1, 4, data.get("company_name", ""))
    _safe_set(table, 2, 4, f"{data.get('province','')}省{data.get('city','')}市")
    _safe_set(table, 3, 4, data.get("address", ""))
    _safe_set(table, 3, 12, data.get("zipcode", ""))
    _safe_set(table, 4, 4, data.get("legal_rep", ""))
    _safe_set(table, 5, 4, data.get("shareholder", ""))
    _safe_set(table, 5, 7, data.get("actual_controller", ""))
    _safe_set(table, 5, 14, data.get("controller_nationality", ""))
    _safe_set(table, 6, 4, data.get("contact", ""))
    _safe_set(table, 6, 7, data.get("phone", ""))
    _safe_set(table, 6, 14, data.get("mobile", ""))
    _safe_set(table, 7, 4, data.get("fax", ""))
    _safe_set(table, 7, 7, data.get("email", ""))
    _safe_set(table, 8, 4, data.get("register_date", ""))
    _safe_set(table, 8, 7, data.get("register_capital", ""))
    _safe_set(table, 9, 4, data.get("credit_code", ""))
    _safe_set(table, 11, 4, data.get("industry_code", ""))
    _safe_set(table, 12, 4, data.get("sub_industry_code", ""))

    # === 二、经济效益和经营情况 (Rows 16-42) ===
    financial_fields = [
        (17, "employees"), (18, "rd_staff"),
        (19, "revenue"), (20, "main_revenue"),
        (21, "revenue_growth"), (22, "cost"),
        (23, "main_cost"), (24, "product_cost"),
        (25, "sales_expense"), (26, "admin_expense"),
        (27, "profit"), (28, "net_profit"),
        (29, "net_profit_growth"), (30, "assets"),
        (31, "net_assets"), (32, "liabilities"),
        (33, "debt_ratio"), (34, "tax"),
        (35, "equity_financing"), (36, "valuation"),
        (37, "bank_loan"), (38, "domestic_bond"),
        (39, "foreign_bond"),
    ]
    for row_idx, field in financial_fields:
        _fill_financial_row(table, row_idx, data, field)

    _safe_set(table, 40, 4, data.get("audit_report_code", ""))

    # === 三、专业化 (Rows 43-52) ===
    _safe_set(table, 44, 4, data.get("market_years", ""))
    _safe_set(table, 45, 4, data.get("main_revenue_ratio", ""))
    _safe_set(table, 46, 4, data.get("revenue_cagr", ""))

    # 主导产品 (Rows 48-50)
    for pi in range(3):
        row = 48 + pi
        p = data.get(f"product_{pi}", {})
        if p:
            _safe_set(table, row, 1, p.get("name", ""))
            _safe_set(table, row, 4, f"收入:{p.get('revenue','')}万元")

    # 标准 (Rows 51-52)
    _safe_set(table, 51, 4, data.get("std_international", ""))
    _safe_set(table, 51, 9, data.get("std_national", ""))
    _safe_set(table, 51, 14, data.get("std_industry", ""))
    _safe_set(table, 52, 4, data.get("std_names", ""))

    # === 六、创新能力 (Rows 72-88) ===
    # 研发机构 (Rows 73-80)
    rd_institutions = data.get("rd_institutions", {})
    for ri, inst_type in enumerate(["tech_academy", "tech_center", "eng_center", "design_center", "key_lab"]):
        row = 73 + ri
        inst = rd_institutions.get(inst_type, {})
        if inst:
            _safe_set(table, row, 7, inst.get("national", ""))
            _safe_set(table, row, 10, inst.get("province", ""))
            _safe_set(table, row, 13, inst.get("self_built", ""))

    _safe_set(table, 78, 7, "□有" if data.get("has_academician_station") else "□无")
    _safe_set(table, 79, 7, "□有" if data.get("has_postdoc_station") else "□无")
    _safe_set(table, 80, 4, data.get("partner_schools", ""))

    # 研发费用 (Rows 81-84)
    _fill_financial_row(table, 82, data, "rd_expense")
    _safe_set(table, 83, 4, data.get("rd_ratio_2023", ""))
    _safe_set(table, 83, 8, data.get("rd_ratio_2024", ""))
    _safe_set(table, 83, 14, data.get("rd_ratio_2025", ""))
    _safe_set(table, 84, 4, data.get("rd_staff_ratio_2023", ""))
    _safe_set(table, 84, 8, data.get("rd_staff_ratio_2024", ""))
    _safe_set(table, 84, 14, data.get("rd_staff_ratio_2025", ""))

    # === 五 & 七-九 文本型字段 ===
    text_fields = {
        # (row, col): field_name
        (56, 4): "market_position_desc",      # 市场地位
        (68, 4): "export_amount",              # 出口额
        (70, 4): "brand_count",                # 自有品牌数
        (71, 4): "brand_revenue",              # 品牌收入
        (90, 4): "industry_chain",             # 所属产业链
        (95, 4): "main_product_name",          # 主导产品名称
        (96, 4): "main_product_category",      # 主导产品类别
        (106, 4): "company_overview",          # 企业总体情况(2000字)
    }
    for (row, col), field in text_fields.items():
        _safe_set(table, row, col, data.get(field, ""))

    doc.save(output_path)
    return output_path
