import base64
import copy
import hashlib
import json
import mimetypes
import os
import re
import struct
import uuid
import shutil
import subprocess
import tempfile
import time
import zlib
from concurrent.futures import ThreadPoolExecutor
from datetime import datetime
from html import escape as html_escape
from io import BytesIO
from pathlib import Path

from flask import abort, current_app, jsonify, render_template, request, redirect, url_for, flash, send_file, session
from flask_login import login_required, current_user
from werkzeug.utils import secure_filename

from models import db, Company, ScoreRecord, ApplicationDraft, ExportJob
from modules.docgen import docgen_bp
from modules.docgen.generator import generate, TEMPLATE_PATH
from modules.docgen.generator_zhuanjing import generate_zhuanjing
from modules.docgen.date_logic import (
    enforce_transformation_wording,
    enforce_temporal_wording,
    evidence_record_date_context,
    event_date_context,
    project_temporal_context,
    system_today,
)
from modules.docgen.product_terms import (
    infer_ps_kind,
    normalize_ps_reference_text,
    ps_statement_case_template,
    ps_type_label,
)
from modules.docgen.relation_table_exporter import export_relation_table, import_relation_table
from modules.docgen.sales_contracts import (
    ensure_sales_contract_codes,
    remap_sales_contract_rows,
    sales_contract_file_sha256,
    selectable_sales_contracts,
)
from modules.docgen.staff_certificates import STAFF_CERTIFICATE_FIELDS, analyze_staff_certificate
from modules.healthcheck.engine import run_health_check, score_result_from_record
from modules.storage import blob_enabled, delete_file, ensure_local_file, persist_file
from modules.ai.analyzer import analyze
from modules.ai.llm_client import call_llm


STAFF_CERTIFICATE_UPLOAD_EXTENSIONS = {
    ".pdf",
    ".jpg",
    ".jpeg",
    ".png",
    ".webp",
    ".bmp",
    ".tif",
    ".tiff",
}

GAOXIN_ATTACHMENT_EXPORT_VERSION = "2026-07-19-v1"
GAOXIN_ATTACHMENT_EXPORT_JOB_TYPE = "gaoxin_attachments_pdf"
GAOXIN_ATTACHMENT_EXPORT_STALE_SECONDS = 15 * 60


def _upsert_application_draft(company, app_data, output_path):
    """同一公司只保留一条申报书记录；再次生成时更新原记录。"""
    draft = ApplicationDraft.query.filter_by(company_id=company.id).first()
    if not draft:
        draft = ApplicationDraft(company_id=company.id, app_type=company.app_type)
        db.session.add(draft)
    draft.app_type = company.app_type
    draft.sections_json = json.dumps(app_data, ensure_ascii=False)
    draft.docx_path = output_path
    return draft


def _load_company_data(company):
    if not company.data_json:
        return {}
    try:
        data = json.loads(company.data_json)
    except (json.JSONDecodeError, TypeError):
        return {}
    return data if isinstance(data, dict) else {}


def _company_health_check(company, data=None):
    """为页面和导出生成同一份实时体检结果。"""
    data = data if isinstance(data, dict) else _load_company_data(company)
    merged_data = _merge_relation_fields(data)
    latest_score = (
        ScoreRecord.query
        .filter_by(company_id=company.id)
        .order_by(ScoreRecord.created_at.desc(), ScoreRecord.id.desc())
        .first()
    )
    return run_health_check(
        merged_data,
        _load_gaoxin_attachments_from_data(data),
        score_result_from_record(latest_score),
    )


def _gaoxin_export_health_warning(company, health):
    """测试阶段允许导出，但把申报体检风险明确返回给用户。"""
    blockers = health.get("export_blockers", []) if isinstance(health, dict) else []
    if not blockers:
        return None
    return {
        "code": "HEALTH_CHECK_WARNING",
        "message": (
            f"当前为测试阶段，仍允许导出 PDF；申报体检发现 {len(blockers)} 个待处理项，"
            "导出文件不代表申报评估已通过。"
        ),
        "health_url": url_for("docgen.assessment", company_id=company.id),
        "blockers": blockers,
    }


def _latest_company_score(company):
    return (
        ScoreRecord.query
        .filter_by(company_id=company.id)
        .order_by(ScoreRecord.created_at.desc(), ScoreRecord.id.desc())
        .first()
    )


def _latest_company_draft(company):
    return (
        ApplicationDraft.query
        .filter_by(company_id=company.id)
        .order_by(ApplicationDraft.created_at.desc(), ApplicationDraft.id.desc())
        .first()
    )


def _score_analysis_from_record(score, data):
    """读取评分时生成的分析，旧记录没有分析时用本地规则补齐。"""
    if not score:
        return None
    try:
        analysis = json.loads(score.ai_analysis) if score.ai_analysis else None
    except (json.JSONDecodeError, TypeError):
        analysis = None
    if analysis:
        return analysis
    try:
        breakdown = json.loads(score.breakdown_json or "[]")
    except (json.JSONDecodeError, TypeError):
        breakdown = []
    pass_score = 71 if score.score_type == "高新技术" else (60 if score.score_type == "小巨人" else 50)
    return analyze(
        {
            "rule_type": score.score_type,
            "total_score": score.total_score,
            "full_score": 100,
            "pass_score": pass_score,
            "passed": (score.total_score or 0) >= pass_score,
            "breakdown": breakdown,
            "warnings": [],
        },
        data,
        use_llm=False,
    )


def _has_meaningful_application_draft(draft):
    if not draft or not draft.sections_json:
        return False
    try:
        sections = json.loads(draft.sections_json)
    except (json.JSONDecodeError, TypeError):
        return False
    if not isinstance(sections, dict):
        return False
    return any(value not in (None, "", [], {}) for value in sections.values())


def _has_legacy_application_input(data):
    """识别统一标记上线前已经保存过的申报书资料。"""
    if not isinstance(data, dict):
        return False
    if data.get("_application_input_saved"):
        return True

    relation = data.get("gaoxin_relation_table")
    if isinstance(relation, dict):
        if str(relation.get("tech_field_path") or "").strip():
            return True
        rows = relation.get("rows")
        if isinstance(rows, list) and any(
            isinstance(row, dict)
            and any(str(row.get(key) or "").strip() for key in (
                "rd_activity",
                "rd_period",
                "ps_name",
                "result_name",
                "technology",
            ))
            for row in rows
        ):
            return True

    explicit_keys = {
        "province",
        "city",
        "tech_field",
        "registration_date",
        "register_date",
        "establish_date",
        "company_established_at",
        "company_register_date",
        "innovation_ip_role",
        "innovation_transform",
        "innovation_rd_mgmt",
        "innovation_talent",
        "gaoxin_system_docs",
    }
    if any(data.get(key) not in (None, "", [], {}) for key in explicit_keys):
        return True
    return any(
        key.startswith(("cv_", "attachment_", "ps_", "rd_"))
        and value not in (None, "", [], {})
        for key, value in data.items()
    )


def _assessment_input_state(company, data=None):
    data = data if isinstance(data, dict) else _load_company_data(company)
    score = _latest_company_score(company)
    draft = _latest_company_draft(company)
    score_ready = score is not None
    application_ready = _has_meaningful_application_draft(draft) or _has_legacy_application_input(data)
    missing = []
    if not score_ready:
        missing.append({
            "key": "score",
            "title": "完成评分",
            "detail": "先提交评分表，评估页才会使用企业的实际评分结果。",
            "url": url_for("scoring.index"),
        })
    if not application_ready:
        application_url = (
            url_for("docgen.gaoxin_relation_table", company_id=company.id)
            if company.app_type == "高新技术"
            else url_for("docgen.fill", company_id=company.id)
        )
        missing.append({
            "key": "application",
            "title": "填写申报书",
            "detail": "先保存申请书、关系表或申报材料，评估页才会判断资料完整度。",
            "url": application_url,
        })
    return {
        "score": score,
        "draft": draft,
        "score_ready": score_ready,
        "application_ready": application_ready,
        "ready": score_ready and application_ready,
        "missing": missing,
    }


def _gaoxin_attachment_export_fingerprint(company, data):
    payload = {
        "version": GAOXIN_ATTACHMENT_EXPORT_VERSION,
        "company_id": company.id,
        "company_name": company.name,
        "data": data if isinstance(data, dict) else {},
    }
    encoded = json.dumps(
        payload,
        ensure_ascii=False,
        sort_keys=True,
        separators=(",", ":"),
        default=str,
    ).encode("utf-8")
    return hashlib.sha256(encoded).hexdigest()


def _gaoxin_attachment_export_relative_path(job):
    return os.path.join(
        "generated_exports",
        str(job.user_id),
        str(job.company_id),
        f"{job.fingerprint}.pdf",
    )


def _gaoxin_attachment_export_local_path(job):
    output_root = os.path.abspath(current_app.config["OUTPUT_FOLDER"])
    relative_path = str(job.result_path or "").strip()
    target = os.path.abspath(os.path.join(output_root, relative_path))
    if target != output_root and not target.startswith(output_root + os.sep):
        abort(403)
    return target


def _update_export_job(job, *, status=None, stage=None, progress=None, error=None):
    if status is not None:
        job.status = status
    if stage is not None:
        job.stage = stage
    if progress is not None:
        job.progress = max(0, min(100, int(progress)))
    if error is not None:
        job.error_message = str(error)[:4000]
    job.updated_at = datetime.utcnow()
    db.session.commit()


def _export_job_is_stale(job):
    reference_time = job.updated_at or job.started_at or job.created_at
    return bool(
        reference_time
        and (datetime.utcnow() - reference_time).total_seconds()
        > GAOXIN_ATTACHMENT_EXPORT_STALE_SECONDS
    )


def _expire_stale_export_job(job):
    if job.status not in {"queued", "running"} or not _export_job_is_stale(job):
        return False
    _update_export_job(
        job,
        status="failed",
        stage="生成任务已中断，可重新发起",
        error="PDF 生成任务长时间没有进度，已结束本次任务，请重新发起导出。",
    )
    return True


def _export_job_payload(job):
    payload = {
        "ok": True,
        "job": {
            "id": job.id,
            "status": job.status,
            "stage": job.stage,
            "progress": int(job.progress or 0),
            "error": job.error_message or "",
            "result_size": int(job.result_size or 0),
            "duration_seconds": round(float(job.duration_seconds or 0), 2),
            "created_at": job.created_at.isoformat() if job.created_at else "",
            "updated_at": job.updated_at.isoformat() if job.updated_at else "",
            "status_url": url_for(
                "docgen.gaoxin_attachments_pdf_job_status",
                company_id=job.company_id,
                job_id=job.id,
            ),
        },
    }
    if job.status in {"queued", "running"}:
        payload["job"]["run_url"] = url_for(
            "docgen.gaoxin_attachments_pdf",
            company_id=job.company_id,
            job=job.id,
        )
    if job.status == "ready":
        payload["job"]["download_url"] = url_for(
            "docgen.gaoxin_attachments_pdf_job_download",
            company_id=job.company_id,
            job_id=job.id,
        )
    return payload


def _owned_export_job(company_id, job_id):
    return ExportJob.query.filter_by(
        id=job_id,
        company_id=company_id,
        user_id=current_user.id,
        job_type=GAOXIN_ATTACHMENT_EXPORT_JOB_TYPE,
    ).first_or_404()


def _company_english_name(company, data=None):
    """Return the maintained English company name, with pinyin as a fallback."""
    from modules.docgen.document_headers import company_name_to_pinyin

    data = data if isinstance(data, dict) else {}
    system_docs = data.get("gaoxin_system_docs") if isinstance(data.get("gaoxin_system_docs"), dict) else {}
    system_base = system_docs.get("base") if isinstance(system_docs.get("base"), dict) else {}
    explicit_name = str(data.get("company_english_name") or system_base.get("company_english_name") or "").strip()
    if explicit_name:
        return explicit_name

    company_name = str(getattr(company, "name", "") or "").strip()
    if not company_name:
        return ""
    return company_name_to_pinyin(company_name)


def _attachment_section_map():
    return {item["key"]: item for item in GAOXIN_ATTACHMENT_SECTIONS}


def _load_gaoxin_attachments_from_data(data):
    saved = data.get("gaoxin_attachments") if isinstance(data.get("gaoxin_attachments"), dict) else {}
    attachments = {}
    for section in GAOXIN_ATTACHMENT_SECTIONS:
        item = saved.get(section["key"]) if isinstance(saved.get(section["key"]), dict) else {}
        files = item.get("files") if isinstance(item.get("files"), list) else []
        attachments[section["key"]] = {"files": [file for file in files if isinstance(file, dict)]}
    # Keep files uploaded before the section was renamed available in the new
    # near-one-year high-tech product income audit section.
    if "hitech_income_audit" not in saved:
        legacy = saved.get("special_audit") if isinstance(saved.get("special_audit"), dict) else {}
        legacy_files = legacy.get("files") if isinstance(legacy.get("files"), list) else []
        attachments["hitech_income_audit"] = {"files": [file for file in legacy_files if isinstance(file, dict)]}
    for key, item in saved.items():
        if key in attachments or not isinstance(item, dict):
            continue
        files = item.get("files") if isinstance(item.get("files"), list) else []
        attachments[key] = {"files": [file for file in files if isinstance(file, dict)]}
    return attachments


def _load_gaoxin_attachments(company):
    return _load_gaoxin_attachments_from_data(_load_company_data(company))


def _ordered_ip_certificate_files(attachments, ip_rows):
    """Return knowledge-property certificate files in the same order as the detail table."""
    ip_files = attachments.get("ip", {}).get("files", []) if isinstance(attachments, dict) else []
    pending_files = [file for file in ip_files if isinstance(file, dict)]
    ordered_files = []
    seen_ids = set()

    def add(file_meta):
        file_id = str(file_meta.get("id") or "")
        path = str(file_meta.get("relative_path") or "")
        identity = file_id or path
        if not path or not identity or identity in seen_ids:
            return
        seen_ids.add(identity)
        ordered_files.append(file_meta)

    for row in ip_rows or []:
        cert_id = str(row.get("cert_id") or "").strip()
        seq = str(row.get("seq") or "").strip()
        for file_meta in pending_files:
            if (
                cert_id and str(file_meta.get("ip_cert_id") or "").strip() == cert_id
            ) or (
                seq and str(file_meta.get("ip_seq") or "").strip() == seq
            ):
                add(file_meta)

    for file_meta in pending_files:
        add(file_meta)
    return ordered_files


def _save_gaoxin_attachment_form_data(company, form_data):
    data = _load_company_data(company)
    for key in list(data.keys()):
        if (
            re.match(r"^cv_\d+_", key)
            or re.match(r"^attachment_rd_staff_\d+_", key)
            or re.match(r"^attachment_staff_month_\d+_", key)
        ):
            data.pop(key, None)
    for key, value in form_data.items():
        if key.startswith("cv_") or key in GAOXIN_SYSTEM_FRAMEWORK_FIELDS or key.startswith("attachment_"):
            data[key] = value
    data["_application_input_saved"] = True
    _sync_staff_month_counts(data)
    company.data_json = json.dumps(data, ensure_ascii=False)
    db.session.commit()


def _attachment_upload_root():
    return os.path.abspath(current_app.config["UPLOAD_FOLDER"])


def _attachment_relative_path(user_id, company_id, section_key, filename):
    return os.path.join("gaoxin_attachments", str(user_id), str(company_id), section_key, filename)


def _safe_attachment_path(relative_path):
    root = _attachment_upload_root()
    target = os.path.abspath(os.path.join(root, relative_path))
    if target != root and not target.startswith(root + os.sep):
        abort(403)
    return ensure_local_file(target, relative_path)


def _source_upload_path(relative_path):
    root = _attachment_upload_root()
    target = os.path.abspath(os.path.join(root, relative_path or ""))
    if target != root and not target.startswith(root + os.sep):
        return None
    return ensure_local_file(target, relative_path)


def _extract_pdf_text(path):
    path = str(path or "")
    if not path:
        return ""
    try:
        try:
            import fitz
        except Exception:
            try:
                import pymupdf as fitz
            except Exception:
                fitz = None
        if fitz:
            doc = fitz.open(path)
            try:
                text = "\n".join((page.get_text() or "") for page in doc).strip()
                if len(text) >= 30:
                    return text
            finally:
                doc.close()
    except Exception:
        current_app.logger.exception("PDF 文本层提取失败")

    try:
        from modules.parser.ip_analyzer import extract_text
        text = extract_text(path)
        if len(str(text or "").strip()) >= 30:
            return text
    except Exception:
        current_app.logger.exception("PDF 通用文本提取失败")

    return _extract_pdf_text_with_cli_ocr(path)


def _extract_pdf_text_with_cli_ocr(path):
    import glob
    import shutil
    import subprocess
    import tempfile

    if not shutil.which("tesseract"):
        return ""
    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            images = []
            if shutil.which("pdftoppm"):
                prefix = os.path.join(tmpdir, "page")
                subprocess.run(
                    ["pdftoppm", "-png", "-r", "220", "-f", "1", "-l", "5", path, prefix],
                    check=True,
                    stdout=subprocess.DEVNULL,
                    stderr=subprocess.DEVNULL,
                    timeout=60,
                )
                images = sorted(glob.glob(os.path.join(tmpdir, "page-*.png")))
            elif shutil.which("qlmanage"):
                subprocess.run(
                    ["qlmanage", "-t", "-s", "2200", "-o", tmpdir, path],
                    check=True,
                    stdout=subprocess.DEVNULL,
                    stderr=subprocess.DEVNULL,
                    timeout=60,
                )
                images = sorted(glob.glob(os.path.join(tmpdir, "*.png")))
            if not images:
                return ""

            texts = []
            for image in images[:5]:
                for lang in ["chi_sim+eng", "eng"]:
                    try:
                        result = subprocess.run(
                            ["tesseract", image, "stdout", "-l", lang, "--psm", "6"],
                            check=True,
                            capture_output=True,
                            text=True,
                            timeout=45,
                        )
                        if result.stdout.strip():
                            texts.append(result.stdout.strip())
                            break
                    except Exception:
                        continue
            return "\n".join(texts).strip()
    except Exception:
        current_app.logger.exception("PDF OCR 提取失败")
        return ""


def _extract_image_text_with_cli_ocr(path):
    if not shutil.which("tesseract"):
        return ""
    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            normalized_path = os.path.join(tmpdir, "certificate.png")
            try:
                from PIL import Image, ImageOps

                with Image.open(path) as image:
                    image = ImageOps.exif_transpose(image).convert("RGB")
                    if max(image.size) < 1800:
                        scale = 1800 / max(image.size)
                        image = image.resize(
                            (round(image.width * scale), round(image.height * scale))
                        )
                    image.save(normalized_path, format="PNG")
            except Exception:
                current_app.logger.exception("人员证书图片预处理失败：%s", path)
                normalized_path = path

            for lang in ["chi_sim+eng", "eng"]:
                try:
                    result = subprocess.run(
                        ["tesseract", normalized_path, "stdout", "-l", lang, "--psm", "6"],
                        check=True,
                        capture_output=True,
                        text=True,
                        timeout=45,
                    )
                    if result.stdout.strip():
                        return result.stdout.strip()
                except Exception:
                    continue
    except Exception:
        current_app.logger.exception("人员证书图片 OCR 提取失败：%s", path)
    return ""


def _extract_staff_certificate_text(path):
    extension = Path(path).suffix.lower()
    if extension == ".pdf":
        return _extract_pdf_text(path)
    return _extract_image_text_with_cli_ocr(path)


def _compact_contract_text(text):
    text = str(text or "")
    text = re.sub(r"(?<=[一-鿿])[ \t]+(?=[一-鿿])", "", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _clean_contract_candidate(value):
    value = str(value or "").strip()
    value = re.sub(r"\s+", "", value)
    value = re.sub(r"^(?:货物|产品|服务|项目|标的|名称|内容|规格|型号|序号|数量|单位|单价|金额|备注|明细)[：:：]?", "", value)
    value = re.split(r"(?:数量|单位|单价|金额|总价|备注|规格型号|型号规格|交付|验收|付款|税率|小计|合计|技术要求)[：:：]?", value, maxsplit=1)[0]
    value = re.sub(r"(?:[A-Za-z]{0,6}[-_]?\d+(?:\.\d+)?|V\d+(?:\.\d+)*|v\d+(?:\.\d+)*|\d+(?:\.\d+)?)(?:台|套|个|件|批|项|元|万元|%)?.*$", "", value)
    value = re.sub(r"(?<=[一-鿿])[A-Za-z0-9_.-]+$", "", value)
    value = value.strip(" ：:，,；;。.【】[]（）()《》<>\"'")
    if len(value) < 2 or len(value) > 40:
        return ""
    noise_words = ["本合同", "购销合同", "销售合同", "采购合同", "合同法", "甲方", "乙方", "双方", "签订", "保证", "付款", "金额", "地址", "电话", "联系人", "开户", "银行", "账号", "人民币", "有限公司", "中国计量大学"]
    if any(word in value for word in noise_words):
        return ""
    if re.search(r"^(?:需方|供方|买方|卖方|合同|协议|日期)$", value):
        return ""
    if re.fullmatch(r"[\dA-Za-z_.\-/]+", value):
        return ""
    return value


def _contract_filename_product_text(filename):
    text = os.path.splitext(os.path.basename(str(filename or "")))[0]
    text = re.sub(r"[【\[].*?[】\]]", "", text)
    text = re.sub(r"(?:双方盖章|盖章|扫描件|合同|购销|销售|采购|协议|订单|中国计量大学|甲方|乙方)", "", text)
    text = re.sub(r"[\s_\-—－]+", "", text)
    return text.strip(" ：:，,；;。.【】[]（）()《》<>\"'")


def _extract_contract_product_candidates(text):
    text = _compact_contract_text(text)
    candidates = []

    def add(value):
        cleaned = _clean_contract_candidate(value)
        if cleaned and cleaned not in candidates:
            candidates.append(cleaned)

    label_patterns = [
        r"(?:产品名称|货物名称|商品名称|设备名称|软件名称|服务名称|项目名称|合同标的|采购内容|销售内容|标的名称|产品/服务名称)\s*[：:]\s*([^\n；;。]+)",
        r"(?:名称及规格|产品及规格|货物及规格|设备及规格|标的物)\s*[：:]\s*([^\n；;。]+)",
    ]
    for pattern in label_patterns:
        for match in re.finditer(pattern, text):
            add(match.group(1))
            if len(candidates) >= 12:
                return candidates

    lines = [line.strip() for line in text.splitlines() if line.strip()]
    header_words = ("产品名称", "货物名称", "商品名称", "设备名称", "服务名称", "项目名称", "名称", "规格", "型号", "数量", "金额")
    for index, line in enumerate(lines):
        if any(word in line for word in header_words) and index + 1 < len(lines):
            for next_line in lines[index + 1:index + 5]:
                parts = re.split(r"\s{2,}|\t|[|｜]", next_line)
                for part in parts[:3]:
                    add(part)
        if len(candidates) >= 12:
            return candidates

    noun_pattern = r"[一-鿿A-Za-z0-9]{2,24}(?:系统|平台|设备|装置|软件|模块|仪|机|器|线|材料|制品|产品|服务|解决方案)"
    for match in re.finditer(noun_pattern, text):
        add(match.group(0))
        if len(candidates) >= 12:
            break
    return candidates


def _extract_sales_contract_info(text, row, filename=""):
    text = _compact_contract_text(text)
    filename_candidates = _extract_contract_product_candidates(_contract_filename_product_text(filename))
    if not text:
        keywords = "；".join(filename_candidates[:8])
        summary = f"合同主要涉及{'、'.join(filename_candidates[:3])}等产品/服务。" if filename_candidates else ""
        return {"summary": summary, "keywords": keywords}
    product_candidates = filename_candidates[:]
    for candidate in _extract_contract_product_candidates(text):
        if candidate not in product_candidates:
            product_candidates.append(candidate)
    important_snippets = []
    for line in text.splitlines():
        if re.search(r"产品名称|货物名称|商品名称|设备名称|服务名称|项目名称|合同标的|采购内容|销售内容|产品内容|服务内容|技术要求|规格型号|标的物", line):
            important_snippets.append(line.strip())
        if len(important_snippets) >= 20:
            break
    excerpt = "\n".join(important_snippets + [text[:12000]])[:14000]
    context = "\n".join(
        f"{label}：{value}"
        for label, key in [
            ("知识产权", "ip_name"),
            ("授权号", "ip_auth_no"),
            ("成果名称", "result_name"),
            ("PS名称", "ps_name"),
            ("RD名称", "rd_activity"),
        ]
        for value in [str((row or {}).get(key) or "").strip()]
        if value
    )
    prompt = f"""请从销售合同文本中提取与高新技术成果转化相关的信息，重点识别合同里的主要产品/服务内容。

当前关系行：
{context or '暂无'}

程序预提取的产品/服务候选：
{'；'.join(product_candidates) or '暂无'}

销售合同关键片段和正文：
{excerpt}

要求：
1. summary 必须优先写清楚合同标的、主要产品/服务名称、产品功能/用途、交付内容和应用场景；如果候选中有合理产品名称，应纳入摘要。
2. keywords 提取 8-16 个关键词，前 3 个优先放主要产品/服务名称或产品类别，其后放功能、技术方向、应用场景、交付内容、行业对象。
3. 不要把甲方、乙方、合同编号、金额、付款、税率、地址、联系人作为关键词。
4. 不要编造金额、日期、客户名称或合同中没有的性能参数。
5. 只输出 JSON，不要 markdown，不要解释。

JSON 格式：
{{"summary": "合同主要产品内容摘要", "keywords": "产品名称1；产品类别；核心功能；应用场景"}}"""
    result = call_llm([
        {"role": "system", "content": "你是高新技术企业申报材料顾问，擅长从销售合同、采购合同和表格型 PDF 中识别主要产品/服务内容。输出必须是可解析 JSON。"},
        {"role": "user", "content": prompt},
    ], temperature=0.1, max_tokens=900, timeout=60, max_attempts=1)
    data = _extract_json_object(result.get("content")) if result.get("success") else None
    if isinstance(data, dict):
        summary = str(data.get("summary") or data.get("摘要") or "").strip()
        keywords = _normalize_relation_value(data.get("keywords") or data.get("关键词") or "")
        candidate_keywords = "；".join(product_candidates[:4])
        keyword_candidates = []
        for item in _as_text_list(keywords):
            if item and item not in keyword_candidates:
                keyword_candidates.append(item)
        merged_keywords = _as_text_list(candidate_keywords) + keyword_candidates
        keywords = "；".join(dict.fromkeys(merged_keywords[:16]))
        if product_candidates and not any(candidate in summary for candidate in product_candidates[:3]):
            summary = f"合同主要涉及{'、'.join(product_candidates[:3])}等产品/服务。{summary}" if summary else f"合同主要涉及{'、'.join(product_candidates[:3])}等产品/服务。"
        return {"summary": summary, "keywords": keywords}
    compact = re.sub(r"\s+", " ", text)[:500]
    words = list(product_candidates[:6])
    for token in re.findall(r"[一-鿿]{2,10}(?:系统|平台|设备|装置|软件|模块|产品|服务)?", compact):
        cleaned = _clean_contract_candidate(token)
        if cleaned and cleaned not in words and cleaned not in ["合同", "甲方", "乙方", "双方", "签订", "销售", "付款", "金额"]:
            words.append(cleaned)
        if len(words) >= 12:
            break
    summary = f"合同主要涉及{'、'.join(product_candidates[:3])}等产品/服务。" if product_candidates else compact
    return {"summary": summary, "keywords": "；".join(words)}


def _ensure_ip_cert_id(cert):
    if isinstance(cert, dict) and not cert.get("id"):
        cert["id"] = uuid.uuid4().hex
    return cert


def sync_ip_cert_pdfs_to_attachments(company, certs):
    certs = [_ensure_ip_cert_id(cert) for cert in (certs or []) if isinstance(cert, dict)]
    data = _load_company_data(company)
    attachments = _load_gaoxin_attachments_from_data(data)
    ip_files = attachments.setdefault("ip", {"files": []}).setdefault("files", [])
    active_cert_ids = {str(cert.get("id") or "") for cert in certs if cert.get("id")}
    kept_files = []
    for file_meta in ip_files:
        if not isinstance(file_meta, dict):
            continue
        auto_synced = file_meta.get("auto_synced") is True and file_meta.get("source") == "score_ip_upload"
        if auto_synced and str(file_meta.get("ip_cert_id") or "") not in active_cert_ids:
            try:
                path = _safe_attachment_path(file_meta.get("relative_path", ""))
                if os.path.exists(path):
                    os.remove(path)
            except Exception:
                pass
            continue
        kept_files.append(file_meta)
    ip_files[:] = kept_files

    existing_by_cert = {str(item.get("ip_cert_id") or ""): item for item in ip_files if item.get("ip_cert_id")}
    existing_by_hash = {str(item.get("source_pdf_sha256") or ""): item for item in ip_files if item.get("source_pdf_sha256")}

    for index, cert in enumerate(certs, start=1):
        source_pdf = cert.get("source_pdf") if isinstance(cert.get("source_pdf"), dict) else None
        if not source_pdf:
            continue
        ip_seq = f"IP{str(index).zfill(2)}"
        cert_id = str(cert.get("id") or "")
        source_hash = str(source_pdf.get("sha256") or "")
        existing = existing_by_cert.get(cert_id) or (existing_by_hash.get(source_hash) if source_hash else None)
        if existing:
            existing["ip_seq"] = ip_seq
            existing["ip_cert_id"] = cert_id
            if source_hash:
                existing["source_pdf_sha256"] = source_hash
            source_pdf["sync_status"] = "synced"
            source_pdf["attachment_file_id"] = existing.get("id")
            source_pdf["attachment_relative_path"] = existing.get("relative_path")
            continue

        source_path = _source_upload_path(source_pdf.get("relative_path"))
        if not source_path or not os.path.exists(source_path):
            continue
        original_filename = source_pdf.get("original_filename") or cert.get("filename") or "知识产权.pdf"
        safe_name = secure_filename(original_filename) or f"ip_{uuid.uuid4().hex}.pdf"
        if not safe_name.lower().endswith(".pdf"):
            safe_name = f"{os.path.splitext(safe_name)[0]}.pdf"
        file_id = uuid.uuid4().hex
        stored_filename = f"{file_id}_{safe_name}"
        relative_path = _attachment_relative_path(company.user_id, company.id, "ip", stored_filename)
        target_path = _safe_attachment_path(relative_path)
        os.makedirs(os.path.dirname(target_path), exist_ok=True)
        shutil.copy2(source_path, target_path)
        persist_file(target_path, relative_path)
        file_meta = {
            "id": file_id,
            "original_filename": original_filename,
            "stored_filename": stored_filename,
            "relative_path": relative_path,
            "uploaded_at": source_pdf.get("uploaded_at") or datetime.utcnow().isoformat(timespec="seconds"),
            "ip_seq": ip_seq,
            "ip_cert_id": cert_id,
            "source_pdf_sha256": source_hash,
            "source": "score_ip_upload",
            "auto_synced": True,
        }
        ip_files.append(file_meta)
        existing_by_cert[cert_id] = file_meta
        if source_hash:
            existing_by_hash[source_hash] = file_meta
        source_pdf["sync_status"] = "synced"
        source_pdf["attachment_file_id"] = file_id
        source_pdf["attachment_relative_path"] = relative_path

    data["gaoxin_attachments"] = attachments
    company.data_json = json.dumps(data, ensure_ascii=False)
    return certs


GAOXIN_FINANCE_YEARS = ["2023", "2024", "2025"]

RD_PROJECT_APPLICATION_TEMPLATE = """科研项目书
请严格按以下标题顺序撰写，并结合已提供的项目名称、周期、技术领域、已有目的、创新点、成果、预算、知识产权和PS关系扩写。

立项背景与必要性
【说明行业或业务背景、现有问题、立项原因和必要性】

拟解决的技术问题
【明确研究对象、应用场景、关键难点和问题边界】

研发目标与考核指标
【分别说明功能、性能、稳定性、适用性和成果归档目标；没有量化数据时写核验方式，不编造数值】

研发内容
【分项说明需求分析、方案设计、关键技术研发、联调测试、优化和归档内容】

技术路线
【按先后顺序说明技术方法、实施步骤和验证闭环】

创新点
【结合已有创新点，说明技术原理、实现方法以及与常规方案的区别】

项目组织与任务分工
【说明负责人、技术研发、测试验证、资料管理和财务归集职责；姓名缺失写待补充】

经费预算
【使用已有预算；费用明细缺失写待补充，不自行分摊金额】

计划进度
【严格在项目起止时间内按顺序安排准备、开发、测试优化和总结验收，不得出现重叠或超出周期的日期】

过程记录与质量控制
【说明立项、评审、研发、测试、整改、变更、工时、费用和成果资料的记录要求】

预期成果与阶段成果
【根据当前项目状态使用计划时、进行时或完成时，结合已有成果内容撰写】

RD-IP-PS关联
【完整写明RD项目、成果名称、知识产权和PS的对应关系，不使用省略号】

验收指标对照
【逐项对照研发目标、完成情况和证明材料；项目未结束时只写阶段检查，不写已通过验收】

验收意见
【仅依据已提供的实际验收记录撰写；项目未结束、日期不完整或没有实际验收记录时明确待验收或待补充】

验收结论
【只有已提供的实际验收记录明确支持时才可以写同意验收；到达计划结束时间本身不代表项目已经完成或通过验收，没有实际记录时必须写待验收或待补充】"""

GAOXIN_ATTACHMENT_SECTIONS = [
    {"no": "1", "key": "commitment", "title": "承诺书", "kind": "pdf", "hint": "导出附件 PDF 时自动生成承诺书，无需上传。"},
    {"no": "2", "key": "application_pdf", "title": "申请书", "kind": "pdf", "hint": "上传签字盖章后的申请书 PDF。"},
    {"no": "3", "key": "business_license", "title": "营业执照", "kind": "pdf", "hint": "上传营业执照扫描件 PDF。"},
    {"no": "4", "key": "ip", "title": "知识产权", "kind": "pdf", "hint": "上传专利、软著等知识产权证明 PDF，可多份。"},
    {"no": "5", "key": "staff_statement", "title": "人员情况说明", "kind": "placeholder", "hint": "模板表格待提供。"},
    {"no": "6", "key": "rd_project_table", "title": "科研项目表", "kind": "placeholder", "hint": "模板表格待提供。"},
    {"no": "7", "key": "rd_expense_audit", "title": "近三年研发费用专项审计报告", "kind": "pdf", "hint": "按年度上传 2023、2024、2025 年研发费用专项审计报告 PDF。"},
    {"no": "8", "key": "hitech_products", "title": "高新技术产品说明", "kind": "placeholder", "hint": "维护产品汇总表、PS 情况说明及每个产品对应的证明附件。"},
    {"no": "9", "key": "hitech_income_audit", "title": "近一年高新产品收入专项审计报告", "kind": "pdf", "hint": "上传近一年高新技术产品（服务）收入专项审计报告 PDF。"},
    {"no": "10", "key": "achievement_transform", "title": "成果转化文件", "kind": "achievement", "hint": "维护成果转化明细，沿用申请书导出数据。"},
    {"no": "11", "key": "system_docs", "title": "制度", "kind": "system", "hint": "维护组织管理制度框架，并跳转生成制度正文。"},
    {"no": "12", "key": "annual_audit", "title": "年度审计报告", "kind": "pdf", "hint": "上传近三年年度审计报告 PDF，可多份。"},
    {"no": "13", "key": "tax_settlement", "title": "汇算清缴文件", "kind": "pdf", "hint": "上传企业所得税汇算清缴文件 PDF，可多份。"},
]

GAOXIN_ATTACHMENT_PDF_KEYS = {
    item["key"]
    for item in GAOXIN_ATTACHMENT_SECTIONS
    if item["kind"] == "pdf" and item["key"] != "commitment"
} | {"staff_statement", "achievement_transform", "hitech_products"}

GAOXIN_ACHIEVEMENT_EVIDENCE_UPLOAD_TYPES = [
    {"key": "patent_cert", "title": "专利证书", "hint": "上传该成果对应的专利证书 PDF。"},
    {"key": "sales_contract", "title": "销售合同", "hint": "上传该成果转化相关销售合同 PDF。"},
    {"key": "sales_invoice", "title": "销售发票", "hint": "上传该成果转化相关销售发票 PDF。"},
    {"key": "product_photo", "title": "产品照片", "hint": "上传产品照片整理成的 PDF。"},
]

GAOXIN_ACHIEVEMENT_EVIDENCE_EDIT_TYPES = [
    {"key": "test_report", "title": "检测报告", "ai_field": "achievement_test_report"},
    {"key": "user_report", "title": "用户使用报告", "ai_field": "achievement_user_report"},
]

GAOXIN_HITECH_PRODUCT_EVIDENCE_TYPES = [
    {"key": "patent_cert", "title": "专利证书", "hint": "上传该产品对应的专利证书 PDF。"},
    {"key": "test_report", "title": "检测报告", "hint": "上传本公司对本技术进行检测形成的检测报告 PDF。"},
    {"key": "sales_contract", "title": "销售合同", "hint": "上传该产品对应的销售合同 PDF。"},
    {"key": "sales_invoice", "title": "销售发票", "hint": "上传该产品对应的销售发票 PDF。"},
    {"key": "product_photo", "title": "产品照片", "hint": "上传产品照片整理成的 PDF。"},
    {"key": "user_report", "title": "用户使用报告", "hint": "上传客户对本公司产品和服务评价形成的用户使用报告 PDF。"},
]

GAOXIN_SYSTEM_FRAMEWORK_FIELDS = {
    "system_rd_files",
    "system_rd_points",
    "system_org_files",
    "system_org_points",
    "system_transform_files",
    "system_transform_points",
    "system_talent_files",
    "system_talent_points",
}

GAOXIN_SYSTEM_DOC_TYPES = [
    {
        "key": "rd_project",
        "title": "科研项目立项管理制度",
        "short": "项目立项",
        "guide": "重点写清研发项目来源、立项申请、可行性论证、预算审核、负责人审批、项目编号、过程记录和验收归档。",
    },
    {
        "key": "rd_accounting",
        "title": "研发投入核算体系制度",
        "short": "投入核算",
        "guide": "重点写清研发投入归集口径、财务部门与研发部门职责、预算控制、费用审批、项目维度核算和年度汇总机制。",
    },
    {
        "key": "rd_auxiliary_ledger",
        "title": "研发费用辅助账制度",
        "short": "辅助账",
        "guide": "重点写清按研发项目建立辅助账、费用科目、凭证附件、月度归集、年度结转、审计核查和留存要求。",
    },
    {
        "key": "rd_center",
        "title": "研发中心管理制度",
        "short": "研发中心",
        "guide": "重点写清研发机构设置、科研条件、设备场地、岗位职责、会议机制、项目组织和研发档案管理。",
    },
    {
        "key": "industry_university_research",
        "title": "产学研合作管理制度",
        "short": "产学研",
        "guide": "重点写清合作对象筛选、合作形式、合同签署、知识产权归属、成果交付、保密要求和合作项目验收。",
    },
    {
        "key": "achievement_transform",
        "title": "科技成果转化组织实施制度",
        "short": "成果转化",
        "guide": "重点写清成果筛选、评审、试制、产品应用、转化效果评价、销售或应用证明材料归档。",
    },
    {
        "key": "achievement_reward",
        "title": "科技成果转化激励奖励制度",
        "short": "转化奖励",
        "guide": "重点写清成果转化奖励对象、奖励条件、评价指标、审批流程、发放方式和争议处理。",
    },
    {
        "key": "innovation_platform",
        "title": "开放式创新创业平台管理制度",
        "short": "创新平台",
        "guide": "重点写清平台定位、开放资源、内部孵化、外部协同、项目入驻、资源使用和成果输出管理。",
    },
    {
        "key": "staff_training",
        "title": "科技人员培养进修与职工技能培训制度",
        "short": "人员培训",
        "guide": "重点写清年度培训计划、科技人员进修、岗位技能培训、培训记录、效果评价和经费支持。",
    },
    {
        "key": "talent_performance",
        "title": "优秀人才引进与人才绩效评价奖励制度",
        "short": "人才绩效",
        "guide": "重点写清人才引进条件、招聘流程、绩效评价维度、研发贡献评价、晋升奖励和持续改进机制。",
    },
]

GAOXIN_SYSTEM_BASE_FIELDS = [
    ("company_name", "企业名称"),
    ("company_english_name", "企业英文名称"),
    ("responsible_person", "企业负责人"),
    ("rd_director", "研发负责人"),
    ("finance_director", "财务负责人"),
    ("rd_center_name", "研发机构名称"),
    ("rd_department", "研发/技术部门"),
    ("applicable_year", "适用年度"),
    ("effective_date", "生效日期"),
    ("tech_field", "技术领域"),
    ("main_business", "主营业务"),
    ("main_products", "主要产品或服务"),
    ("main_rd_projects", "主要研发项目"),
    ("evidence_project_code", "研发项目编号"),
    ("evidence_project_name", "佐证材料项目名称"),
    ("evidence_project_source", "项目来源"),
    ("evidence_project_period", "项目周期"),
    ("evidence_budget_amount", "研发预算/经费"),
    ("evidence_project_goal", "研发目标"),
    ("evidence_project_content", "核心研发内容"),
    ("evidence_project_innovation", "关键技术或创新点"),
    ("evidence_acceptance_standard", "验收或考核指标"),
    ("evidence_expected_result", "预期或阶段成果"),
    ("evidence_record_date", "记录/审批日期"),
    ("evidence_participants", "参会/培训/参与人员"),
    ("evidence_archive_no", "归档编号"),
    ("evidence_archive_location", "归档位置"),
    ("generation_notes", "制度正文和佐证文件补充信息"),
    ("staff_total", "职工总数"),
    ("tech_staff", "科技人员数"),
]


GAOXIN_SYSTEM_EVIDENCE_FILE_TEMPLATES = {
    "rd_project": [
        ("研发项目立项申请表", "集中记录项目来源、目标内容、周期预算、预期成果、验收指标和审批意见。"),
    ],
    "rd_accounting": [
        ("研发投入归集审批表", "按研发项目归集主要费用、归集依据、预算金额和财务复核意见。"),
        ("研发投入月度汇总表", "按月汇总各研发项目主要投入情况，支撑年度研发费用统计。"),
        ("研发投入年度复核表", "年度末复核研发投入口径、辅助账和申报数据一致性。"),
    ],
    "rd_auxiliary_ledger": [
        ("研发费用辅助账建账表", "明确辅助账项目、费用科目、建账日期、责任人和财务复核要求。"),
        ("研发费用辅助账月度登记表", "逐月登记项目费用、凭证编号、金额、归集说明和复核状态。"),
        ("研发费用年度结转确认表", "确认年度辅助账结转、差异调整和审计留痕。"),
    ],
    "rd_center": [
        ("研发中心组织架构及岗位职责表", "归档研发机构设置、岗位分工、人员配置和职责边界。"),
        ("研发设备场地台账", "记录研发设备、场地、软件工具、使用部门和维护情况。"),
        ("研发中心例会纪要", "记录研发中心项目协调、问题处理、资源安排和会议决议。"),
        ("研发档案移交登记表", "记录制度、项目、知识产权、试验记录等研发档案移交和保管信息。"),
    ],
    "industry_university_research": [
        ("产学研合作立项评审表", "评审合作单位资质、合作目标、投入资源、成果交付和风险控制。"),
        ("产学研合作协议要点登记表", "登记合同名称、合作内容、知识产权归属、保密条款和验收安排。"),
        ("产学研合作过程记录表", "记录沟通会议、技术交流、阶段交付、问题闭环和责任人。"),
        ("产学研成果验收确认表", "确认合作成果、验收意见、交付资料和后续转化安排。"),
    ],
    "achievement_transform": [
        ("科技成果转化项目登记表", "登记成果来源、转化对象、应用场景、责任部门和转化计划。"),
        ("成果转化评审记录", "记录技术成熟度、应用价值、试制条件、风险和评审结论。"),
        ("试制试用及应用证明表", "记录样品、试用、客户或内部应用、反馈和改进情况。"),
        ("成果转化效果评价表", "汇总转化收入、应用效果、知识产权支撑和归档材料。"),
    ],
    "achievement_reward": [
        ("成果转化奖励申请表", "记录奖励事项、成果贡献、申请人、奖励依据和审批意见。"),
        ("成果贡献评价表", "评价人员在研发、转化、市场应用、知识产权中的贡献。"),
        ("奖励审批发放登记表", "记录奖励标准、审批链、发放方式、凭证和签收情况。"),
        ("成果奖励争议处理记录", "记录异议事项、调查过程、处理意见和归档结论。"),
    ],
    "innovation_platform": [
        ("创新创业平台项目入驻申请表", "记录内部孵化或外部协同项目的入驻条件、资源需求和审批意见。"),
        ("平台资源开放使用登记表", "登记设备、场地、数据、专家等资源开放使用和归还情况。"),
        ("创新平台协同活动记录", "记录技术交流、辅导、路演、测试验证等平台活动。"),
        ("平台项目成果输出确认表", "确认平台项目阶段成果、转化去向和归档附件。"),
    ],
    "staff_training": [
        ("科技人员年度培训计划表", "制定培训主题、对象、时间、预算、组织部门和预期效果。"),
        ("科技人员培养进修审批表", "审批外部进修、专业培训、资格认证和费用支持。"),
        ("职工技能培训签到及记录表", "记录培训时间、地点、讲师、参训人员、课程内容和签到。"),
        ("培训效果评价表", "评价培训结果、岗位应用、改进建议和后续培养安排。"),
    ],
    "talent_performance": [
        ("优秀人才引进审批表", "记录人才需求、岗位条件、面试评价、薪酬建议和审批意见。"),
        ("研发人才绩效评价表", "评价研发任务、项目贡献、成果转化、知识产权和协作表现。"),
        ("人才奖励晋升审批表", "记录奖励晋升依据、评价结果、审批链和执行情况。"),
        ("人才持续培养跟踪表", "跟踪培养目标、阶段表现、培训安排和改进措施。"),
    ],
}


def _system_doc_type_map():
    return {item["key"]: item for item in GAOXIN_SYSTEM_DOC_TYPES}


def _collect_relation_names(data, key):
    rows = ((data.get("gaoxin_relation_table") or {}).get("rows") or []) if isinstance(data, dict) else []
    names = []
    seen = set()
    for row in rows:
        value = str(row.get(key) or "").strip()
        if value and value not in seen:
            seen.add(value)
            names.append(value)
    return "；".join(names[:8])


def _collect_rd_staff_names(data):
    """Return maintained R&D staff names in a stable, de-duplicated order."""
    names = []

    def add(value):
        name = str(value or "").strip()
        if name and name not in names:
            names.append(name)

    for index in range(100):
        add(data.get(f"attachment_rd_staff_{index}_name"))

    hr_staff_rows = data.get("hr_staff_rows") if isinstance(data, dict) else []
    if isinstance(hr_staff_rows, list):
        for row in hr_staff_rows:
            if not isinstance(row, dict):
                continue
            if str(row.get("是否科技人员") or "").strip() == "是":
                add(row.get("姓名"))
    return names


def _collect_december_2025_rd_staff_rows(data):
    """Return effective, named rows from the maintained December 2025 roster."""
    data = data if isinstance(data, dict) else {}
    field_map = {
        "seq": "seq",
        "name": "name",
        "id_no": "id_no",
        "contract": "contract",
        "social_security": "social_security",
        "education": "education",
        "title": "title",
        "is_tech": "is_tech",
    }
    rows = []
    for index in range(100):
        name = str(data.get(f"attachment_rd_staff_{index}_name") or "").strip()
        if not name:
            continue
        row = {}
        for source_key, target_key in field_map.items():
            value = data.get(f"attachment_rd_staff_{index}_{source_key}")
            row[target_key] = str(value or "").strip()
        row["seq"] = row["seq"] or str(len(rows) + 1)
        row["is_tech"] = row["is_tech"] or "是"
        rows.append(row)

    if rows:
        return rows

    hr_staff_rows = data.get("hr_staff_rows")
    if not isinstance(hr_staff_rows, list):
        return []
    hr_field_map = {
        "seq": "序号",
        "name": "姓名",
        "id_no": "身份证号",
        "contract": "是否签订合同",
        "social_security": "是否缴纳社保",
        "education": "学历",
        "title": "职称",
        "is_tech": "是否科技人员",
    }
    for source_row in hr_staff_rows:
        if not isinstance(source_row, dict):
            continue
        name = str(source_row.get("姓名") or "").strip()
        if not name:
            continue
        row = {
            target_key: str(source_row.get(source_key) or "").strip()
            for target_key, source_key in hr_field_map.items()
        }
        row["seq"] = row["seq"] or str(len(rows) + 1)
        row["is_tech"] = row["is_tech"] or "是"
        rows.append(row)
    return rows


def _sync_staff_month_counts(data, rows=None):
    """Keep every 2025 monthly staff count aligned with the December roster."""
    data = data if isinstance(data, dict) else {}
    rows = rows if isinstance(rows, list) else _collect_december_2025_rd_staff_rows(data)
    total_count = len(rows)
    tech_count = sum(1 for row in rows if str(row.get("is_tech") or "是").strip() == "是")
    total_fields = ("start_total", "end_total", "avg_total")
    tech_fields = ("start_tech", "end_tech", "avg_tech")

    for index in range(12):
        data[f"attachment_staff_month_{index}_month"] = f"2025年{index + 1}月"
        for field in total_fields:
            data[f"attachment_staff_month_{index}_{field}"] = total_count
        for field in tech_fields:
            data[f"attachment_staff_month_{index}_{field}"] = tech_count
        data.pop(f"attachment_staff_month_{index}_year_avg_total", None)
        data.pop(f"attachment_staff_month_{index}_year_avg_tech", None)

    # All monthly figures currently follow the December roster, so their
    # twelve-month averages equal the synchronized monthly averages.
    data["attachment_staff_year_avg_total"] = total_count
    data["attachment_staff_year_avg_tech"] = tech_count
    return data


def _build_staff_statement_from_rd_list(data, company_name):
    """Build a factual staff statement from the December 2025 R&D roster."""
    rows = _collect_december_2025_rd_staff_rows(data)
    if not rows:
        return ""

    def status_text(field, yes_label, no_label):
        yes_count = sum(1 for row in rows if row.get(field) == "是")
        no_count = sum(1 for row in rows if row.get(field) == "否")
        blank_count = len(rows) - yes_count - no_count
        if yes_count == len(rows):
            return f"名单所列{len(rows)}人均{yes_label}"
        parts = []
        if yes_count:
            parts.append(f"{yes_label}{yes_count}人")
        if no_count:
            parts.append(f"{no_label}{no_count}人")
        if blank_count:
            parts.append(f"尚未填写{blank_count}人")
        return "、".join(parts)

    def distribution_text(field):
        counts = {}
        for row in rows:
            value = str(row.get(field) or "").strip()
            if value:
                counts[value] = counts.get(value, 0) + 1
        return "、".join(f"{value}{count}人" for value, count in counts.items())

    names = "、".join(row["name"] for row in rows)
    contract_text = status_text("contract", "已签订劳动合同", "未签订劳动合同")
    social_text = status_text("social_security", "已缴纳社会保险", "未缴纳社会保险")
    tech_text = status_text("is_tech", "明确标注为科技人员", "明确标注为非科技人员")
    tech_count = sum(1 for row in rows if row.get("is_tech") == "是")
    tech_ratio = tech_count / len(rows) * 100
    tech_ratio_text = f"{tech_ratio:.2f}".rstrip("0").rstrip(".")
    education_text = distribution_text("education")
    title_text = distribution_text("title")

    detail_parts = []
    if education_text:
        detail_parts.append(f"学历分布为{education_text}")
    if title_text:
        detail_parts.append(f"职称分布为{title_text}")
    if not detail_parts:
        detail_parts.append("学历及职称等未填写信息不作推定")

    company_label = str(company_name or "").strip() or "本公司"
    return "\n".join([
        f"截至2025年12月31日，{company_label}依据《2025年12月份研发人员名单》对研发相关人员情况进行核对，名单共登记{len(rows)}人。",
        f"名单人员为：{names}。",
        f"劳动合同情况：{contract_text}；社会保险情况：{social_text}；科技人员属性：{tech_text}。",
        f"科技人员共{tech_count}人，占总人数的{tech_ratio_text}%。",
        f"{'；'.join(detail_parts)}。上述统计以本附件所列名单为依据，未填写事项以劳动合同、社保记录及人员档案等佐证材料为准。",
    ])


def _rd_project_staff_assignment(staff_names, project_index):
    """Assign a repeatable leader and supporting roles from the R&D staff list."""
    names = [str(name or "").strip() for name in staff_names or [] if str(name or "").strip()]
    if not names:
        return {
            "leader": "待补充",
            "technical": "待补充",
            "testing": "待补充",
            "records": "待补充",
        }

    offset = int(project_index or 0) % len(names)
    rotated = names[offset:] + names[:offset]

    def role_names(start, count=1):
        selected = []
        for step in range(count):
            name = rotated[(start + step) % len(rotated)]
            if name not in selected:
                selected.append(name)
        return "、".join(selected)

    return {
        "leader": rotated[0],
        "technical": role_names(1, min(2, len(rotated))),
        "testing": role_names(3 if len(rotated) > 3 else 1),
        "records": role_names(4 if len(rotated) > 4 else 2 if len(rotated) > 2 else 0),
    }


def _format_rd_project_no(rd_code, project_index):
    """Normalize the relation-table RD sequence for formal project documents."""
    match = re.search(r"\d+", str(rd_code or ""))
    if match:
        return f"RD{int(match.group(0)):02d}"
    return f"RD{int(project_index or 0) + 1:02d}"


def _balanced_cover_title_lines(value, max_line_units=15):
    """Split a project title into visually balanced cover lines."""
    text = re.sub(r"\s+", " ", str(value or "").strip())
    if not text:
        return ["项目名称待补充"]

    def character_units(character):
        if character.isspace():
            return 0.35
        return 0.55 if ord(character) < 128 else 1.0

    total_units = sum(character_units(character) for character in text)
    line_count = max(1, min(3, int((total_units + max_line_units - 0.01) // max_line_units)))
    if line_count == 1:
        return [text]

    closing_punctuation = set("，。；：、）》】〕〉!?！？")
    opening_punctuation = set("《（【〔〈")
    lines = []
    start = 0
    remaining_units = total_units
    for line_index in range(line_count - 1):
        remaining_lines = line_count - line_index
        target_units = remaining_units / remaining_lines
        current_units = 0.0
        split_at = start
        while split_at < len(text):
            next_units = current_units + character_units(text[split_at])
            if current_units and next_units > target_units:
                break
            current_units = next_units
            split_at += 1

        if split_at <= start:
            split_at = start + 1
        while split_at < len(text) and text[split_at] in closing_punctuation:
            current_units += character_units(text[split_at])
            split_at += 1
        while split_at > start + 1 and text[split_at - 1] in opening_punctuation:
            split_at -= 1
            current_units -= character_units(text[split_at])

        line = text[start:split_at].strip()
        if line:
            lines.append(line)
        start = split_at
        remaining_units -= current_units

    final_line = text[start:].strip()
    if final_line:
        lines.append(final_line)
    return lines or [text]


def _collect_rd_project_rows(data):
    rows = ((data.get("gaoxin_relation_table") or {}).get("rows") or []) if isinstance(data, dict) else []
    projects = []
    seen = set()
    rd_staff_names = _collect_rd_staff_names(data)
    for row in rows:
        year = str(row.get("year") or "").strip()
        rd_code = str(row.get("rd_code") or "").strip()
        rd_activity = str(row.get("rd_activity") or "").strip()
        rd_period = str(row.get("rd_period") or "").strip()
        key = "|".join([year, rd_code, rd_activity, rd_period])
        if not key.strip("|") or key in seen:
            continue
        seen.add(key)
        data_index = len(projects)
        for candidate_index in range(100):
            candidate_code = str(data.get(f"rd_{candidate_index}_no") or "").strip()
            candidate_name = str(data.get(f"rd_{candidate_index}_name") or "").strip()
            if (rd_code and candidate_code == rd_code) or (rd_activity and candidate_name == rd_activity):
                data_index = candidate_index
                break

        rd_identity = rd_code or rd_activity
        related_rows = [
            item for item in rows
            if (str(item.get("rd_code") or item.get("rd_activity") or "").strip() == rd_identity)
        ]

        def unique_values(*field_names):
            values = []
            for related in related_rows:
                for field_name in field_names:
                    value = str(related.get(field_name) or "").strip()
                    if value and value not in values:
                        values.append(value)
            return values

        ip_labels = []
        ps_labels = []
        result_names = unique_values("result_name")
        technologies = unique_values("technology", "sales_contract_keywords")
        for related in related_rows:
            ip_label = _relation_label(related.get("ip_code", ""), related.get("ip_name", ""))
            if ip_label and ip_label not in ip_labels:
                ip_labels.append(ip_label)
            ps_label = _relation_label(related.get("ps_code", ""), related.get("ps_name", ""))
            if ps_label and ps_label not in ps_labels:
                ps_labels.append(ps_label)

        prefix = f"rd_{data_index}_"
        temporal = project_temporal_context(rd_period)
        result_text = enforce_temporal_wording(data.get(f"{prefix}result", ""), temporal)
        budget = str(data.get(f"{prefix}budget") or data.get(f"{prefix}total") or "").strip()
        budget_display = budget if not budget or "万" in budget else f"{budget}万元"
        staff_assignment = _rd_project_staff_assignment(rd_staff_names, len(projects))
        projects.append({
            "index": len(projects),
            "data_index": data_index,
            "year": year,
            # The printed project number is always the RD sequence maintained
            # in the RD-IP-PS achievement relation table.
            "project_no": _format_rd_project_no(rd_code, len(projects)),
            "rd_code": rd_code,
            "rd_activity": rd_activity,
            "cover_title_lines": _balanced_cover_title_lines(rd_activity),
            "rd_period": rd_period,
            "purpose": str(data.get(f"{prefix}purpose") or "").strip(),
            "innovation": str(data.get(f"{prefix}innovation") or "").strip(),
            "result": str(result_text or "").strip(),
            "field": str(data.get(f"{prefix}field") or "").strip(),
            "source": str(data.get(f"{prefix}source") or "自主研发").strip(),
            "budget": budget,
            "budget_display": budget_display or "待补充",
            "organization": str(data.get(f"{prefix}org_method") or "").strip(),
            "leader": staff_assignment["leader"],
            "staff_names": rd_staff_names,
            "staff_assignment": staff_assignment,
            "ip_labels": ip_labels or _as_text_list(data.get(f"{prefix}ip_no")),
            "ps_labels": ps_labels or _as_text_list(data.get(f"{prefix}ps_no")),
            "result_names": result_names,
            "technologies": technologies,
            "temporal": temporal,
            "status": temporal["status"],
            "status_display": temporal["status_display"],
            "stages": temporal["stages"],
        })
    return projects


def _as_text_list(value):
    if isinstance(value, list):
        items = value
    else:
        items = re.split(r"[；;、,，\n]+", str(value or ""))
    result = []
    for item in items:
        text = str(item or "").strip()
        if text and text not in result:
            result.append(text)
    return result


def _collect_hitech_product_rows(data):
    relation_table = (data.get("gaoxin_relation_table") or {}) if isinstance(data, dict) else {}
    rows = relation_table.get("rows") or []
    tech_field_path = str(relation_table.get("tech_field_path") or data.get("tech_field") or "").strip()
    products = []
    product_map = {}

    for row in rows:
        ps_code = str(row.get("ps_code") or "").strip()
        ps_name = str(row.get("ps_name") or "").strip()
        key = ps_code or ps_name
        if not key:
            continue
        product = product_map.get(key)
        if not product:
            product = {
                "index": len(products),
                "ps_code": ps_code or f"PS{str(len(products) + 1).zfill(2)}",
                "ps_name": ps_name,
                "field": tech_field_path,
                "revenue": "",
                "source": "",
                "tech": "",
                "advantage": "",
                "ip_support": "",
                "rds": [],
                "ips": [],
                "technologies": [],
                "results": [],
                "proofs": ["专利证书", "销售合同", "销售发票", "产品照片", "检测报告/企业质检报告", "用户使用报告"],
            }
            product_map[key] = product
            products.append(product)

        rd_label = _relation_label(row.get("rd_code", ""), row.get("rd_activity", ""))
        if rd_label and rd_label not in product["rds"]:
            product["rds"].append(rd_label)

        ip_label = _relation_label(row.get("ip_code", ""), row.get("ip_name", ""), 15)
        if ip_label:
            ip_item = {
                "code": str(row.get("ip_code") or "").strip(),
                "name": str(row.get("ip_name") or "").strip(),
                "auth_no": str(row.get("ip_auth_no") or "").strip(),
                "label": ip_label,
            }
            if all(existing.get("label") != ip_label for existing in product["ips"]):
                product["ips"].append(ip_item)

        technology = str(row.get("sales_contract_keywords") or row.get("sales_contract_summary") or row.get("technology") or "").strip()
        if technology and technology not in product["technologies"]:
            product["technologies"].append(technology)
        result_name = str(row.get("result_name") or "").strip()
        if result_name and result_name not in product["results"]:
            product["results"].append(result_name)

    for index in range(20):
        ps_name = str(data.get(f"ps_{index}_name") or "").strip()
        if not ps_name:
            continue
        ps_code = str(data.get(f"ps_{index}_no") or f"PS{str(index + 1).zfill(2)}").strip()
        key = ps_code or ps_name
        if key in product_map or ps_name in product_map:
            continue
        product = {
            "index": len(products),
            "ps_code": ps_code,
            "ps_name": ps_name,
            "field": str(data.get(f"ps_{index}_field") or tech_field_path or "").strip(),
            "revenue": str(data.get(f"ps_{index}_revenue") or "").strip(),
            "source": str(data.get(f"ps_{index}_source") or "").strip(),
            "tech": str(data.get(f"ps_{index}_tech") or "").strip(),
            "advantage": str(data.get(f"ps_{index}_advantage") or "").strip(),
            "ip_support": str(data.get(f"ps_{index}_ip_support") or "").strip(),
            "rds": _as_text_list(data.get(f"ps_{index}_rds")),
            "ips": [{"code": "", "name": ip, "auth_no": "", "label": ip} for ip in _as_text_list(data.get(f"ps_{index}_ip_no"))],
            "technologies": [],
            "results": [],
            "proofs": ["专利证书", "销售合同", "销售发票", "产品照片", "检测报告/企业质检报告", "用户使用报告"],
        }
        product_map[key] = product
        products.append(product)

    for index, product in enumerate(products):
        product["field"] = product.get("field") or data.get(f"ps_{index}_field") or tech_field_path
        product["revenue"] = str(data.get(f"ps_{index}_revenue") or "").strip()
        product["revenue_year"] = GAOXIN_FINANCE_YEARS[-1]
        product["source"] = str(data.get(f"ps_{index}_source") or "").strip()
        product["tech"] = str(data.get(f"ps_{index}_tech") or "").strip()
        product["advantage"] = str(data.get(f"ps_{index}_advantage") or "").strip()
        product["ip_support"] = str(data.get(f"ps_{index}_ip_support") or "").strip()
        saved_rds = _as_text_list(data.get(f"ps_{index}_rds"))
        saved_ips = _as_text_list(data.get(f"ps_{index}_ip_no"))
        for rd in saved_rds:
            if rd not in product["rds"]:
                product["rds"].append(rd)
        for ip in saved_ips:
            if all(existing.get("label") != ip for existing in product["ips"]):
                product["ips"].append({"code": "", "name": ip, "auth_no": "", "label": ip})
        product["ps_kind"] = infer_ps_kind(
            product.get("ps_name"),
            data.get(f"ps_{index}_type"),
        )
        product["type_label"] = ps_type_label(
            product.get("ps_name"),
            product["ps_kind"],
        )
        product["subject_label"] = f"本{product['type_label']}"
        for key in ["source", "tech", "advantage", "ip_support"]:
            product[key] = normalize_ps_reference_text(
                product.get(key),
                product.get("ps_name"),
                product["ps_kind"],
            )
        if product["ps_kind"] == "service":
            product["proofs"] = [
                "专利证书",
                "销售合同",
                "销售发票",
                "服务成果材料",
                "检测报告/企业质检报告",
                "用户使用报告",
            ]
    return products


def _hitech_product_context_text(products):
    lines = []
    for product in products:
        ip_lines = []
        for ip in product.get("ips", []):
            auth = f"，授权号/登记号：{ip.get('auth_no')}" if ip.get("auth_no") else ""
            ip_lines.append(f"{ip.get('label')}{auth}")
        lines.append(
            "\n".join([
                f"{product.get('type_label')}编号：{product.get('ps_code')}",
                f"{product.get('type_label')}名称：{product.get('ps_name')}",
                f"PS类型：{product.get('type_label')}",
                f"技术领域：{product.get('field')}",
                f"上年度销售收入：{product.get('revenue') or '未填写'}万元",
                f"关联RD：{'；'.join(product.get('rds') or [])}",
                f"关联知识产权：{'；'.join(ip_lines)}",
                f"核心技术/成果：{'；'.join((product.get('technologies') or []) + (product.get('results') or []))}",
                f"证明材料：{'；'.join(product.get('proofs') or [])}",
            ])
        )
    return "\n\n".join(lines) if lines else "暂无PS数据。"


def _hitech_product_summary_labels(products):
    kinds = {
        product.get("ps_kind") or infer_ps_kind(product.get("ps_name"))
        for product in products or []
        if str(product.get("ps_name") or "").strip()
    }
    if kinds == {"service"}:
        label = "服务"
        ps_kind = "service"
    elif kinds == {"product"}:
        label = "产品"
        ps_kind = "product"
    else:
        label = "产品及服务"
        ps_kind = "mixed"
    return {
        "ps_kind": ps_kind,
        "type_label": label,
        "title": f"高新技术{label}汇总表",
        "count_label": f"{label}数量" if label in {"产品", "服务"} else "PS数量",
        "name_label": f"高新技术{label}名称" if label in {"产品", "服务"} else "PS名称",
        "data_label": f"{label}基础数据" if label in {"产品", "服务"} else "PS基础数据",
    }


def _normalize_generated_ps_text(text, products):
    text = str(text or "")
    if not text:
        return text
    kinds = {
        product.get("ps_kind") or infer_ps_kind(product.get("ps_name"))
        for product in products or []
        if str(product.get("ps_name") or "").strip()
    }
    if len(kinds) == 1:
        product = next(
            (
                item
                for item in products
                if str(item.get("ps_name") or "").strip()
            ),
            {},
        )
        return normalize_ps_reference_text(
            text,
            product.get("ps_name"),
            next(iter(kinds)),
        )
    combined = r"产品\s*(?:[（(]\s*服务\s*[）)]|[/／]\s*服务|或\s*服务)"
    return re.sub(combined, "产品及服务", text)


def _normalize_generated_ps_structure(value, products):
    if isinstance(value, str):
        return _normalize_generated_ps_text(value, products)
    if isinstance(value, dict):
        return {
            key: _normalize_generated_ps_structure(item, products)
            for key, item in value.items()
        }
    if isinstance(value, list):
        return [_normalize_generated_ps_structure(item, products) for item in value]
    if isinstance(value, tuple):
        return tuple(_normalize_generated_ps_structure(item, products) for item in value)
    return value


def _normalize_system_generated_text(text, products):
    text = _normalize_generated_ps_text(text, products)
    text = text.replace(
        "验收通过后",
        "完成验收程序并形成实际验收记录后",
    )
    text = re.sub(
        r"归档编号可按公司档案规则执行[，,]?如[：:]?\s*\d+",
        "归档编号按公司档案规则填写，未提供时填写“待填写”",
        text,
    )
    text = text.replace("2023100023", "待填写")
    return text


def _natural_sort_key(value):
    parts = re.split(r"(\d+)", str(value or ""))
    return [int(part) if part.isdigit() else part for part in parts]


def _truthy_record_flag(value):
    return str(value or "").strip().lower() in {
        "1",
        "true",
        "yes",
        "y",
        "是",
        "已提供",
        "已核实",
    }


def _achievement_transformation_record_supported(data, achievement_index):
    """Only an explicit verification flag can unlock completed transformation claims."""
    keys = [
        f"cv_{achievement_index}_transformation_record_supported",
        f"cv_{achievement_index}_record_supported",
        f"achievement_{achievement_index}_transformation_verified",
    ]
    return any(_truthy_record_flag((data or {}).get(key)) for key in keys)


def _neutral_achievement_description(achievement):
    """Build a complete description without treating a planned end date as proof."""
    temporal = achievement.get("temporal") or project_temporal_context(
        achievement.get("period")
    )
    status_display = temporal.get("status_display") or "待补充"
    type_label = achievement.get("type_label") or ps_type_label(
        achievement.get("ps_name"),
        achievement.get("ps_kind"),
    )
    rd_label = _relation_label(
        achievement.get("rd_code", ""),
        achievement.get("rd_name", ""),
    ) or "待补充"
    ip_label = _relation_label(
        achievement.get("ip_code", ""),
        achievement.get("ip_name", ""),
    ) or "待补充"
    if achievement.get("ip_auth_no"):
        ip_label = f"{ip_label}（授权号/登记号：{achievement['ip_auth_no']}）"
    ps_label = achievement.get("ps") or _relation_label(
        "",
        achievement.get("ps_name", ""),
    ) or "待补充"

    technology = enforce_temporal_wording(
        achievement.get("technology") or "",
        temporal,
    ).strip()
    if technology:
        technical_note = (
            f"该成果与上述RD、IP及{type_label}建立资料对应关系，"
            f"现有资料登记的技术内容包括：{technology}。"
        )
    else:
        technical_note = (
            f"该成果与上述RD、IP及{type_label}建立资料对应关系，"
            "具体技术内容待依据研发任务书、过程记录、测试资料和成果文件补充核实。"
        )

    if temporal.get("status") == "计划中":
        status_note = (
            "项目尚未到计划开始时间，当前仅记录成果规划和资料对应关系，"
            "不得据此认定成果已经形成、转化或投入应用。"
        )
    elif temporal.get("status") == "研发中":
        status_note = (
            "项目处于研发周期内，当前仅记录研发进展、阶段材料和后续转化安排，"
            "不得据此认定已经完成成果转化。"
        )
    elif temporal.get("status") == "已完成":
        status_note = (
            "项目已到计划结束时间，但计划日期不等于实际完成日期，"
            "不得仅依据计划周期对成果形成、验收结论、转化实施或实际应用作出完成式认定。"
        )
    else:
        status_note = (
            "项目周期信息不完整，当前仅保留资料对应关系，"
            "实际成果形成、转化和应用状态待补充日期及证明材料后核实。"
        )

    lines = [
        f"科技成果名称：{achievement.get('result_name') or '待补充'}",
        f"项目时间状态：{status_display}",
        f"关联研发项目：{rd_label}",
        f"计划周期：{achievement.get('period') or '待补充'}",
        f"关联知识产权：{ip_label}",
        f"关联{type_label}：{ps_label}",
        f"技术关联说明：{technical_note}",
        f"时间及状态说明：{status_note}",
        (
            "核验说明：实际成果形成、转化实施、应用效果和经济效益，应根据研发记录、"
            "验收记录、合同、发票、检测报告、用户使用报告及归档材料逐项核实。"
        ),
        (
            "成果转化核验材料：研发过程及成果文件、知识产权证明、验收或测试记录、"
            "销售合同及发票、应用或用户反馈材料，以及能够证明实际转化日期和效果的其他资料。"
        ),
    ]
    return "\n".join(lines)


def _collect_achievement_transform_rows(data):
    relation = (((data or {}).get("gaoxin_relation_table") or {}).get("rows") or []) if isinstance(data, dict) else []
    achievements = []
    seen = set()
    for row in relation:
        result_no = str(row.get("result_no") or "").strip()
        result_name = str(row.get("result_name") or "").strip()
        key = result_no or result_name
        if not key or key in seen:
            continue
        seen.add(key)
        period = str(row.get("rd_period") or row.get("year") or "").strip()
        temporal = project_temporal_context(period)
        achievements.append({
            "index": len(achievements),
            "result_no": result_no or f"成果{str(len(achievements) + 1).zfill(2)}",
            "result_name": result_name,
            "rd_code": str(row.get("rd_code") or "").strip(),
            "rd_name": str(row.get("rd_activity") or "").strip(),
            "ps": _relation_label(row.get("ps_code", ""), row.get("ps_name", "")),
            "ps_name": str(row.get("ps_name") or "").strip(),
            "ps_kind": infer_ps_kind(row.get("ps_name")),
            "type_label": ps_type_label(row.get("ps_name")),
            "period": period,
            "temporal": temporal,
            "status": temporal["status"],
            "ip": _relation_label(row.get("ip_code", ""), row.get("ip_name", ""), 15),
            "ip_code": str(row.get("ip_code") or "").strip(),
            "ip_name": str(row.get("ip_name") or "").strip(),
            "ip_auth_no": str(row.get("ip_auth_no") or "").strip(),
            "technology": str(row.get("sales_contract_keywords") or row.get("sales_contract_summary") or row.get("technology") or "").strip(),
            "desc": "",
        })

    achievements.sort(key=lambda item: _natural_sort_key(item.get("result_no")))
    for index, achievement in enumerate(achievements):
        achievement["index"] = index
        record_supported = _achievement_transformation_record_supported(data, index)
        achievement["transformation_record_supported"] = record_supported
        saved_desc = data.get(f"cv_{index}_desc") or ""
        achievement["desc"] = (
            enforce_transformation_wording(
                saved_desc,
                achievement["temporal"],
                record_supported=True,
            ).strip()
            if record_supported
            else _neutral_achievement_description(achievement)
        )
        achievement["desc"] = normalize_ps_reference_text(
            achievement["desc"],
            achievement.get("ps_name") or achievement.get("ps"),
            achievement.get("ps_kind"),
        )
        if not achievement["result_name"]:
            achievement["result_name"] = str(data.get(f"cv_{index}_result_name") or "").strip()
    return achievements


def _achievement_evidence_field_name(achievement_index, evidence_key):
    return f"achievement_evidence_{achievement_index}_{evidence_key}"


def _achievement_evidence_edit_map():
    return {item["key"]: item for item in GAOXIN_ACHIEVEMENT_EVIDENCE_EDIT_TYPES}


def _achievement_user_report_fields(content_text, company, achievement):
    labels = {
        "使用单位": "customer_name",
        "被评价单位": "company_name",
        "产品（服务）名称": "product_name",
        "产品(服务)名称": "product_name",
        "产品名称": "product_name",
        "服务名称": "product_name",
        "使用时间": "use_period",
        "应用场景": "use_scene",
        "功能适用性": "function_fit",
        "运行稳定性": "stability",
        "操作便利性": "usability",
        "服务响应": "service_response",
        "综合评价": "overall_evaluation",
        "客户意见及建议": "suggestion",
        "使用单位（盖章）": "customer_seal",
        "使用单位(盖章)": "customer_seal",
        "日期": "report_date",
    }
    fields = {key: "" for key in set(labels.values())}
    current_key = ""
    label_pattern = "|".join(re.escape(label) for label in sorted(labels, key=len, reverse=True))
    for raw_line in str(content_text or "").replace("\r", "").splitlines():
        line = raw_line.strip()
        if not line or line == "用户使用报告":
            continue
        match = re.match(rf"^(?:{label_pattern})\s*[：:]\s*(.*)$", line)
        if match:
            label = re.match(rf"^({label_pattern})", line).group(1)
            current_key = labels[label]
            fields[current_key] = match.group(1).strip()
        elif current_key:
            fields[current_key] = f"{fields[current_key]} {line}".strip()

    compact_content = re.sub(r"\s+", " ", str(content_text or "")).strip()
    if compact_content and not any(fields.values()):
        sentences = [item.strip() for item in re.split(r"(?<=[。！？])", compact_content) if item.strip()]

        def pick_sentence(*keywords):
            for sentence in sentences:
                if any(keyword in sentence for keyword in keywords):
                    return sentence[:120]
            return ""

        fields["use_scene"] = pick_sentence("应用场景", "用户现场", "实际业务", "使用场景")
        fields["function_fit"] = pick_sentence("满足", "适用", "使用需求", "应用需要")
        fields["stability"] = pick_sentence("稳定", "可靠")
        fields["usability"] = pick_sentence("便利", "方便", "效率")
        fields["service_response"] = pick_sentence("服务响应", "技术支持", "技术人员")
        fields["suggestion"] = pick_sentence("建议", "改进")
        fields["overall_evaluation"] = pick_sentence("综合评价", "整体评价", "评价认为") or compact_content[:120]

    for key, value in fields.items():
        normalized = str(value or "").strip()
        if re.fullmatch(r"[_＿\s]+", normalized) or re.fullmatch(r"[年月日_＿\s]+", normalized):
            fields[key] = ""
    fields["company_name"] = fields["company_name"] or company.name
    fields["product_name"] = fields["product_name"] or achievement.get("ps") or achievement.get("result_name") or ""
    return fields


def _achievement_test_report_fields(content_text, company, achievement):
    labels = {
        "检测单位": "company_name",
        "检测对象": "product_name",
        "检测对象（产品/服务）": "product_name",
        "检测对象(产品/服务)": "product_name",
        "被检测技术": "technology_name",
        "检测技术": "technology_name",
        "检测目的": "test_purpose",
        "检测依据": "test_basis",
        "检测方法": "test_method",
        "检测项目1": "item_1",
        "检测项目一": "item_1",
        "检测结果1": "result_1",
        "检测结果一": "result_1",
        "检测项目2": "item_2",
        "检测项目二": "item_2",
        "检测结果2": "result_2",
        "检测结果二": "result_2",
        "检测项目3": "item_3",
        "检测项目三": "item_3",
        "检测结果3": "result_3",
        "检测结果三": "result_3",
        "检测结论": "test_conclusion",
        "检测单位（盖章）": "company_seal",
        "检测单位(盖章)": "company_seal",
        "日期": "report_date",
    }
    fields = {key: "" for key in set(labels.values())}
    current_key = ""
    label_pattern = "|".join(re.escape(label) for label in sorted(labels, key=len, reverse=True))
    for raw_line in str(content_text or "").replace("\r", "").splitlines():
        line = raw_line.strip()
        if not line or line == "检测报告":
            continue
        match = re.match(rf"^(?:{label_pattern})\s*[：:]\s*(.*)$", line)
        if match:
            label = re.match(rf"^({label_pattern})", line).group(1)
            current_key = labels[label]
            fields[current_key] = match.group(1).strip()
        elif current_key:
            fields[current_key] = f"{fields[current_key]} {line}".strip()

    for key, value in fields.items():
        normalized = str(value or "").strip()
        if re.fullmatch(r"[_＿\s]+", normalized) or re.fullmatch(r"[年月日_＿\s]+", normalized):
            fields[key] = ""
    fields["company_name"] = fields["company_name"] or company.name
    fields["product_name"] = fields["product_name"] or achievement.get("ps") or achievement.get("result_name") or ""
    fields["technology_name"] = fields["technology_name"] or achievement.get("result_name") or ""
    return fields


def _achievement_evidence_upload_map():
    return {item["key"]: item for item in GAOXIN_ACHIEVEMENT_EVIDENCE_UPLOAD_TYPES}


def _achievement_evidence_files(attachments, achievement_index, evidence_key):
    files = attachments.get("achievement_transform", {}).get("files", []) if isinstance(attachments, dict) else []
    index_text = str(achievement_index)
    return [
        file for file in files
        if str(file.get("achievement_index") or "") == index_text
        and str(file.get("achievement_evidence_type") or "") == evidence_key
    ]


def _hitech_product_evidence_map():
    return {item["key"]: item for item in GAOXIN_HITECH_PRODUCT_EVIDENCE_TYPES}


def _hitech_product_evidence_files(attachments, product_index, evidence_key=None):
    files = attachments.get("hitech_products", {}).get("files", []) if isinstance(attachments, dict) else []
    index_text = str(product_index)
    return [
        file for file in files
        if str(file.get("product_index") or "") == index_text
        and (evidence_key is None or str(file.get("product_evidence_type") or "") == evidence_key)
    ]


def _sync_achievement_patent_cert_files(company, data, achievement_index, achievement):
    attachments = _load_gaoxin_attachments_from_data(data)
    ip_files = attachments.get("ip", {}).get("files", [])
    target_section = attachments.setdefault("achievement_transform", {"files": []})
    target_files = target_section.setdefault("files", [])
    index_text = str(achievement_index)
    existing_source_ids = {
        str(file.get("source_ip_file_id") or "")
        for file in target_files
        if str(file.get("achievement_index") or "") == index_text
        and str(file.get("achievement_evidence_type") or "") == "patent_cert"
    }
    ip_code = str(achievement.get("ip_code") or "").strip()
    ip_name = str(achievement.get("ip_name") or "").strip()
    ip_auth_no = str(achievement.get("ip_auth_no") or "").strip()
    fallback_summary, fallback_rows = _build_ip_attachment_rows(_load_ip_details(company))
    ip_rows = _ip_attachment_rows_from_data(data, fallback_rows)
    matched_row_cert_ids = set()
    matched_row_seqs = set()
    for row in ip_rows:
        seq = str(row.get("seq") or "").strip()
        name = str(row.get("name") or "").strip()
        patent_no = str(row.get("patent_no") or "").strip()
        if (ip_code and seq == ip_code) or (ip_name and name == ip_name) or (ip_auth_no and patent_no == ip_auth_no):
            if row.get("cert_id"):
                matched_row_cert_ids.add(str(row.get("cert_id")))
            if seq:
                matched_row_seqs.add(seq)

    matched_ip_files = []
    seen_file_ids = set()
    for file in ip_files:
        if not isinstance(file, dict):
            continue
        source_id = str(file.get("id") or "")
        if not source_id or source_id in seen_file_ids:
            continue
        file_seq = str(file.get("ip_seq") or "").strip()
        file_cert_id = str(file.get("ip_cert_id") or "").strip()
        filename = str(file.get("original_filename") or "").strip()
        if (
            (ip_code and file_seq == ip_code)
            or (file_seq and file_seq in matched_row_seqs)
            or (file_cert_id and file_cert_id in matched_row_cert_ids)
            or (ip_auth_no and ip_auth_no in filename)
            or (ip_name and ip_name in filename)
        ):
            matched_ip_files.append(file)
            seen_file_ids.add(source_id)

    changed = False
    for source_file in matched_ip_files:
        source_id = str(source_file.get("id") or "")
        if not source_id or source_id in existing_source_ids:
            continue
        relative_source_path = source_file.get("relative_path")
        if not relative_source_path:
            continue
        source_path = _safe_attachment_path(relative_source_path)
        if not os.path.exists(source_path):
            continue
        safe_name = secure_filename(source_file.get("original_filename") or source_file.get("stored_filename") or "patent_certificate.pdf") or f"patent_{uuid.uuid4().hex}.pdf"
        if not safe_name.lower().endswith(".pdf"):
            safe_name = f"{os.path.splitext(safe_name)[0]}.pdf"
        file_id = uuid.uuid4().hex
        stored_filename = f"{file_id}_{safe_name}"
        relative_path = _attachment_relative_path(company.user_id, company.id, "achievement_transform", stored_filename)
        target_path = _safe_attachment_path(relative_path)
        os.makedirs(os.path.dirname(target_path), exist_ok=True)
        shutil.copy2(source_path, target_path)
        persist_file(target_path, relative_path)
        target_files.append({
            "id": file_id,
            "original_filename": source_file.get("original_filename") or safe_name,
            "stored_filename": stored_filename,
            "relative_path": relative_path,
            "uploaded_at": datetime.utcnow().isoformat(timespec="seconds"),
            "achievement_index": index_text,
            "achievement_evidence_type": "patent_cert",
            "source": "achievement_ip_auto_sync",
            "auto_synced": True,
            "source_ip_file_id": source_id,
            "source_ip_seq": source_file.get("ip_seq", ""),
            "source_ip_cert_id": source_file.get("ip_cert_id", ""),
        })
        existing_source_ids.add(source_id)
        changed = True

    if changed:
        data["gaoxin_attachments"] = attachments
        company.data_json = json.dumps(data, ensure_ascii=False)
        db.session.commit()
    return attachments


def _render_html_pdf(html, download_name, redirect_endpoint, **redirect_values):
    render_started = time.perf_counter()
    company_name = str(redirect_values.pop("_header_company_name", "") or "").strip()
    company_english_name = str(redirect_values.pop("_header_company_english_name", "") or "").strip()
    header_skip_first_page = bool(redirect_values.pop("_header_skip_first_page", False))
    try:
        with tempfile.TemporaryDirectory(prefix="html-pdf-") as output_dir:
            pdf_path = os.path.join(output_dir, "document.pdf")
            _render_pdf_file(
                current_app._get_current_object(),
                html,
                pdf_path,
                download_name,
            )
            if company_name or company_english_name:
                _stamp_pdf_file_headers(
                    pdf_path,
                    company_name,
                    company_english_name,
                    skip_first_page=header_skip_first_page,
                )
            pdf_bytes = Path(pdf_path).read_bytes()
        duration = time.perf_counter() - render_started
        current_app.logger.info(
            "%s PDF 生成完成 duration=%.3fs bytes=%s",
            download_name,
            duration,
            len(pdf_bytes),
        )
        response = send_file(
            BytesIO(pdf_bytes),
            mimetype="application/pdf",
            as_attachment=False,
            download_name=download_name,
        )
        response.headers["X-PDF-Render-Duration"] = f"{duration:.3f}"
        return response
    except Exception as exc:
        current_app.logger.exception(
            "%s PDF 生成失败 duration=%.3fs",
            download_name,
            time.perf_counter() - render_started,
        )
        flash(f"PDF 生成失败：{exc}", "danger")
        return redirect(url_for(redirect_endpoint, **redirect_values))


def _pdf_cjk_font_path():
    configured = str(os.getenv("PDF_CJK_FONT") or "").strip()
    return configured if configured and os.path.isfile(configured) else ""


def _stamp_generated_pdf_pages(document, page_indexes, company_name, company_english_name):
    """Replace flow headers and stamp a true bilingual header on generated pages."""
    try:
        import fitz
    except ImportError:
        import pymupdf as fitz

    chinese_name = str(company_name or "").strip()
    english_name = str(company_english_name or "").strip()
    if not chinese_name and not english_name:
        return

    text_color = tuple(value / 255 for value in (55, 65, 81))
    line_color = tuple(value / 255 for value in (156, 163, 175))
    mm = 72 / 25.4
    text_top = 12.25 * mm
    text_bottom = 20.25 * mm
    line_y = 22 * mm
    header_text = " | ".join(part for part in [chinese_name, english_name] if part)
    cjk_font_path = _pdf_cjk_font_path()
    header_font = (
        fitz.Font(fontfile=cjk_font_path)
        if cjk_font_path
        else fitz.Font("china-s")
    )
    latin_font = fitz.Font("helv")

    overlays = {}

    def overlay_for(page):
        page_key = (round(page.rect.width, 2), round(page.rect.height, 2))
        if page_key in overlays:
            return overlays[page_key]
        overlay_document = fitz.open()
        overlay_page = overlay_document.new_page(
            width=page.rect.width,
            height=page.rect.height,
        )
        header_font_name = "generated-header-cjk" if cjk_font_path else "china-s"
        if cjk_font_path:
            overlay_page.insert_font(
                fontname=header_font_name,
                fontfile=cjk_font_path,
            )
        landscape = page.rect.width > page.rect.height
        horizontal_margin = (12 if landscape else 14) * mm
        available_width = page.rect.width - (2 * horizontal_margin)
        preferred_font_size = 7.5
        header_lines = []
        if chinese_name and english_name:
            line_specs = (
                (chinese_name, header_font, header_font_name),
                (english_name, latin_font, "helv"),
            )
        elif chinese_name:
            line_specs = ((chinese_name, header_font, header_font_name),)
        else:
            line_specs = ((header_text, latin_font, "helv"),)
        for text, line_font, line_font_name in line_specs:
            text_width = line_font.text_length(
                text,
                fontsize=preferred_font_size,
            )
            font_size = preferred_font_size
            if text_width > available_width:
                font_size *= available_width / text_width
            header_lines.append(
                (text, max(font_size * 0.995, 1.0), line_font, line_font_name)
            )

        line_heights = [
            (line_font.ascender - line_font.descender) * font_size
            for _, font_size, line_font, _ in header_lines
        ]
        total_text_height = sum(line_heights)
        current_y = text_top + max(
            0,
            ((text_bottom - text_top) - total_text_height) / 2,
        )
        for (
            text,
            font_size,
            line_font,
            line_font_name,
        ), line_height in zip(header_lines, line_heights):
            text_width = line_font.text_length(text, fontsize=font_size)
            text_x = horizontal_margin + max(
                0,
                (available_width - text_width) / 2,
            )
            baseline_y = current_y + (line_font.ascender * font_size)
            overlay_page.insert_text(
                fitz.Point(text_x, baseline_y),
                text,
                fontname=line_font_name,
                fontsize=font_size,
                color=text_color,
                overlay=True,
            )
            current_y += line_height
        overlay_page.draw_line(
            fitz.Point(horizontal_margin, line_y),
            fitz.Point(page.rect.width - horizontal_margin, line_y),
            color=line_color,
            width=0.5,
        )
        if cjk_font_path:
            overlay_document.subset_fonts()
        overlays[page_key] = overlay_document
        return overlay_document

    try:
        for page_index in sorted(set(int(index) for index in page_indexes)):
            if page_index < 0 or page_index >= document.page_count:
                continue
            page = document[page_index]
            has_old_header = False
            for block in page.get_text("blocks"):
                block_text = str(block[4] or "")
                if (
                    (english_name and english_name in block_text)
                    or (chinese_name and chinese_name in block_text)
                ) and block[1] < 70:
                    old_header_rect = fitz.Rect(
                        0,
                        max(0, block[1] - 3),
                        page.rect.width,
                        block[3] + 10,
                    )
                    page.add_redact_annot(old_header_rect, fill=(1, 1, 1))
                    has_old_header = True
            if has_old_header:
                page.apply_redactions()
            page.show_pdf_page(
                page.rect,
                overlay_for(page),
                0,
                overlay=True,
            )
    finally:
        for overlay_document in overlays.values():
            overlay_document.close()


def _stamp_pdf_file_headers(
    pdf_path,
    company_name,
    company_english_name,
    *,
    skip_first_page=False,
):
    try:
        import fitz
    except ImportError:
        import pymupdf as fitz

    source = fitz.open(pdf_path)
    stamped_path = f"{pdf_path}.stamped.pdf"
    try:
        _stamp_generated_pdf_pages(
            source,
            range(1 if skip_first_page else 0, source.page_count),
            company_name,
            company_english_name,
        )
        source.save(stamped_path, garbage=4, deflate=True)
    finally:
        source.close()
    os.replace(stamped_path, pdf_path)


def _chrome_executable(app):
    configured = str(app.config.get("CHROME_BIN") or "").strip()
    candidates = [
        configured,
        "/Applications/Google Chrome.app/Contents/MacOS/Google Chrome",
        shutil.which("google-chrome"),
        shutil.which("google-chrome-stable"),
        shutil.which("chromium"),
        shutil.which("chromium-browser"),
    ]
    for candidate in candidates:
        if candidate and os.path.isfile(candidate) and os.access(candidate, os.X_OK):
            return candidate
    return None


def _css_length_to_points(value):
    match = re.fullmatch(r"\s*([0-9]+(?:\.[0-9]+)?)\s*(mm|cm|in|pt|px)\s*", value, re.IGNORECASE)
    if not match:
        raise ValueError(f"不支持的 PDF 页边距：{value}")
    amount = float(match.group(1))
    unit = match.group(2).lower()
    factors = {
        "mm": 72 / 25.4,
        "cm": 72 / 2.54,
        "in": 72,
        "pt": 1,
        "px": 72 / 96,
    }
    return amount * factors[unit]


def _story_page_layout(html):
    """Read the simple @page rules used by export templates."""
    try:
        import fitz
    except ImportError:
        import pymupdf as fitz

    page_rule_match = re.search(r"@page\s*\{(.*?)\}", html, re.IGNORECASE | re.DOTALL)
    page_rule = page_rule_match.group(1) if page_rule_match else ""
    page_rect = fitz.paper_rect("a4")
    if re.search(r"\bsize\s*:\s*(?:a4\s+)?landscape\b", page_rule, re.IGNORECASE):
        page_rect = fitz.Rect(0, 0, page_rect.height, page_rect.width)

    margins = [28 * 72 / 25.4, 16 * 72 / 25.4, 18 * 72 / 25.4, 16 * 72 / 25.4]
    margin_match = re.search(r"\bmargin\s*:\s*([^;]+)", page_rule, re.IGNORECASE)
    if margin_match:
        values = [
            _css_length_to_points(value)
            for value in re.findall(
                r"[0-9]+(?:\.[0-9]+)?\s*(?:mm|cm|in|pt|px)",
                margin_match.group(1),
                re.IGNORECASE,
            )
        ]
        if len(values) == 1:
            margins = values * 4
        elif len(values) == 2:
            margins = [values[0], values[1], values[0], values[1]]
        elif len(values) == 3:
            margins = [values[0], values[1], values[2], values[1]]
        elif len(values) == 4:
            margins = values

    top, right, bottom, left = margins
    content_rect = fitz.Rect(
        left,
        top,
        page_rect.width - right,
        page_rect.height - bottom,
    )
    if content_rect.width < 72 or content_rect.height < 72:
        raise ValueError("PDF 页边距过大，正文区域不足")
    return page_rect, content_rect


def _prepare_pymupdf_story_html(html, content_width):
    """Add Story-compatible table sizing without changing Chrome output."""
    from lxml import etree
    from lxml import html as lxml_html

    def transparent_png_data_uri(width):
        width = max(1, int(width))
        signature = b"\x89PNG\r\n\x1a\n"

        def chunk(chunk_type, payload):
            checksum = zlib.crc32(chunk_type + payload) & 0xFFFFFFFF
            return (
                struct.pack(">I", len(payload))
                + chunk_type
                + payload
                + struct.pack(">I", checksum)
            )

        header = struct.pack(">IIBBBBB", width, 1, 8, 6, 0, 0, 0)
        transparent_row = b"\x00" + (b"\x00\x00\x00\x00" * width)
        png = (
            signature
            + chunk(b"IHDR", header)
            + chunk(b"IDAT", zlib.compress(transparent_row, 9))
            + chunk(b"IEND", b"")
        )
        return f"data:image/png;base64,{base64.b64encode(png).decode('ascii')}"

    def table_column_weights(table, column_count):
        raw_weights = table.get("data-pymupdf-widths", "")
        if raw_weights:
            try:
                parsed_weights = [
                    float(value.strip())
                    for value in raw_weights.split(",")
                ]
            except (TypeError, ValueError):
                parsed_weights = []
            if (
                len(parsed_weights) == column_count
                and all(weight > 0 for weight in parsed_weights)
            ):
                total_weight = sum(parsed_weights)
                return tuple(weight / total_weight for weight in parsed_weights)

        if column_count == 2:
            return (0.2, 0.8)
        if column_count == 3:
            return (0.2, 0.6, 0.2)
        if column_count == 4:
            return (0.16, 0.34, 0.16, 0.34)
        return (1 / column_count,) * column_count

    def split_long_text(text, max_chars=180):
        normalized = str(text or "").strip()
        if len(normalized) <= max_chars:
            return [normalized]

        segments = []
        for paragraph in re.split(r"\n+", normalized):
            paragraph = paragraph.strip()
            if not paragraph:
                continue
            sentence_parts = re.split(r"(?<=[。！？；])", paragraph)
            segments.extend(part.strip() for part in sentence_parts if part.strip())

        chunks = []
        current = ""
        for segment in segments:
            while len(segment) > max_chars:
                if current:
                    chunks.append(current)
                    current = ""
                chunks.append(segment[:max_chars])
                segment = segment[max_chars:]
            if not segment:
                continue
            if current and len(current) + len(segment) > max_chars:
                chunks.append(current)
                current = segment
            else:
                current += segment
        if current:
            chunks.append(current)
        return chunks or [normalized]

    def split_oversized_table_rows(table):
        for row in list(table.xpath(".//tr[not(ancestor::table[2])]")):
            cells = row.xpath("./th|./td")
            if not cells:
                continue
            cell_texts = [" ".join(cell.text_content().split()) for cell in cells]
            longest_index = max(range(len(cells)), key=lambda index: len(cell_texts[index]))
            chunks = split_long_text(cell_texts[longest_index])
            if len(chunks) <= 1:
                continue

            parent = row.getparent()
            insert_at = parent.index(row)
            for chunk_index, chunk in enumerate(chunks):
                split_row = copy.deepcopy(row)
                split_cells = split_row.xpath("./th|./td")
                for cell_index, cell in enumerate(split_cells):
                    cell.clear()
                    if cell_index == longest_index:
                        cell.text = chunk
                    elif chunk_index == 0:
                        cell.text = cell_texts[cell_index]
                    else:
                        cell.text = ""
                parent.insert(insert_at + chunk_index, split_row)
            parent.remove(row)

    root = lxml_html.document_fromstring(html)
    head_nodes = root.xpath("//head")
    head = head_nodes[0] if head_nodes else etree.SubElement(root, "head")
    compatibility_style = etree.Element("style")
    compatibility_style.text = """
      .pymupdf-table-sizer,
      .pymupdf-table-sizer td {
        border: 0 !important;
        padding: 0 !important;
        height: 0 !important;
        min-height: 0 !important;
        line-height: 0 !important;
        background-color: #ffffff !important;
      }
      .pymupdf-table-sizer img {
        display: block !important;
        margin: 0 !important;
        padding: 0 !important;
      }
      [data-pymupdf-page-break-before] {
        break-before: auto !important;
        page-break-before: auto !important;
      }
      .cover-page {
        break-after: auto !important;
        page-break-after: auto !important;
      }
      .rd-project-document table.cover-code-row {
        margin-bottom: 40mm !important;
      }
      tr {
        break-inside: auto !important;
        page-break-inside: auto !important;
      }
    """
    head.append(compatibility_style)

    # Story ignores percentage, fixed CSS and HTML width attributes on tables.
    # Story honors intrinsic image dimensions even though it ignores percentage
    # and fixed CSS widths. Transparent 1px-high PNGs therefore provide stable
    # column sizing without leaving text or visible marks in the PDF.
    spacer_count = max(1, int(round(float(content_width) / 0.75)))
    for table in root.xpath("//table"):
        split_oversized_table_rows(table)
        rows = table.xpath(".//tr[not(ancestor::table[2])]")
        column_count = 1
        for row in rows:
            logical_columns = 0
            for cell in row.xpath("./th|./td"):
                try:
                    logical_columns += max(1, int(cell.get("colspan") or 1))
                except (TypeError, ValueError):
                    logical_columns += 1
            column_count = max(column_count, logical_columns)

        sizer_row = etree.Element("tr", {"class": "pymupdf-table-sizer"})
        weights = table_column_weights(table, column_count)
        allocated = 0
        for column_index, weight in enumerate(weights):
            remaining_columns = column_count - column_index
            if remaining_columns == 1:
                column_spacers = max(1, spacer_count - allocated)
            else:
                column_spacers = max(1, int(round(spacer_count * weight)))
                allocated += column_spacers
            sizer_cell = etree.SubElement(
                sizer_row,
                "td",
                {"aria-hidden": "true"},
            )
            etree.SubElement(
                sizer_cell,
                "img",
                {
                    "aria-hidden": "true",
                    "alt": "",
                    "data-pymupdf-spacer-width": str(column_spacers),
                    "src": transparent_png_data_uri(column_spacers),
                },
            )

        body_nodes = table.xpath("./tbody")
        if body_nodes:
            body_nodes[0].insert(0, sizer_row)
        else:
            first_row = table.xpath("./tr")
            if first_row:
                table.insert(table.index(first_row[0]), sizer_row)
            else:
                table.append(sizer_row)

    return etree.tostring(root, encoding="unicode", method="html")


def _split_pymupdf_story_documents(html):
    """Split explicitly marked document parts into separate Story inputs."""
    from lxml import etree
    from lxml import html as lxml_html

    root = lxml_html.document_fromstring(html)
    body_nodes = root.xpath("//body")
    if not body_nodes:
        return [html]

    body = body_nodes[0]
    children = list(body)
    groups = []
    current_group = []
    for child in children:
        if child.get("data-pymupdf-page-break-before") is not None and current_group:
            groups.append(current_group)
            current_group = []
        current_group.append(child)
    if current_group:
        groups.append(current_group)
    if len(groups) <= 1:
        return [html]

    head_nodes = root.xpath("//head")
    head_html = (
        etree.tostring(head_nodes[0], encoding="unicode", method="html")
        if head_nodes
        else "<head></head>"
    )
    body_attributes = " ".join(
        f'{name}="{html_escape(value, quote=True)}"'
        for name, value in body.attrib.items()
    )
    body_open = f"<body {body_attributes}>" if body_attributes else "<body>"
    return [
        (
            "<!doctype html><html lang=\"zh-CN\">"
            + head_html
            + body_open
            + "".join(
                etree.tostring(child, encoding="unicode", method="html")
                for child in group
            )
            + "</body></html>"
        )
        for group in groups
    ]


def _remove_pymupdf_story_repeated_backgrounds(document):
    """Remove Story's stale short background fills from continuation pages."""
    number = rb"-?(?:\d+(?:\.\d*)?|\.\d+)"
    rect_operator = re.compile(
        rb"(?P<x>"
        + number
        + rb")\s+(?P<y>"
        + number
        + rb")\s+(?P<width>"
        + number
        + rb")\s+(?P<height>"
        + number
        + rb")\s+re\s+f\s+"
    )
    color_group = re.compile(
        rb"(?P<color>"
        + number
        + rb"\s+"
        + number
        + rb"\s+"
        + number
        + rb"\s+rg\s+)"
        rb"(?P<rects>(?:(?:"
        + number
        + rb"\s+){4}re\s+f\s+)+)"
    )
    known_backgrounds = (
        (233 / 255, 237 / 255, 242 / 255),  # #e9edf2
        (227 / 255, 232 / 255, 238 / 255),  # #e3e8ee
        (225 / 255, 230 / 255, 236 / 255),  # #e1e6ec
        (237 / 255, 240 / 255, 244 / 255),  # #edf0f4
        (223 / 255, 231 / 255, 239 / 255),  # #dfe7ef
    )

    removed_rects = 0
    for page_index in range(1, document.page_count):
        page = document[page_index]
        word_centers = [
            ((word[0] + word[2]) / 2, (word[1] + word[3]) / 2)
            for word in page.get_text("words")
        ]
        for xref in page.get_contents():
            stream = document.xref_stream(xref)
            clip_match = re.search(rb"\bre\s+W\s+n\s+", stream)
            if not clip_match:
                continue

            artifact_start = clip_match.end()
            cursor = artifact_start
            cleaned_groups = []
            while True:
                match = color_group.match(stream, cursor)
                if not match:
                    break
                color_values = tuple(
                    float(value)
                    for value in re.findall(number, match.group("color"))[:3]
                )
                if not any(
                    all(abs(actual - expected) < 0.002 for actual, expected in zip(color_values, target))
                    for target in known_backgrounds
                ):
                    break

                retained_rects = []
                for rect_match in rect_operator.finditer(match.group("rects")):
                    x = float(rect_match.group("x"))
                    y = float(rect_match.group("y"))
                    width = float(rect_match.group("width"))
                    height = float(rect_match.group("height"))
                    left, right = sorted((x, x + width))
                    top, bottom = sorted((y, y + height))
                    is_short_story_fragment = 7.5 <= bottom - top <= 10.5
                    has_centered_text = any(
                        left <= center_x <= right and top <= center_y <= bottom
                        for center_x, center_y in word_centers
                    )
                    if is_short_story_fragment and not has_centered_text:
                        removed_rects += 1
                    else:
                        retained_rects.append(rect_match.group(0))

                if retained_rects:
                    cleaned_groups.append(
                        match.group("color") + b"".join(retained_rects)
                    )
                cursor = match.end()

            if cursor > artifact_start:
                document.update_stream(
                    xref,
                    stream[:artifact_start] + b"".join(cleaned_groups) + stream[cursor:],
                )

    return removed_rects


def _render_pdf_with_pymupdf(html, pdf_path):
    try:
        import fitz
    except ImportError:
        try:
            import pymupdf as fitz
        except ImportError as exc:
            raise RuntimeError("无 Chrome 环境生成 PDF 需要 PyMuPDF") from exc

    page_rect, content_rect = _story_page_layout(html)
    writer = fitz.DocumentWriter(pdf_path)

    def rect_function(_rect_number, _filled):
        return page_rect, content_rect, None

    try:
        for story_document in _split_pymupdf_story_documents(html):
            story_html = _prepare_pymupdf_story_html(
                story_document,
                content_rect.width,
            )
            story = fitz.Story(html=story_html)
            story.write(writer, rect_function)
    finally:
        writer.close()

    optimized_path = f"{pdf_path}.optimized.pdf"
    document = fitz.open(pdf_path)
    try:
        if document.page_count == 0:
            raise RuntimeError("PyMuPDF 未生成任何 PDF 页面")
        _remove_pymupdf_story_repeated_backgrounds(document)
        document.subset_fonts()
        document.save(optimized_path, garbage=4, deflate=True, clean=True)
    finally:
        document.close()
    os.replace(optimized_path, pdf_path)


def _render_pdf_file(app, html, pdf_path, label, render_info=None):
    """Render HTML with Chrome when available, otherwise use bundled PyMuPDF."""
    output_dir = os.path.dirname(pdf_path)
    os.makedirs(output_dir, exist_ok=True)
    chrome = _chrome_executable(app)
    renderer = "chrome"

    if chrome:
        html_path = os.path.join(output_dir, f"{Path(pdf_path).stem}.html")
        Path(html_path).write_text(html, encoding="utf-8")
        timeout = max(15, int(app.config.get("PDF_RENDER_TIMEOUT", 90)))
        command = [
            chrome,
            "--headless",
            "--disable-gpu",
            "--no-sandbox",
            f"--print-to-pdf={pdf_path}",
            "--no-pdf-header-footer",
            html_path,
        ]
        try:
            subprocess.run(command, timeout=timeout, check=True, capture_output=True)
            if not os.path.isfile(pdf_path) or os.path.getsize(pdf_path) == 0:
                raise RuntimeError("Chrome 未生成有效 PDF 文件")
            app.logger.info("%s PDF 使用 Chrome 生成", label)
        except Exception:
            app.logger.warning("%s Chrome PDF 生成失败，改用 PyMuPDF", label, exc_info=True)
            try:
                os.remove(pdf_path)
            except FileNotFoundError:
                pass
            _render_pdf_with_pymupdf(html, pdf_path)
            renderer = "pymupdf"
    else:
        app.logger.info("%s 未找到 Chrome/Chromium，使用 PyMuPDF 生成 PDF", label)
        _render_pdf_with_pymupdf(html, pdf_path)
        renderer = "pymupdf"

    if not os.path.isfile(pdf_path) or os.path.getsize(pdf_path) == 0:
        raise RuntimeError(f"{label} PDF 生成结果为空")
    if render_info is not None:
        render_info["renderer"] = renderer
    return pdf_path


def _validate_export_pdf_layout(pdf_path, label, *, check_shapes=True):
    """Reject generated PDFs whose body content enters the header/footer safety areas."""
    try:
        import fitz
    except ImportError:
        try:
            import pymupdf as fitz
        except ImportError as exc:
            raise RuntimeError("PDF 排版检查不可用：缺少 PyMuPDF") from exc

    # Generated pages reserve 28 mm at the top. The stamped rule sits at 22 mm,
    # leaving roughly 6 mm before normal body content begins.
    header_limit = 75
    header_text_limit = 45
    footer_limit = 32
    issues = []
    document = fitz.open(pdf_path)
    try:
        for page_index, page in enumerate(document, start=1):
            page_height = page.rect.height
            is_portrait = page_height >= page.rect.width
            words = page.get_text("words", sort=True)
            if is_portrait:
                unsafe_words = [
                    word for word in words
                    if word[1] < header_limit and word[3] > header_text_limit
                ]
                if unsafe_words:
                    preview = "、".join(str(word[4]) for word in unsafe_words[:3])
                    issues.append(f"第{page_index}页正文进入页眉安全区（{preview}）")

                if check_shapes:
                    header_shapes = [
                        drawing["rect"] for drawing in page.get_drawings()
                        if drawing["rect"].y0 < header_limit and drawing["rect"].y1 > 70
                    ]
                    if header_shapes:
                        issues.append(f"第{page_index}页表格或边框进入页眉安全区")

            overflow_words = [word for word in words if word[3] > page_height - footer_limit]
            if overflow_words:
                preview = "、".join(str(word[4]) for word in overflow_words[:3])
                issues.append(f"第{page_index}页内容超出页面底部安全区（{preview}）")
            if check_shapes:
                overflow_shapes = [
                    drawing["rect"] for drawing in page.get_drawings()
                    if drawing["rect"].y1 > page_height - footer_limit
                ]
                if overflow_shapes:
                    issues.append(f"第{page_index}页表格或边框超出页面底部安全区")
    finally:
        document.close()

    if issues:
        raise RuntimeError(f"{label}排版检查未通过：{'；'.join(issues[:5])}")


def _render_export_pdf_file(app, html, output_dir, label):
    """Render one static export document without an unnecessary virtual wait."""
    render_id = uuid.uuid4().hex
    pdf_path = os.path.join(output_dir, f"{render_id}.pdf")
    render_info = {}
    try:
        _render_pdf_file(app, html, pdf_path, label, render_info=render_info)
    except Exception:
        app.logger.exception("%s PDF 生成失败", label)
        return None
    try:
        _validate_export_pdf_layout(
            pdf_path,
            label,
            check_shapes=render_info.get("renderer") != "pymupdf",
        )
    except Exception:
        app.logger.exception("%s PDF 排版检查未通过", label)
        return None
    return pdf_path


def _html_max_table_columns(html):
    """Return the largest logical column count used by a generated HTML table."""
    from lxml import html as lxml_html

    root = lxml_html.document_fromstring(html)
    maximum = 0
    for row in root.xpath("//table//tr[not(ancestor::table[2])]"):
        columns = 0
        for cell in row.xpath("./th|./td"):
            try:
                columns += max(1, int(cell.get("colspan") or 1))
            except (TypeError, ValueError):
                columns += 1
        maximum = max(maximum, columns)
    return maximum


def _generated_document_needs_landscape(html, minimum_columns=6):
    """Use landscape pages for generated documents whose tables are too wide."""
    return _html_max_table_columns(html) >= max(2, int(minimum_columns))


def _ensure_landscape_page_rule(html):
    """Promote an A4 generated document to landscape while preserving margins."""
    if re.search(
        r"@page\s*\{[^}]*\bsize\s*:\s*(?:a4\s+)?landscape\b",
        html,
        re.IGNORECASE | re.DOTALL,
    ):
        return html
    return re.sub(
        r"(@page\s*\{[^}]*\bsize\s*:\s*)A4\b",
        r"\1A4 landscape",
        html,
        count=1,
        flags=re.IGNORECASE | re.DOTALL,
    )


def _ordered_attachment_section_ranges(section_start_pages, page_count, ordered_numbers):
    """Validate section markers and return non-overlapping ranges in export order."""
    missing_sections = [
        section_no
        for section_no in ordered_numbers
        if section_no not in section_start_pages
    ]
    if missing_sections:
        raise RuntimeError(
            f"附件板块分页标记缺失：{', '.join(missing_sections)}"
        )

    ordered_start_pages = [
        section_start_pages[section_no]
        for section_no in ordered_numbers
    ]
    if (
        ordered_start_pages != sorted(ordered_start_pages)
        or len(set(ordered_start_pages)) != len(ordered_start_pages)
    ):
        raise RuntimeError("附件板块顺序异常，必须严格按2至13排列")

    ranges = []
    for index, section_no in enumerate(ordered_numbers):
        start_page = section_start_pages[section_no]
        end_page = (
            section_start_pages[ordered_numbers[index + 1]] - 1
            if index + 1 < len(ordered_numbers)
            else page_count - 1
        )
        ranges.append((section_no, start_page, end_page))
    return ranges


def _combine_export_documents(
    documents,
    company_name,
    company_english_name,
    *,
    page_rule,
    title,
):
    """Combine same-orientation documents into one render and optimization pass."""
    from lxml import etree
    from lxml import html as lxml_html

    style_blocks = []
    seen_style_blocks = set()
    document_blocks = []
    for index, document in enumerate(documents):
        root = lxml_html.document_fromstring(document["html"])
        for style in root.xpath("//style/text()"):
            style_text = str(style)
            if not style_text.strip() or style_text in seen_style_blocks:
                continue
            seen_style_blocks.add(style_text)
            style_blocks.append(style_text)
        body_nodes = root.xpath("//body")
        if not body_nodes:
            raise ValueError(f"{document['label']}缺少可打印正文")
        body = body_nodes[0]
        body_classes = [
            class_name
            for class_name in str(body.get("class") or "").split()
            if class_name
        ]
        for header in body.xpath(
            ".//header[contains(concat(' ', normalize-space(@class), ' '), ' document-header ')]"
        ):
            parent = header.getparent()
            if parent is not None:
                parent.remove(header)
        marker = f"GAOXINPDFDOC{index:04d}X{uuid.uuid4().hex[:12].upper()}"
        document["marker"] = marker
        document_classes = "".join(f" {class_name}" for class_name in body_classes)
        body_groups = []
        current_group = []
        for child in body:
            if (
                child.get("data-pymupdf-page-break-before") is not None
                and current_group
            ):
                body_groups.append(current_group)
                current_group = []
            current_group.append(child)
        if current_group:
            body_groups.append(current_group)

        for group_index, group in enumerate(body_groups):
            is_first_block = index == 0 and group_index == 0
            first_class = " batch-document-first" if is_first_block else ""
            pymupdf_break = (
                "" if is_first_block else ' data-pymupdf-page-break-before="true"'
            )
            marker_html = (
                f'<div class="batch-document-marker">{marker}</div>'
                if group_index == 0
                else ""
            )
            group_html = "".join(
                etree.tostring(child, encoding="unicode", method="html")
                for child in group
            )
            document_blocks.append(
                f'<section class="batch-document{first_class}{document_classes}"{pymupdf_break}>'
                f"{marker_html}{group_html}</section>"
            )

    combined_styles = "\n".join(style_blocks)
    return f"""<!doctype html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8">
  <title>{html_escape(title)}</title>
  <style>
  {combined_styles}
  @page {{ {page_rule} }}
  body {{ padding-top: 0 !important; }}
  thead {{ display: table-header-group; }}
  tr {{ break-inside: avoid; page-break-inside: avoid; }}
  .batch-document {{ break-before: page; page-break-before: always; padding-top: 0; }}
  .batch-document-first {{ break-before: auto; page-break-before: auto; }}
  .batch-document-marker {{ height: 1px; overflow: hidden; color: #fff; font: 1px/1px Arial, sans-serif; }}
  </style>
</head>
<body>
  {''.join(document_blocks)}
</body>
</html>"""


def _combine_portrait_export_documents(documents, company_name, company_english_name):
    return _combine_export_documents(
        documents,
        company_name,
        company_english_name,
        page_rule="size: A4; margin: 28mm 16mm 18mm;",
        title="高新技术企业认定附件内部材料",
    )


def _combine_landscape_export_documents(documents, company_name, company_english_name):
    return _combine_export_documents(
        documents,
        company_name,
        company_english_name,
        page_rule="size: A4 landscape; margin: 28mm 12mm 14mm;",
        title="高新技术企业认定附件横向内部材料",
    )


def _assign_portrait_document_page_ranges(pdf_path, documents):
    try:
        import fitz
    except ImportError:
        import pymupdf as fitz

    source_pdf = fitz.open(pdf_path)
    try:
        marker_pages = {}
        pending_markers = {document["marker"] for document in documents}
        for page_index, page in enumerate(source_pdf):
            page_text = page.get_text() or ""
            for marker in list(pending_markers):
                if marker in page_text:
                    marker_pages[marker] = page_index
                    pending_markers.remove(marker)
            if not pending_markers:
                break
        if pending_markers:
            raise RuntimeError(f"内部材料分页标记识别失败（{len(pending_markers)}项）")

        start_pages = [marker_pages[document["marker"]] for document in documents]
        if start_pages != sorted(start_pages) or len(set(start_pages)) != len(start_pages):
            raise RuntimeError("内部材料分页顺序异常")
        for index, document in enumerate(documents):
            document["pdf_path"] = pdf_path
            document["from_page"] = start_pages[index]
            document["to_page"] = (
                start_pages[index + 1] - 1 if index + 1 < len(start_pages) else source_pdf.page_count - 1
            )
    finally:
        source_pdf.close()


def _generated_insert_header_page_indexes(insert_start, insert_end, *, skip_first_page=False):
    first_header_page = int(insert_start) + (1 if skip_first_page else 0)
    return range(min(first_header_page, int(insert_end)), int(insert_end))


def _portrait_export_document_style_signature(document):
    from lxml import html as lxml_html

    root = lxml_html.document_fromstring(document["html"])
    styles = tuple(
        str(style)
        for style in root.xpath("//style/text()")
        if str(style).strip()
    )
    body_nodes = root.xpath("//body")
    body_class = str(body_nodes[0].get("class") or "") if body_nodes else ""
    return styles, body_class


def _portrait_export_document_batches(documents, batch_size):
    normalized_batch_size = max(1, int(batch_size))
    batches = []
    batches_by_signature = {}
    for document in documents:
        signature = _portrait_export_document_style_signature(document)
        signature_batches = batches_by_signature.setdefault(signature, [])
        if not signature_batches or len(signature_batches[-1]) >= normalized_batch_size:
            batch = []
            signature_batches.append(batch)
            batches.append(batch)
        signature_batches[-1].append(document)
    return batches


def _render_export_document_batch(
    app,
    batch,
    output_dir,
    company_name,
    company_english_name,
    *,
    orientation,
    label,
):
    batch_started = time.perf_counter()
    if orientation == "landscape":
        combined_html = _combine_landscape_export_documents(
            batch,
            company_name,
            company_english_name,
        )
    else:
        combined_html = _combine_portrait_export_documents(
            batch,
            company_name,
            company_english_name,
        )
    pdf_path = _render_export_pdf_file(
        app,
        combined_html,
        output_dir,
        label,
    )
    if not pdf_path:
        raise RuntimeError(f"{label} PDF 生成失败")
    _assign_portrait_document_page_ranges(pdf_path, batch)
    app.logger.info(
        "%s PDF 生成完成 documents=%s duration=%.2fs first=%s last=%s",
        label,
        len(batch),
        time.perf_counter() - batch_started,
        batch[0]["label"],
        batch[-1]["label"],
    )
    return pdf_path


def _render_portrait_export_document_batches(
    app,
    documents,
    output_dir,
    company_name,
    company_english_name,
    batch_size,
):
    batches = _portrait_export_document_batches(documents, batch_size)
    pdf_paths = []
    for batch_index, batch in enumerate(batches, start=1):
        label = f"附件内部材料（第{batch_index}/{len(batches)}批）"
        pdf_path = _render_export_document_batch(
            app,
            batch,
            output_dir,
            company_name,
            company_english_name,
            orientation="portrait",
            label=label,
        )
        pdf_paths.append(pdf_path)
    return pdf_paths


def _prepare_export_attachment_files(app, references):
    """Resolve uploaded export files concurrently and reuse duplicate paths."""
    upload_root = os.path.abspath(app.config["UPLOAD_FOLDER"])
    references_by_path = {}
    for reference in references:
        relative_path = str(reference.get("relative_path") or "").strip()
        if relative_path:
            references_by_path.setdefault(relative_path, []).append(reference)

    def resolve(relative_path):
        target = os.path.abspath(os.path.join(upload_root, relative_path))
        if target != upload_root and not target.startswith(upload_root + os.sep):
            app.logger.warning("导出附件路径越界，已跳过：%s", relative_path)
            return None
        try:
            resolved = ensure_local_file(target, relative_path)
            return resolved if os.path.isfile(resolved) else None
        except Exception:
            app.logger.exception("导出附件准备失败：%s", relative_path)
            return None

    configured_workers = max(
        1,
        int(app.config.get("PDF_ATTACHMENT_DOWNLOAD_WORKERS", 6)),
    )
    worker_count = min(configured_workers, len(references_by_path))
    if not worker_count:
        return 0

    with ThreadPoolExecutor(
        max_workers=worker_count,
        thread_name_prefix="gaoxin-attachment",
    ) as executor:
        resolved_paths = dict(
            zip(
                references_by_path,
                executor.map(resolve, references_by_path),
            )
        )

    resolved_count = 0
    for relative_path, path_references in references_by_path.items():
        resolved_path = resolved_paths.get(relative_path)
        for reference in path_references:
            reference["pdf_path"] = resolved_path
            if resolved_path:
                resolved_count += 1
            else:
                app.logger.warning(
                    "%s文件不存在，已跳过：%s",
                    reference.get("label") or "附件",
                    relative_path,
                )
    return resolved_count


def _system_doc_base_defaults(company, data):
    merged = _merge_relation_fields(data or {}) if isinstance(data, dict) else {}
    project_name = _collect_relation_names(merged, "rd_activity") or merged.get("rd_0_name") or ""

    def first_value(*keys):
        for key in keys:
            value = merged.get(key)
            if value not in (None, "", []):
                return value
        return ""

    rd_budget = first_value("evidence_budget_amount", "rd_0_budget", "rd_0_total")
    rd_result = first_value("evidence_expected_result", "rd_0_result")
    rd_purpose = first_value("rd_0_purpose", "rd_0_org_method")
    rd_innovation = first_value("rd_0_innovation", "rd_0_field")
    return {
        "company_name": merged.get("company_name") or company.name,
        "company_english_name": _company_english_name(company, data),
        "responsible_person": merged.get("legal_rep") or merged.get("contact") or "",
        "rd_director": merged.get("rd_director") or merged.get("tech_director") or "",
        "finance_director": merged.get("finance_director") or "",
        "rd_center_name": merged.get("rd_center_name") or "研发中心",
        "rd_department": merged.get("rd_department") or "研发中心",
        "applicable_year": merged.get("year3_label") or "2025",
        "effective_date": merged.get("system_effective_date") or "",
        "tech_field": merged.get("tech_field") or "",
        "main_business": merged.get("business_scope") or "",
        "main_products": _collect_relation_names(merged, "ps_name"),
        "main_rd_projects": project_name,
        "evidence_project_code": merged.get("evidence_project_code") or merged.get("rd_0_no") or "RD01",
        "evidence_project_name": merged.get("evidence_project_name") or project_name,
        "evidence_project_source": merged.get("evidence_project_source") or "企业年度研发计划",
        "evidence_project_period": merged.get("evidence_project_period") or merged.get("rd_0_period") or "",
        "evidence_budget_amount": f"{rd_budget}万元" if rd_budget and "万" not in str(rd_budget) else rd_budget,
        "evidence_project_goal": merged.get("evidence_project_goal") or rd_purpose,
        "evidence_project_content": merged.get("evidence_project_content") or rd_purpose,
        "evidence_project_innovation": merged.get("evidence_project_innovation") or rd_innovation,
        "evidence_acceptance_standard": merged.get("evidence_acceptance_standard") or rd_result,
        "evidence_expected_result": rd_result,
        "evidence_record_date": merged.get("evidence_record_date") or "",
        "evidence_participants": merged.get("evidence_participants") or "",
        "evidence_archive_no": merged.get("evidence_archive_no") or "",
        "evidence_archive_location": merged.get("evidence_archive_location") or "研发档案",
        "generation_notes": merged.get("generation_notes") or "",
        "staff_total": merged.get("staff_total") or merged.get("hr_total") or "",
        "tech_staff": merged.get("tech_staff") or merged.get("hr_tech") or "",
    }


def _load_gaoxin_system_docs(company):
    data = _load_company_data(company)
    relation_rows = ((data.get("gaoxin_relation_table") or {}).get("rows") or []) if isinstance(data, dict) else []
    saved = data.get("gaoxin_system_docs") if isinstance(data.get("gaoxin_system_docs"), dict) else {}
    saved_base = saved.get("base") if isinstance(saved.get("base"), dict) else {}
    defaults = _system_doc_base_defaults(company, data)
    base = {
        key: _expand_relation_references(
            saved_base.get(key) if saved_base.get(key) not in (None, "") else defaults.get(key, ""),
            relation_rows,
        )
        for key, _ in GAOXIN_SYSTEM_BASE_FIELDS
    }
    if defaults.get("tech_field"):
        base["tech_field"] = defaults["tech_field"]
    saved_docs = saved.get("docs") if isinstance(saved.get("docs"), dict) else {}
    docs = {
        item["key"]: re.sub(
            r"(?m)^落款单位：",
            "单位：",
            _expand_relation_reference_text(saved_docs.get(item["key"], "") or "", relation_rows),
        )
        for item in GAOXIN_SYSTEM_DOC_TYPES
    }
    saved_evidence = saved.get("evidence") if isinstance(saved.get("evidence"), dict) else {}
    evidence = {
        item["key"]: _expand_relation_reference_text(saved_evidence.get(item["key"], "") or "", relation_rows)
        for item in GAOXIN_SYSTEM_DOC_TYPES
    }
    products = _collect_hitech_product_rows(_sync_gaoxin_finance_years(_merge_relation_fields(data)))
    base = _normalize_generated_ps_structure(base, products)
    docs = {
        key: _normalize_system_generated_text(value, products)
        for key, value in docs.items()
    }
    evidence = {
        key: _normalize_system_generated_text(value, products)
        for key, value in evidence.items()
    }
    return {"base": base, "docs": docs, "evidence": evidence}


def _save_gaoxin_system_docs(company, form):
    data = _load_company_data(company)
    base = {key: str(form.get(key, "") or "").strip() for key, _ in GAOXIN_SYSTEM_BASE_FIELDS}
    company_name = base.get("company_name") or company.name
    effective_date = base.get("effective_date", "")
    docs = {
        item["key"]: _append_system_doc_signature(
            form.get(f"doc_{item['key']}", ""),
            company_name,
            effective_date,
        )
        for item in GAOXIN_SYSTEM_DOC_TYPES
    }
    evidence = {
        item["key"]: _normalize_system_doc_text(form.get(f"evidence_{item['key']}", ""))
        for item in GAOXIN_SYSTEM_DOC_TYPES
    }
    data["gaoxin_system_docs"] = {"base": base, "docs": docs, "evidence": evidence}
    company.data_json = json.dumps(data, ensure_ascii=False)
    db.session.commit()
    return data["gaoxin_system_docs"]


def _format_system_doc_context(base):
    labels = dict(GAOXIN_SYSTEM_BASE_FIELDS)
    lines = []
    for key, label in labels.items():
        value = str((base or {}).get(key) or "").strip()
        if value:
            lines.append(f"{label}：{value}")
    return "\n".join(lines) if lines else "企业基础信息暂未完善。"


def _system_doc_time_context(base):
    """Build the shared temporal rules for system documents and evidence."""
    base = base or {}
    project_period = str(
        base.get("evidence_project_period")
        or base.get("applicable_year")
        or ""
    ).strip()
    project_temporal = project_temporal_context(project_period)
    effective_date = event_date_context(base.get("effective_date"))
    record_date = event_date_context(base.get("evidence_record_date"))
    if (
        record_date["date_valid"]
        and not record_date["is_future"]
        and record_date.get("start_iso")
    ):
        record_project = project_temporal_context(
            project_period,
            as_of=record_date["start_iso"],
        )
    else:
        record_project = project_temporal_context("")
    return {
        "project": project_temporal,
        "record_project": record_project,
        "effective_date": effective_date,
        "record_date": record_date,
        "prompt": f"""
统一时间规则：
1. 本次生成的系统基准日期为 {project_temporal['as_of_display']}。
2. 佐证项目周期为“{project_temporal['period_display']}”，截至系统基准日期的状态为“{project_temporal['status_display']}”；项目开始时间为“{project_temporal['start_display']}”，结束时间为“{project_temporal['end_display']}”。
3. 记录/审批日期为“{record_date['display']}”。在该记录日期回看，项目状态为“{record_project['status_display']}”；佐证材料中的实施、测试、成果、应用和归档表述必须以该记录日期时的项目状态为准。{record_project['tense_instruction']}
4. 制度生效日期为“{effective_date['display']}”，状态为“{effective_date['status']}”；记录/审批日期为“{record_date['display']}”，状态为“{record_date['status']}”。
5. 立项审批可以早于项目开始，但实施、测试、成果形成、验收和归档必须按真实时间先后排列；未来日期对应的事项只能写“计划、拟开展、待填写”，不得写成已经发生。
6. 计划结束日期仅表示项目已到计划节点，不代表实际完成、通过验收、完成成果转化或已经归档；这些结论必须有对应的真实记录支持。
7. “记录/审批日期”是共享输入，不得机械套用到所有佐证表。费用归集、项目实施、测试、成果、验收和归档类表单若日期早于项目开始，应写“待根据实际记录填写”。
8. 未提供的日期必须留空或写“待补充”，不得自行编造年月日；2023、2024、2025等财务申报年度属于固定业务年度，不得按当前年份替换。
""",
    }


def _normalize_system_doc_text(text):
    text = str(text or "").replace("\r\n", "\n").replace("\r", "\n")
    if text.startswith("```"):
        lines = text.splitlines()
        if len(lines) >= 2:
            text = "\n".join(lines[1:-1])
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = text.replace("**", "")
    lines = [line.rstrip() for line in text.split("\n")]
    normalized = []
    blank_seen = False
    for line in lines:
        if not line.strip():
            if normalized and not blank_seen:
                normalized.append("")
            blank_seen = True
            continue
        normalized.append(line.strip())
        blank_seen = False
    return "\n".join(normalized).strip()


def _append_system_doc_signature(text, company_name, effective_date=""):
    cleaned = _normalize_system_doc_text(text)
    company_name = str(company_name or "").strip()
    effective_date = str(effective_date or "").strip()
    if company_name:
        cleaned = re.sub(r"\n*(?:落款单位|单位)：.*(?:\n日期：.*)?\s*$", "", cleaned).strip()
        cleaned = re.sub(rf"\n*{re.escape(company_name)}\s*\n*(?:日期[:：].*)?\s*$", "", cleaned).strip()
        signature = f"单位：{company_name}"
        if effective_date:
            signature += f"\n日期：{effective_date}"
        else:
            signature += "\n日期：    年    月    日"
        cleaned = f"{cleaned}\n\n{signature}" if cleaned else signature
    return cleaned


def _split_system_doc_signature(text):
    """Separate a system document body from its right-aligned signature."""
    normalized = _normalize_system_doc_text(text)
    match = re.search(
        r"(?:^|\n)(?:落款单位|单位)：(?P<company>[^\n]*)(?:\n日期：(?P<date>[^\n]*))?\s*$",
        normalized,
    )
    if not match:
        return normalized, "", ""
    body = normalized[:match.start()].strip()
    return body, f"单位：{match.group('company').strip()}", (
        f"日期：{match.group('date').strip()}" if match.group("date") is not None else ""
    )


def _safe_docx_name(value):
    name = re.sub(r"[\\/:*?\"<>|\r\n]+", "_", str(value or "").strip())
    name = re.sub(r"\s+", "", name).strip("._ ")
    return name or "未命名文件"


def _evidence_file_templates(doc_key, evidence_text=""):
    templates = list(GAOXIN_SYSTEM_EVIDENCE_FILE_TEMPLATES.get(doc_key) or [])
    if templates:
        return templates
    lines = [line.strip(" 0123456789.、") for line in _normalize_system_doc_text(evidence_text).splitlines() if line.strip()]
    generated = []
    for line in lines[:6]:
        title = re.split(r"[：:；;，,。]", line, 1)[0].strip()
        if title and len(title) <= 32:
            generated.append((title, line))
    return generated or [("制度佐证材料登记表", "用于记录该制度执行过程中形成的佐证材料、责任人和归档情况。")]


def _system_evidence_table_widths(headers):
    """Return readable Story column ratios based on evidence table semantics."""
    narrow_keywords = ("序号", "日期", "签到", "签字", "确认")
    compact_keywords = (
        "审核人",
        "责任人",
        "复核人",
        "移交人",
        "接收人",
        "部门",
        "岗位",
        "科目",
        "金额",
        "编号",
    )
    wide_keywords = (
        "内容",
        "事项",
        "意见",
        "记录",
        "依据",
        "原因",
        "要点",
        "结论",
        "成果",
        "附件",
        "用途",
        "名称",
        "项目",
        "评价",
        "贡献",
        "资料",
    )

    weights = []
    for header in headers or []:
        label = str(header or "").strip()
        if any(keyword in label for keyword in narrow_keywords):
            weight = 0.65
        elif any(keyword in label for keyword in compact_keywords):
            weight = 0.9
        elif any(keyword in label for keyword in wide_keywords):
            weight = 1.55
        else:
            weight = 1.1
        weights.append(weight)

    if not weights:
        return ""
    total = sum(weights)
    return ",".join(f"{weight / total * 100:.2f}" for weight in weights)


def _system_evidence_pdf_context(base, doc_type, file_title, purpose, evidence_text=""):
    """Build concise, table-first content for one system evidence attachment."""
    base = base or {}
    time_context = _system_doc_time_context(base)
    current_temporal = time_context["project"]
    record_context = evidence_record_date_context(
        file_title,
        base.get("evidence_record_date"),
        base.get("evidence_project_period") or base.get("applicable_year"),
    )
    temporal = record_context["record_project"]

    def value(key, fallback=""):
        return str(base.get(key) or fallback).strip()

    def at_record(value_text):
        return enforce_temporal_wording(value_text, temporal)

    safe_purpose = at_record(
        str(purpose or evidence_text or "记录制度执行情况并形成申报佐证。").strip()
    )

    company_name = value("company_name")
    responsible_person = value("responsible_person", "企业负责人")
    rd_director = value("rd_director", "研发负责人")
    finance_director = value("finance_director", "财务负责人")
    rd_center = value("rd_center_name", value("rd_department", "研发中心"))
    rd_department = value("rd_department", rd_center)
    applicable_year = value("applicable_year", "申报年度")
    tech_field = value("tech_field", "企业所属高新技术领域")
    main_business = value("main_business", "企业主营业务")
    main_products = value("main_products", "主要产品或服务")
    main_rd_projects = value("main_rd_projects", "年度研发项目")
    project_code = value("evidence_project_code", "RD01")
    project_name = value("evidence_project_name")
    project_source = value("evidence_project_source", "企业年度研发计划")
    project_period = value("evidence_project_period", f"{applicable_year}年度")
    budget_amount = value("evidence_budget_amount")
    project_goal = at_record(value("evidence_project_goal"))
    project_content = at_record(value("evidence_project_content"))
    project_innovation = at_record(value("evidence_project_innovation"))
    acceptance_standard = at_record(value("evidence_acceptance_standard"))
    expected_result = at_record(value("evidence_expected_result"))
    record_date = record_context["display"]
    participants = value("evidence_participants")
    archive_no = "待填写"
    archive_location = value("evidence_archive_location", "研发档案")
    archive_date = "待根据实际归档记录填写"
    generation_notes = at_record(value("generation_notes"))
    project_values = _split_relation_values(main_rd_projects)
    product_values = _split_relation_values(main_products)
    primary_project = project_name or (project_values[0] if project_values else main_rd_projects)
    primary_product = product_values[0] if product_values else main_products
    approval_opinion = "待根据实际审批记录填写"
    if temporal["status"] == "计划中":
        result_name = f"{primary_project}预期成果"
    elif temporal["status"] == "研发中":
        result_name = f"{primary_project}阶段成果"
    elif temporal["status"] == "已完成":
        result_name = f"{primary_project}计划成果（实际形成情况待核实）"
    else:
        result_name = f"{primary_project}成果（待核实）"

    tables = [{
        "title": "一、基础信息",
        "kind": "key-value",
        "rows": [
            ["企业名称", company_name, "适用年度", applicable_year],
            ["对应制度", doc_type.get("title", ""), "表单名称", file_title],
            ["责任部门", rd_department, "责任人", rd_director],
            ["记录日期", record_date, "归档编号", archive_no],
            ["归档位置", archive_location, "复核人", finance_director],
            [
                "项目周期",
                project_period,
                "时间状态",
                f"当前：{current_temporal['status_display']}；记录日：{record_context['status_display']}",
            ],
        ],
    }]

    if any(word in file_title for word in ["立项申请", "项目登记", "入驻申请"]):
        tables.append({
            "title": "二、项目及审批信息",
            "kind": "key-value",
            "rows": [
                ["项目编号", project_code, "项目名称", primary_project],
                ["项目来源", project_source, "实施周期", project_period],
                ["技术领域", tech_field, "对应产品/服务", primary_product],
                ["承担部门", rd_department, "项目负责人", rd_director],
                ["研发预算", budget_amount, "预期成果", expected_result],
                ["研发目标", project_goal, "验收指标", acceptance_standard],
                ["核心研发内容", project_content, "关键技术/创新点", project_innovation],
            ],
        })
        tables.append({
            "title": "三、审批确认",
            "headers": ["审核环节", "审核人", "日期", "意见"],
            "rows": [
                ["研发部门初审", rd_director, record_date, approval_opinion],
                ["财务预算复核", finance_director, record_date, approval_opinion],
                ["企业负责人审批", responsible_person, record_date, approval_opinion],
            ],
        })
    elif any(word in file_title for word in ["预算", "投入", "费用", "辅助账", "凭证", "分摊", "结转", "复核"]):
        tables.append({
            "title": "二、研发费用及台账记录",
            "headers": ["科目", "对应项目", "归集依据", "责任人", "复核人", "金额/编号"],
            "rows": [
                ["人员人工费用", project_code, "工时记录、工资表", rd_director, finance_director, ""],
                ["直接投入费用", project_code, "领料单、采购发票", rd_director, finance_director, ""],
                ["折旧与摊销", project_code, "设备或软件使用记录", rd_director, finance_director, ""],
                ["设计试验费用", project_code, "试验记录、测试报告", rd_director, finance_director, ""],
                ["其他相关费用", project_code, "研发相关凭证", rd_director, finance_director, ""],
            ],
        })
        tables.append({
            "title": "三、复核确认",
            "kind": "key-value",
            "rows": [
                ["预算金额", budget_amount, "归集口径", "按研发项目归集"],
                ["财务复核人", finance_director, "复核日期", record_date],
                ["差异调整", "", "年度归档位置", archive_location],
            ],
        })
    elif any(word in file_title for word in ["组织架构", "岗位职责"]):
        tables.append({
            "title": "二、组织架构及职责",
            "headers": ["岗位/角色", "人员或部门", "主要职责", "对应事项", "签字"],
            "rows": [
                ["企业负责人", responsible_person, "统筹制度执行和资源保障", f"{applicable_year}年度研发管理", ""],
                ["研发负责人", rd_director, "组织研发项目、成果转化和档案管理", primary_project, ""],
                ["财务负责人", finance_director, "研发费用归集、辅助账复核", f"{applicable_year}年度研发费用", ""],
                ["研发/技术部门", rd_department, "研发、试验、技术资料整理", tech_field, ""],
            ],
        })
    elif any(word in file_title for word in ["设备", "场地", "资源开放", "台账"]):
        tables.append({
            "title": "二、设备场地及资源记录",
            "headers": ["序号", "设备/场地/资源", "用途", "所属部门", "责任人", "使用记录", "维护情况"],
            "rows": [
                ["1", "", f"支撑{primary_project}", rd_department, rd_director, "", ""],
                ["2", "", f"支撑{primary_product}测试或开发", rd_department, rd_director, "", ""],
                ["3", "", "研发资料及成果管理", rd_department, rd_director, "", ""],
            ],
        })
    elif any(word in file_title for word in ["会议", "例会", "纪要", "活动记录", "过程记录"]):
        tables.append({
            "title": "二、会议或活动信息",
            "kind": "key-value",
            "rows": [
                ["主题", primary_project, "组织部门", rd_department],
                ["主持人", rd_director, "参与人员", participants],
                ["时间", record_date, "地点", rd_center],
                ["主要事项", generation_notes or primary_project, "记录人", ""],
            ],
        })
        tables.append({
            "title": "三、任务跟踪",
            "headers": ["序号", "决议/任务", "责任人", "完成期限", "完成情况", "归档附件"],
            "rows": [
                ["1", expected_result or primary_project, rd_director, "", "", ""],
                ["2", "资料复核与归档", finance_director, "", "", ""],
                ["3", "", "", "", "", ""],
            ],
        })
    elif any(word in file_title for word in ["合作", "产学研", "协议"]):
        tables.append({
            "title": "二、合作事项记录",
            "headers": ["合作事项", "合作单位", "合作内容", "成果归属", "责任人", "阶段成果", "验收资料"],
            "rows": [
                [f"{tech_field}方向合作", "", f"围绕{primary_project}开展技术协作", "按协议约定", rd_director, "", ""],
                [f"{primary_product}应用协同", "", f"支撑{main_business}", "按协议约定", rd_director, "", ""],
            ],
        })
    elif any(word in file_title for word in ["成果", "转化", "试制", "试用", "应用证明", "输出确认", "验收"]):
        tables.append({
            "title": "二、成果及转化记录",
            "headers": ["成果名称", "来源项目", "对应产品/服务", "转化形式", "应用场景", "责任人", "证明附件"],
            "rows": [
                [result_name, primary_project, primary_product, "按实际情况填写", main_business, rd_director, ""],
                [expected_result or f"{tech_field}相关预期或阶段成果", primary_project, primary_product, "按实际情况填写", main_business, rd_director, ""],
            ],
        })
        tables.append({
            "title": "三、效果评价",
            "headers": ["评价维度", "简要记录", "数据或附件来源", "复核人"],
            "rows": [
                ["技术效果", "填写性能、效率或质量提升情况", "测试记录、验收报告", rd_director],
                ["应用效果", "填写对主营业务或客户应用的支撑", "销售、应用或反馈资料", responsible_person],
                ["知识产权", "填写关联专利、软著或技术秘密", "证书或知识产权台账", rd_director],
            ],
        })
    elif any(word in file_title for word in ["奖励", "贡献", "绩效", "晋升"]):
        tables.append({
            "title": "二、评价及奖励记录",
            "headers": ["评价/奖励对象", "部门", "关联项目", "贡献事项", "评价指标", "奖励/晋升", "审批意见"],
            "rows": [
                ["", rd_department, primary_project, "研发或成果转化贡献", "项目贡献、成果产出、协作表现", "", ""],
                ["", rd_department, primary_project, f"支撑{primary_product}", "技术难度、质量改进、转化效果", "", ""],
            ],
        })
    elif any(word in file_title for word in ["培训", "进修", "签到", "培养"]):
        tables.append({
            "title": "二、培训信息",
            "kind": "key-value",
            "rows": [
                ["培训主题", generation_notes or tech_field, "组织部门", rd_department],
                ["培训对象", participants, "培训负责人", rd_director],
                ["培训时间", record_date, "培训地点", rd_center],
                ["培训目的", expected_result, "经费预算", budget_amount],
            ],
        })
        tables.append({
            "title": "三、签到及评价",
            "headers": ["序号", "姓名", "部门/岗位", "培训内容", "签到", "评价结果", "备注"],
            "rows": [
                ["1", "", rd_department, generation_notes or tech_field, "", "", ""],
                ["2", "", rd_department, generation_notes or tech_field, "", "", ""],
                ["3", "", "", "", "", "", ""],
            ],
        })
    elif any(word in file_title for word in ["人才", "引进"]):
        tables.append({
            "title": "二、人才引进及评价",
            "headers": ["岗位/人才需求", "所属方向", "需求原因", "评价要点", "考察结论", "审批意见"],
            "rows": [
                [f"{tech_field}研发人才", primary_project, f"支撑{main_business}研发能力建设", "专业能力、项目经验、成果贡献", "", ""],
                [f"{primary_product}技术人才", primary_project, "支撑成果转化和产品服务优化", "技术匹配度、协作能力、稳定性", "", ""],
            ],
        })
    elif any(word in file_title for word in ["档案", "移交", "归档"]):
        tables.append({
            "title": "二、档案移交及归档",
            "headers": ["序号", "档案类别", "档案名称", "对应制度/项目", "移交人", "接收人", "归档位置", "确认"],
            "rows": [
                ["1", "制度文件", doc_type.get("title", ""), doc_type.get("title", ""), rd_director, "", archive_location, ""],
                ["2", "项目资料", primary_project, doc_type.get("title", ""), rd_director, "", archive_location, ""],
                ["3", "财务/台账", f"{applicable_year}年度相关台账", doc_type.get("title", ""), finance_director, "", archive_location, ""],
                ["4", "成果/人员", f"{primary_product}相关材料", doc_type.get("title", ""), rd_director, "", archive_location, ""],
            ],
        })
    else:
        tables.append({
            "title": "二、执行记录",
            "headers": ["序号", "事项", "简要记录", "责任部门", "责任人", "日期", "附件"],
            "rows": [
                ["1", file_title, safe_purpose, rd_department, rd_director, record_date, ""],
                ["2", "资料复核", "材料真实性、完整性复核结果待填写", "财务/行政", finance_director, record_date, ""],
                ["3", "负责人确认", "归档及申报使用意见待填写", "管理层", responsible_person, record_date, ""],
            ],
        })

    if not any(table.get("title") == "三、审批确认" for table in tables):
        tables.append({
            "title": "四、归档确认",
            "kind": "key-value",
            "rows": [
                ["经办人", rd_director, "部门负责人", rd_director],
                ["财务/行政复核", finance_director, "企业负责人", responsible_person],
                ["归档日期", archive_date, "归档位置", archive_location],
            ],
        })

    for table in tables:
        if table.get("kind") != "key-value":
            table["pymupdf_widths"] = _system_evidence_table_widths(
                table.get("headers") or []
            )

    return {
        "purpose": safe_purpose,
        "tables": tables,
    }


def _set_docx_cell_text(cell, text, set_run_font, size=10):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    set_run_font(paragraph.add_run(str(text or "")), size)


def _add_docx_company_header(doc, company_name, company_english_name, set_run_font):
    """Apply a consistent bilingual header to each section of an exported document."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt, RGBColor

    chinese_name = str(company_name or "").strip()
    english_name = str(company_english_name or "").strip()
    header_text = "  |  ".join(part for part in [chinese_name, english_name] if part)
    if not header_text:
        return

    for section in doc.sections:
        section.header_distance = Cm(0.65)
        header = section.header
        header.is_linked_to_previous = False
        paragraph = header.paragraphs[0]
        paragraph.clear()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(header_text)
        set_run_font(run, 8, False)
        run.font.color.rgb = RGBColor(107, 114, 128)


def _add_evidence_form_table(doc, set_run_font, base, doc_type, file_title, purpose, evidence_text=""):
    from docx.shared import Pt

    time_context = _system_doc_time_context(base)
    current_temporal = time_context["project"]
    record_context = evidence_record_date_context(
        file_title,
        base.get("evidence_record_date"),
        base.get("evidence_project_period") or base.get("applicable_year"),
    )
    temporal = record_context["record_project"]

    def at_record(value_text):
        return enforce_temporal_wording(value_text, temporal)

    company_name = str(base.get("company_name") or "").strip()
    responsible_person = str(base.get("responsible_person") or "企业负责人").strip()
    rd_director = str(base.get("rd_director") or "研发负责人").strip()
    finance_director = str(base.get("finance_director") or "财务负责人").strip()
    rd_center = str(base.get("rd_center_name") or base.get("rd_department") or "研发中心").strip()
    rd_department = str(base.get("rd_department") or rd_center).strip()
    tech_field = str(base.get("tech_field") or "企业所属高新技术领域").strip()
    applicable_year = str(base.get("applicable_year") or "申报年度").strip()
    main_business = str(base.get("main_business") or "企业主营业务").strip()
    main_products = str(base.get("main_products") or "主要PS").strip()
    ps_kind = infer_ps_kind(main_products)

    def normalize_output(value):
        return normalize_ps_reference_text(value, main_products, ps_kind)

    main_rd_projects = str(base.get("main_rd_projects") or "年度研发项目").strip()
    project_code = str(base.get("evidence_project_code") or "RD01").strip()
    project_name = str(base.get("evidence_project_name") or "").strip()
    project_source = str(base.get("evidence_project_source") or "企业年度研发计划").strip()
    project_period = str(base.get("evidence_project_period") or f"{applicable_year}年度").strip()
    budget_amount = str(base.get("evidence_budget_amount") or "").strip()
    project_goal = at_record(str(base.get("evidence_project_goal") or "").strip())
    project_content = at_record(str(base.get("evidence_project_content") or "").strip())
    project_innovation = at_record(str(base.get("evidence_project_innovation") or "").strip())
    acceptance_standard = at_record(str(base.get("evidence_acceptance_standard") or "").strip())
    expected_result = at_record(str(base.get("evidence_expected_result") or "").strip())
    record_date = record_context["display"]
    participants = str(base.get("evidence_participants") or "").strip()
    archive_no = "待填写"
    archive_location = str(base.get("evidence_archive_location") or "研发档案").strip()
    archive_date = "待根据实际归档记录填写"
    generation_notes = at_record(str(base.get("generation_notes") or "").strip())
    primary_project = project_name or (_split_relation_values(main_rd_projects)[0] if _split_relation_values(main_rd_projects) else main_rd_projects)
    primary_product = _split_relation_values(main_products)[0] if _split_relation_values(main_products) else main_products
    approval_opinion = "待根据实际审批记录填写"
    if temporal["status"] == "计划中":
        result_name = f"{primary_project}预期成果"
    elif temporal["status"] == "研发中":
        result_name = f"{primary_project}阶段成果"
    elif temporal["status"] == "已完成":
        result_name = f"{primary_project}计划成果（实际形成情况待核实）"
    else:
        result_name = f"{primary_project}成果（待核实）"

    def add_title(text, size=12):
        p = doc.add_paragraph()
        set_run_font(p.add_run(normalize_output(text)), size, True)
        return p

    def add_para(text, size=10):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(21)
        set_run_font(p.add_run(normalize_output(text)), size)
        return p

    def add_kv_table(rows):
        table = doc.add_table(rows=len(rows), cols=4)
        table.style = "Table Grid"
        for row, values in zip(table.rows, rows):
            padded = list(values)[:4] + [""] * max(0, 4 - len(values))
            for cell, value in zip(row.cells, padded):
                _set_docx_cell_text(cell, normalize_output(value), set_run_font, 10)
        return table

    def add_grid(headers, rows, size=9):
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        for cell, header in zip(table.rows[0].cells, headers):
            _set_docx_cell_text(cell, normalize_output(header), set_run_font, size)
        for row_values in rows:
            row = table.add_row().cells
            padded = list(row_values)[:len(headers)] + [""] * max(0, len(headers) - len(row_values))
            for cell, value in zip(row, padded):
                _set_docx_cell_text(cell, normalize_output(value), set_run_font, size)
        return table

    add_title("一、基础信息")
    add_kv_table([
        ("企业名称", company_name, "适用年度", applicable_year),
        ("研发机构", rd_center, "研发/技术部门", rd_department),
        ("研发负责人", rd_director, "财务负责人", finance_director),
        ("企业负责人", responsible_person, "归档编号", archive_no),
        ("归档位置", archive_location, "记录日期", record_date),
        (
            "项目周期",
            project_period,
            "时间状态",
            f"当前：{current_temporal['status_display']}；记录日：{record_context['status_display']}",
        ),
    ])

    title = file_title
    add_title("二、核心内容")
    if any(word in title for word in ["立项申请", "项目登记", "入驻申请"]):
        add_kv_table([
            ("项目编号", project_code, "项目名称", primary_project),
            ("项目来源", project_source, "实施周期", project_period),
            ("技术领域", tech_field, "对应产品/服务", primary_product),
            ("承担部门", rd_department, "项目负责人", rd_director),
            ("研发预算", budget_amount, "预期成果", expected_result),
            ("研发目标", project_goal, "验收指标", acceptance_standard),
            ("核心研发内容", project_content, "关键技术/创新点", project_innovation),
        ])
        add_title("三、审批确认")
        add_grid(["审核环节", "审核人", "日期", "意见"], [
            ["研发部门初审", rd_director, record_date, approval_opinion],
            ["财务预算复核", finance_director, record_date, approval_opinion],
            ["企业负责人审批", responsible_person, record_date, approval_opinion],
        ])
    elif any(word in title for word in ["可行性", "评审记录", "论证"]):
        add_grid(["论证维度", "论证内容", "责任人", "结论"], [
            ["技术可行性", generation_notes or tech_field, rd_director, "待根据实际评审记录填写"],
            ["业务必要性", main_business, rd_director, "待根据实际评审记录填写"],
            ["资源保障", rd_center, rd_director, "待根据实际评审记录填写"],
            ["预算保障", budget_amount, finance_director, "待根据实际评审记录填写"],
            ["评审结论", expected_result, responsible_person, "待根据实际评审记录填写"],
        ])
    elif any(word in title for word in ["预算", "投入", "费用", "辅助账", "凭证", "分摊", "结转"]):
        add_grid(["费用/台账科目", "对应项目", "归集依据", "责任人", "复核人", "金额/编号", "备注"], [
            ["人员人工费用", project_code, "工时记录/工资表", rd_director, finance_director, "", ""],
            ["直接投入费用", project_code, "领料单/采购发票", rd_director, finance_director, "", ""],
            ["折旧与摊销", project_code, "设备或软件使用记录", rd_director, finance_director, "", ""],
            ["设计试验费用", project_code, "试验记录/测试报告", rd_director, finance_director, "", ""],
            ["其他相关费用", project_code, "研发相关凭证", rd_director, finance_director, "", ""],
        ])
        add_title("三、财务复核")
        add_kv_table([
            ("预算金额", budget_amount, "费用归集口径", "按项目归集"),
            ("财务复核人", finance_director, "复核日期", record_date),
            ("差异调整说明", "", "年度归档位置", archive_location),
        ])
    elif any(word in title for word in ["组织架构", "岗位职责"]):
        add_grid(["岗位/角色", "人员或部门", "职责内容", "对应研发管理事项", "签字确认"], [
            ["企业负责人", responsible_person, "统筹研发管理制度执行和资源保障", f"{company_name}{applicable_year}年度高新申报组织管理", ""],
            ["研发负责人", rd_director, f"组织{rd_center}研发项目、成果转化和研发档案管理", primary_project, ""],
            ["财务负责人", finance_director, "负责研发费用归集、辅助账复核和财务凭证留存", f"{applicable_year}年度研发费用管理", ""],
            ["研发/技术部门", rd_department, f"承担{tech_field}方向研发、试验、技术资料整理和归档", primary_project, ""],
        ])
    elif any(word in title for word in ["设备", "场地", "台账"]):
        add_grid(["序号", "设备/场地/软件名称", "用途", "所属部门", "管理责任人", "使用记录", "维护情况"], [
            ["1", "", f"服务于{primary_project}及{tech_field}方向研发", rd_department, rd_director, "", ""],
            ["2", "", f"支撑{primary_product}相关测试、开发或验证", rd_department, rd_director, "", ""],
            ["3", "", f"用于{company_name}研发资料管理和成果转化记录", rd_department, rd_director, "", ""],
        ])
    elif any(word in title for word in ["会议", "例会", "纪要", "活动记录", "过程记录"]):
        add_kv_table([
            ("会议/活动主题", primary_project, "组织部门", rd_department),
            ("主持人", rd_director, "参会/参与人员", participants),
            ("时间", record_date, "地点", rd_center),
            ("讨论事项", generation_notes or primary_project, "记录人", ""),
        ])
        add_title("三、决议及跟踪")
        add_grid(["序号", "决议/任务", "责任人", "完成期限", "完成情况", "归档附件"], [
            ["1", expected_result or primary_project, rd_director, "", "", ""],
            ["2", "财务或行政资料归档", finance_director, "", "", ""],
            ["3", "", "", "", "", ""],
        ])
    elif any(word in title for word in ["合作", "产学研", "协议"]):
        add_grid(["合作事项", "合作单位", "合作内容", "知识产权/成果归属", "企业责任人", "阶段成果", "验收资料"], [
            [f"{company_name}{tech_field}方向合作", "", f"围绕{primary_project}开展技术交流、测试验证或成果转化", "按协议约定", rd_director, "", ""],
            [f"{primary_product}应用协同", "", f"支撑{main_business}相关产品或服务优化", "按协议约定", rd_director, "", ""],
        ])
    elif any(word in title for word in ["成果", "转化", "试制", "试用", "应用证明"]):
        add_grid(["成果名称", "来源研发项目", "对应产品/服务", "转化形式", "应用场景", "责任人", "证明材料"], [
            [result_name, primary_project, primary_product, "按实际情况填写", main_business, rd_director, ""],
            [expected_result or f"{tech_field}相关预期或阶段成果", primary_project, primary_product, "按实际情况填写", main_business, rd_director, ""],
        ])
        add_title("三、转化效果记录")
        add_grid(["评价维度", "记录内容", "数据或附件来源", "复核人"], [
            ["技术效果", f"说明成果对{primary_product}性能、效率、质量或服务能力的提升", "测试记录/验收报告", rd_director],
            ["经济或应用效果", f"说明成果对{company_name}主营业务的支撑", "销售/应用/客户反馈资料", responsible_person],
            ["知识产权支撑", "填写关联专利、软著、标准或技术秘密材料", "知识产权证书或台账", rd_director],
        ])
    elif any(word in title for word in ["奖励", "贡献", "绩效", "晋升"]):
        add_grid(["被评价/奖励对象", "所属部门", "关联研发项目", "贡献事项", "评价指标", "建议奖励/晋升", "审批意见"], [
            ["", rd_department, primary_project, f"参与{tech_field}方向研发或成果转化", "项目贡献、成果产出、协作表现", "", ""],
            ["", rd_department, primary_project, f"支撑{primary_product}相关技术应用", "技术难度、质量改进、转化效果", "", ""],
        ])
        add_title("三、审批确认")
        add_kv_table([
            ("研发负责人意见", rd_director, "财务/行政复核", finance_director),
            ("企业负责人审批", responsible_person, "执行日期", ""),
        ])
    elif any(word in title for word in ["培训", "进修", "签到", "培养"]):
        add_kv_table([
            ("培训主题", generation_notes or tech_field, "组织部门", rd_department),
            ("培训对象", participants, "培训负责人", rd_director),
            ("培训时间", record_date, "培训地点", rd_center),
            ("培训目的", expected_result, "经费来源", budget_amount),
        ])
        add_title("三、签到及评价")
        add_grid(["序号", "姓名", "部门/岗位", "培训内容", "签到", "考核/评价结果", "备注"], [
            ["1", "", rd_department, generation_notes or tech_field, "", "", ""],
            ["2", "", rd_department, generation_notes or tech_field, "", "", ""],
            ["3", "", "", "", "", "", ""],
        ])
    elif any(word in title for word in ["人才", "引进"]):
        add_grid(["岗位/人才需求", "所属方向", "需求原因", "评价要点", "面试/考察结论", "审批意见"], [
            [f"{tech_field}研发人才", primary_project, f"支撑{company_name}{main_business}研发能力建设", "专业能力、项目经验、成果贡献", "", ""],
            [f"{primary_product}产品或服务技术人才", primary_project, "支撑成果转化和产品服务优化", "技术匹配度、协作能力、稳定性", "", ""],
        ])
    elif any(word in title for word in ["档案", "移交", "归档", "验收"]):
        add_grid(["序号", "档案类别", "档案名称", "对应制度/项目", "移交人", "接收人", "归档位置", "完整性确认"], [
            ["1", "制度文件", doc_type.get("title", ""), doc_type.get("title", ""), rd_director, "", "", ""],
            ["2", "项目资料", primary_project, doc_type.get("title", ""), rd_director, "", "", ""],
            ["3", "财务/台账资料", f"{applicable_year}年度相关台账", doc_type.get("title", ""), finance_director, "", "", ""],
            ["4", "成果/人员资料", f"{primary_product}相关证明材料", doc_type.get("title", ""), rd_director, "", "", ""],
        ])
    else:
        add_grid(["序号", "事项名称", "企业化记录内容", "责任部门", "责任人", "日期", "附件"], [
            ["1", file_title, enforce_temporal_wording(f"围绕{company_name}{main_business}、{tech_field}和{primary_project}形成", temporal), rd_department, rd_director, record_date, ""],
            ["2", "财务或行政复核", "材料完整性复核结果待填写", "财务/行政", finance_director, record_date, ""],
            ["3", "负责人确认", "归档及申报使用意见待填写", "管理层", responsible_person, record_date, ""],
        ])

    if not any(word in title for word in ["立项申请", "项目登记", "入驻申请"]):
        add_title("四、归档确认")
        add_kv_table([
            ("经办人", rd_director, "部门负责人", rd_director),
            ("财务/行政复核", finance_director, "企业负责人", responsible_person),
            ("归档日期", archive_date, "归档位置", archive_location),
        ])

def _sync_gaoxin_finance_years(data):
    synced = dict(data or {})
    for index, year in enumerate(GAOXIN_FINANCE_YEARS, start=1):
        prefix = f"year{index}"
        synced[prefix] = year
        synced[f"{prefix}_label"] = year
        net_assets = synced.get(f"fin_{year}_net_assets")
        sales = synced.get(f"fin_{year}_sales") or synced.get(f"fin_{year}_revenue") or synced.get(f"fin_{year}_main_revenue")
        profit = synced.get(f"fin_{year}_profit") or synced.get(f"fin_{year}_net_profit") or synced.get(f"fin_{year}_operating_profit")
        if net_assets not in (None, ""):
            synced[f"{prefix}_net_assets"] = net_assets
        if sales not in (None, ""):
            synced[f"{prefix}_sales"] = sales
        if profit not in (None, ""):
            synced[f"{prefix}_profit"] = profit
    return synced


def _load_high_tech_field_options():
    path = Path(__file__).resolve().parent / "high_tech_fields.json"
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return []
    return data if isinstance(data, list) else []


def _normalize_relation_rows(rows):
    normalized = []
    for raw in rows or []:
        row = {
            "year": str(raw.get("year", "")).strip(),
            "rd_code": str(raw.get("rd_code", "")).strip(),
            "rd_activity": str(raw.get("rd_activity", "")).strip(),
            "rd_period": str(raw.get("rd_period", "")).strip(),
            "ip_code": str(raw.get("ip_code", "")).strip(),
            "ip_name": str(raw.get("ip_name", "")).strip(),
            "ip_auth_no": str(raw.get("ip_auth_no", "")).strip(),
            "ps_code": str(raw.get("ps_code", "")).strip(),
            "ps_name": str(raw.get("ps_name", "")).strip(),
            "result_no": str(raw.get("result_no", "")).strip(),
            "result_name": str(raw.get("result_name", "")).strip(),
            "sales_contract_file_id": str(raw.get("sales_contract_file_id", "")).strip(),
            "sales_contract_code": str(raw.get("sales_contract_code", "")).strip(),
            "sales_contract_filename": str(raw.get("sales_contract_filename", "")).strip(),
            "sales_contract_summary": str(raw.get("sales_contract_summary", "")).strip(),
            "sales_contract_keywords": str(raw.get("sales_contract_keywords", "")).strip(),
        }
        if not any(row.values()):
            continue
        row["ps_display"] = "-".join(part for part in [row["ps_code"], row["ps_name"]] if part)
        normalized.append(row)
    return normalized


def _split_relation_values(value):
    return [item.strip() for item in re.split(r"[；;、,，\n]+", str(value or "")) if item.strip()]


def _relation_source_details(row):
    """Return the usable patent and contract facts for one relation row."""
    patent_name = str(row.get("ip_name") or "").strip()
    contract_content = str(
        row.get("sales_contract_keywords") or row.get("sales_contract_summary") or ""
    ).strip()
    return patent_name, contract_content


def _validate_relation_rows(rows):
    errors = []
    ip_results = {}
    result_names = {}
    name_usage = {}
    rd_ps_usage = {}

    def add_name_usage(kind, value, label):
        name = str(value or "").strip()
        if not name:
            return
        usages = name_usage.setdefault(name, [])
        current = (kind, label)
        if current not in usages:
            usages.append(current)

    def relation_identity(row, code_key, name_key):
        return "-".join(part for part in [row.get(code_key), row.get(name_key)] if part) or row.get(code_key) or row.get(name_key) or ""

    for index, row in enumerate(rows, start=1):
        label = f"第{index}行"
        if not row.get("year"):
            errors.append(f"{label}：请填写年份")
        if not row.get("rd_code") and not row.get("rd_activity"):
            errors.append(f"{label}：请填写 RD 序号或研发活动")
        patent_name, contract_content = _relation_source_details(row)
        if not patent_name and not contract_content:
            errors.append(f"{label}：请至少提供专利名称或销售合同内容（合同关键词/摘要）")
        if not row.get("ps_code") and not row.get("ps_name"):
            errors.append(f"{label}：请填写对应 PS")
        if not row.get("result_no"):
            errors.append(f"{label}：请填写成果序号")
        if not row.get("result_name"):
            errors.append(f"{label}：请填写成果名称")

        ip_keys = _split_relation_values(row.get("ip_code")) or _split_relation_values(row.get("ip_auth_no"))
        result_key = row.get("result_no") or row.get("result_name")
        if ip_keys and result_key:
            for ip_key in ip_keys:
                old = ip_results.get(ip_key)
                if old and old != result_key:
                    errors.append(f"{label}：同一知识产权不能对应多个成果")
                ip_results[ip_key] = result_key

        result_no = row.get("result_no")
        result_name = row.get("result_name")
        if result_no and result_name:
            old_name = result_names.get(result_no)
            if old_name and old_name != result_name:
                errors.append(f"{label}：同一成果序号不能对应多个成果名称")
            result_names[result_no] = result_name

        add_name_usage("RD名称", row.get("rd_activity"), label)
        add_name_usage("PS名称", row.get("ps_name"), label)
        add_name_usage("成果名称", row.get("result_name"), label)

        rd_key = relation_identity(row, "rd_code", "rd_activity")
        ps_key = relation_identity(row, "ps_code", "ps_name")
        if rd_key and ps_key:
            rd_ps_usage.setdefault(rd_key, []).append((ps_key, label))

    for rd_key, usages in rd_ps_usage.items():
        ps_keys = []
        for ps_key, _ in usages:
            if ps_key not in ps_keys:
                ps_keys.append(ps_key)
        if len(ps_keys) > 1:
            usage_text = "、".join(f"{label}{ps_key}" for ps_key, label in usages)
            errors.append(f"同一个 RD “{rd_key}” 下的成果和专利必须同属于一个 PS：{usage_text}")

    for name, usages in name_usage.items():
        kinds = {kind for kind, _ in usages}
        if len(kinds) > 1:
            usage_text = "、".join(f"{label}{kind}" for kind, label in usages)
            errors.append(f"名称“{name}”在成果名称、RD名称、PS名称之间重复：{usage_text}")

    if not rows:
        errors.append("请至少填写一行 RD-IP-PS-成果关联关系")
    return errors


def _relation_payload():
    if request.is_json:
        return request.get_json(silent=True) or {}
    rows_json = request.form.get("rows", "[]")
    try:
        rows = json.loads(rows_json)
    except (json.JSONDecodeError, TypeError):
        rows = []
    return {"rows": rows}


def _save_relation_table(company, rows, tech_field_path=""):
    data = _load_company_data(company)
    data["gaoxin_relation_table"] = {"rows": rows, "tech_field_path": str(tech_field_path or "").strip()}
    data["_application_input_saved"] = True
    company.data_json = json.dumps(data, ensure_ascii=False)
    db.session.commit()


def _recover_relation_sales_contract_files(company, contracts, relation_rows):
    """Restore legacy relation-row files whose attachment metadata was lost."""
    if not company:
        return contracts
    known_ids = {
        str(item.get("id") or "").strip()
        for item in contracts
        if isinstance(item, dict) and str(item.get("id") or "").strip()
    }
    relative_dir = _attachment_relative_path(
        company.user_id,
        company.id,
        "relation_sales_contract",
        "",
    ).rstrip(os.sep)
    absolute_dir = _safe_attachment_path(relative_dir)
    if not os.path.isdir(absolute_dir):
        return contracts

    stored_names = os.listdir(absolute_dir)
    for row in relation_rows or []:
        if not isinstance(row, dict):
            continue
        file_id = str(row.get("sales_contract_file_id") or "").strip()
        if not file_id or file_id in known_ids:
            continue
        stored_filename = next(
            (name for name in stored_names if name.startswith(f"{file_id}_")),
            "",
        )
        if not stored_filename:
            continue
        original_filename = str(row.get("sales_contract_filename") or "").strip()
        if not original_filename:
            original_filename = stored_filename[len(file_id) + 1:]
        contracts.append({
            "id": file_id,
            "original_filename": original_filename,
            "stored_filename": stored_filename,
            "relative_path": os.path.join(relative_dir, stored_filename),
            "source": "relation_sales_contract",
            "year": str(row.get("year") or "").strip(),
            "summary": str(row.get("sales_contract_summary") or "").strip(),
            "keywords": str(row.get("sales_contract_keywords") or "").strip(),
            "recovered_from_relation_row": True,
        })
        known_ids.add(file_id)
    return contracts


def _relation_sales_contracts(data, company=None):
    attachments = _load_gaoxin_attachments_from_data(data)
    section = attachments.setdefault("relation_sales_contract", {"files": []})
    contracts = section.setdefault("files", [])
    relation_rows = ((data.get("gaoxin_relation_table") or {}).get("rows") or [])
    _recover_relation_sales_contract_files(company, contracts, relation_rows)
    for item in contracts:
        if (
            not isinstance(item, dict)
            or item.get("sha256")
            or item.get("blob_url")
            or item.get("blob_etag")
        ):
            continue
        relative_path = str(item.get("relative_path") or "").strip()
        path = _safe_attachment_path(relative_path) if relative_path else ""
        if path and os.path.isfile(path):
            item["sha256"] = sales_contract_file_sha256(path)
    ensure_sales_contract_codes(contracts, relation_rows)
    remap_sales_contract_rows(relation_rows, contracts)
    return attachments, contracts


def _relation_sales_contract_options(contracts):
    return sorted(
        [
            {
                "id": str(item.get("id") or "").strip(),
                "code": str(item.get("contract_code") or "").strip(),
                "year": str(item.get("year") or "").strip(),
                "original_filename": str(item.get("original_filename") or "").strip(),
                "summary": str(item.get("summary") or "").strip(),
                "keywords": str(item.get("keywords") or "").strip(),
            }
            for item in selectable_sales_contracts(contracts)
            if isinstance(item, dict) and str(item.get("id") or "").strip()
        ],
        key=lambda item: (
            item["year"],
            item["code"],
            item["original_filename"],
        ),
    )


GAOXIN_HR_STAFF_HEADERS = [
    "序号",
    "姓名",
    "身份证号",
    "是否签订合同",
    "是否缴纳社保",
    "学历",
    "职称",
    "是否科技人员",
]

GAOXIN_HR_STAFF_HEADER_ALIASES = {
    "序号": {"序号", "编号"},
    "姓名": {"姓名", "人员姓名", "员工姓名"},
    "身份证号": {"身份证号", "身份证号码", "证件号码"},
    "是否签订合同": {"是否签订合同", "是否签订劳动合同", "劳动合同"},
    "是否缴纳社保": {"是否缴纳社保", "社保", "社保缴纳"},
    "学历": {"学历", "最高学历", "学历学位"},
    "职称": {"职称", "专业技术职称", "技术职称"},
    "是否科技人员": {"是否科技人员", "科技人员", "是否研发人员"},
}


def _normalize_hr_staff_value(value):
    if value in (None, ""):
        return ""
    if hasattr(value, "strftime"):
        return value.strftime("%Y-%m-%d")
    text = str(value).strip()
    if text.endswith(".0") and text[:-2].isdigit():
        return text[:-2]
    return text


def _classify_hr_education(value):
    text = str(value or "").strip()
    if "博士" in text:
        return "edu_phd"
    if "硕士" in text or "研究生" in text:
        return "edu_master"
    if "本科" in text or "学士" in text:
        return "edu_bachelor"
    return "edu_diploma" if text else ""


def _classify_hr_title(value):
    text = str(value or "").strip()
    if any(word in text for word in ["高级", "正高", "副高"]):
        if "技工" in text or "技能" in text:
            return "title_tech"
        return "title_senior"
    if "中级" in text:
        return "title_mid"
    if "初级" in text or "助理" in text:
        return "title_junior"
    if "技工" in text or "技师" in text:
        return "title_tech"
    return ""


def _summarize_hr_staff_rows(rows):
    summary = {
        "hr_total": len(rows),
        "hr_fulltime": len(rows),
        "edu_phd": 0,
        "edu_master": 0,
        "edu_bachelor": 0,
        "edu_diploma": 0,
        "title_senior": 0,
        "title_mid": 0,
        "title_junior": 0,
        "title_tech": 0,
    }
    for row in rows:
        edu_key = _classify_hr_education(row.get("学历"))
        title_key = _classify_hr_title(row.get("职称"))
        if edu_key:
            summary[edu_key] += 1
        if title_key:
            summary[title_key] += 1
    tech_count = sum(1 for row in rows if str(row.get("是否科技人员", "")).strip() == "是")
    if tech_count:
        summary["hr_tech"] = tech_count
        summary["tech_staff"] = tech_count
    return summary


def _import_hr_staff_excel(upload):
    from openpyxl import load_workbook

    wb = load_workbook(upload, data_only=True)
    ws = wb.active
    headers = [_normalize_hr_staff_value(cell.value) for cell in ws[1]]
    indexes = {}
    for canonical, aliases in GAOXIN_HR_STAFF_HEADER_ALIASES.items():
        matched = next((index for index, header in enumerate(headers) if header in aliases), None)
        if matched is not None:
            indexes[canonical] = matched
    if "姓名" not in indexes:
        raise ValueError("缺少表头：姓名")

    rows = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        item = {
            header: _normalize_hr_staff_value(
                row[indexes[header]]
                if header in indexes and indexes[header] < len(row)
                else ""
            )
            for header in GAOXIN_HR_STAFF_HEADERS
        }
        if not any(item.values()):
            continue
        if not item.get("姓名"):
            continue
        if not item.get("序号"):
            item["序号"] = str(len(rows) + 1)
        rows.append(item)
    return rows


def _create_hr_staff_template():
    from io import BytesIO
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    wb = Workbook()
    ws = wb.active
    ws.title = "企业人员情况"
    header_fill = PatternFill("solid", fgColor="D9EAF7")
    thin = Side(style="thin", color="B7C9D6")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    for col, header in enumerate(GAOXIN_HR_STAFF_HEADERS, start=1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.font = Font(name="Arial", bold=True, color="000000")
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = border
    examples = [
        [1, "张三", "420100199001010011", "是", "是", "本科", "中级工程师", "是"],
        [2, "李四", "420100199202020022", "是", "是", "硕士", "高级工程师", "是"],
    ]
    for row in examples:
        ws.append(row)
    widths = [10, 14, 24, 16, 16, 14, 18, 16]
    for idx, width in enumerate(widths, start=1):
        ws.column_dimensions[chr(64 + idx)].width = width
    ws.freeze_panes = "A2"
    stream = BytesIO()
    wb.save(stream)
    stream.seek(0)
    return stream


def _save_gaoxin_book_data(company, form_data, finance_data=None):
    existing = _load_company_data(company)
    relation_table = existing.get("gaoxin_relation_table")
    system_docs = existing.get("gaoxin_system_docs")
    hr_staff_rows = existing.get("hr_staff_rows")
    gaoxin_attachments = existing.get("gaoxin_attachments")
    data = dict(form_data)
    if finance_data:
        data.update(finance_data)
    data = _sync_gaoxin_finance_years(data)
    if relation_table:
        data["gaoxin_relation_table"] = relation_table
    if system_docs:
        data["gaoxin_system_docs"] = system_docs
    if hr_staff_rows:
        data["hr_staff_rows"] = hr_staff_rows
    if gaoxin_attachments:
        data["gaoxin_attachments"] = gaoxin_attachments
    for key, value in existing.items():
        if re.match(r"^cv_\d+_", key) or key in GAOXIN_SYSTEM_FRAMEWORK_FIELDS:
            data[key] = value
    data["_application_input_saved"] = True
    company.data_json = json.dumps(data, ensure_ascii=False)
    db.session.commit()


def _relation_label(code, name, max_len=None):
    """Return the complete RD/IP/PS label without shortening its name."""
    code = str(code or "").strip()
    name = str(name or "").strip()
    if code and name:
        return f"{code} - {name}"
    return code or name


def _expand_relation_reference_text(text, relation_rows):
    """Restore legacy shortened RD/IP/PS references from the relation table."""
    text = str(text or "")
    if not text or not relation_rows or not re.search(r"…|\.{3,}", text):
        return text

    reference_names = {}
    for row in relation_rows:
        if not isinstance(row, dict):
            continue
        for code_key, name_key in [
            ("rd_code", "rd_activity"),
            ("ip_code", "ip_name"),
            ("ps_code", "ps_name"),
        ]:
            code = str(row.get(code_key) or "").strip()
            name = str(row.get(name_key) or "").strip()
            if code and name:
                reference_names.setdefault(code, name)

    ellipsis = r"(?:…+|\.{3,})"
    for code, name in sorted(reference_names.items(), key=lambda item: len(item[0]), reverse=True):
        escaped_code = re.escape(code)
        text = re.sub(
            rf"(?<![A-Za-z0-9]){escaped_code}(\s*[-－—]\s*)[^\n，。；：:（）()]{{0,120}}?{ellipsis}",
            lambda match, code=code, name=name: f"{code}{match.group(1)}{name}",
            text,
        )

        def replace_quoted(match, code=code, name=name):
            opening = match.group(1)
            closing = "”" if opening.strip() == "“" else '"'
            return f"{code}{opening}{name}{closing}"

        text = re.sub(
            rf"(?<![A-Za-z0-9]){escaped_code}(\s*[“\"])[^\n”\"]{{0,120}}?{ellipsis}(?:[”\"])?",
            replace_quoted,
            text,
        )

    # Older generated content can wrap the RD/IP/PS code onto the previous
    # line, leaving only a shortened name such as "基于EMS的电网在线继电保护定…".
    # Match known name prefixes as well so those saved values are restored.
    for name in sorted(set(reference_names.values()), key=len, reverse=True):
        if len(name) < 7:
            continue
        prefix_alternatives = "|".join(
            re.escape(name[:prefix_length])
            for prefix_length in range(len(name) - 1, 5, -1)
        )
        text = re.sub(
            rf"(?:{prefix_alternatives}){ellipsis}",
            lambda _match, full_name=name: full_name,
            text,
        )
    return text


def _expand_relation_references(value, relation_rows):
    if isinstance(value, str):
        return _expand_relation_reference_text(value, relation_rows)
    if isinstance(value, list):
        return [_expand_relation_references(item, relation_rows) for item in value]
    if isinstance(value, dict):
        return {key: _expand_relation_references(item, relation_rows) for key, item in value.items()}
    return value


def _load_ip_details(company):
    ip_details = []
    try:
        from modules.parser.routes import _get_ip_certs
        ip_details = _get_ip_certs()
    except Exception:
        ip_details = []

    if not ip_details and company.ip_certs_json:
        try:
            ip_details = json.loads(company.ip_certs_json)
        except (json.JSONDecodeError, TypeError):
            ip_details = []

    return ip_details if isinstance(ip_details, list) else []


def _build_relation_ip_options(ip_details):
    options = []
    seen = set()
    for index, ip in enumerate(ip_details or [], start=1):
        if not isinstance(ip, dict):
            continue
        details = (ip.get("parsed") or {}).get("details") or {}
        name = str(details.get("name") or "").strip()
        auth_no = str(details.get("patent_no") or details.get("grant_no") or "").strip()
        app_date = str(details.get("app_date") or details.get("application_date") or details.get("apply_date") or "").strip()
        code = f"IP{str(index).zfill(2)}"
        key = auth_no or name or code
        if key in seen:
            continue
        seen.add(key)
        options.append({"code": code, "name": name, "auth_no": auth_no, "app_date": app_date})
    return options


def _ip_type_label(value):
    return {
        "invention": "发明",
        "utility": "实用新型",
        "design": "外观设计",
        "copyright": "软著",
    }.get(str(value or ""), str(value or ""))


def _build_ip_attachment_rows(ip_details):
    rows = []
    summary = {
        "ip_invention": 0,
        "ip_defense": 0,
        "ip_plant": 0,
        "ip_crop": 0,
        "ip_pharma": 0,
        "ip_tcm": 0,
        "ip_ic": 0,
        "ip_utility": 0,
        "ip_design": 0,
        "ip_copyright": 0,
    }
    for index, ip in enumerate(ip_details or [], start=1):
        parsed = ip.get("parsed") or {}
        details = parsed.get("details") or {}
        patent_type = parsed.get("patent_type") or details.get("patent_type") or ""
        label = _ip_type_label(patent_type)
        if patent_type == "invention" or label == "发明":
            summary["ip_invention"] += 1
        elif patent_type == "utility" or label == "实用新型":
            summary["ip_utility"] += 1
        elif patent_type == "design" or label == "外观设计":
            summary["ip_design"] += 1
        elif patent_type == "copyright" or label in ("软著", "软件著作权"):
            summary["ip_copyright"] += 1
        rows.append({
            "seq": f"IP{str(index).zfill(2)}",
            "cert_id": ip.get("id", ""),
            "name": details.get("name", ""),
            "type": label,
            "status": details.get("legal_status") or ("授权" if details.get("grant_no") or details.get("grant_date") else ""),
            "patent_no": details.get("patent_no", "") or details.get("grant_no", ""),
            "app_date": details.get("app_date", ""),
            "applicant": details.get("applicant", ""),
            "grant_date": details.get("grant_date", ""),
            "filename": ip.get("filename", ""),
        })
    return summary, rows


def _ip_attachment_summary_from_data(data, fallback):
    summary = dict(fallback or {})
    for key in [
        "ip_invention", "ip_defense", "ip_plant", "ip_crop", "ip_pharma",
        "ip_tcm", "ip_ic", "ip_utility", "ip_design", "ip_copyright",
    ]:
        value = data.get(key)
        if value not in (None, ""):
            summary[key] = value
    return summary


def _ip_attachment_rows_from_data(data, fallback_rows):
    rows = []
    for index in range(100):
        row = {
            "seq": data.get(f"ip_{index}_seq") or f"IP{str(index + 1).zfill(2)}",
            "cert_id": data.get(f"ip_{index}_cert_id", ""),
            "name": data.get(f"ip_{index}_name", ""),
            "type": data.get(f"ip_{index}_type", ""),
            "status": data.get(f"ip_{index}_status", ""),
            "patent_no": data.get(f"ip_{index}_patent_no", ""),
            "app_date": data.get(f"ip_{index}_app_date", ""),
            "applicant": data.get(f"ip_{index}_applicant", ""),
            "grant_date": data.get(f"ip_{index}_date", ""),
            "filename": "",
        }
        # A default type/status alone does not represent a real patent row.
        # Require an identifying field so unused form rows stay out of exports.
        if any(
            str(row.get(key) or "").strip()
            for key in ["cert_id", "name", "patent_no", "app_date", "applicant", "grant_date"]
        ):
            rows.append(row)
    return rows or fallback_rows


def _ip_attachment_context(company, data=None):
    data = data if isinstance(data, dict) else _load_company_data(company)
    ip_details = _load_ip_details(company)
    fallback_summary, fallback_rows = _build_ip_attachment_rows(ip_details)
    summary = _ip_attachment_summary_from_data(data, fallback_summary)
    rows = _ip_attachment_rows_from_data(data, fallback_rows)
    return {"summary": summary, "rows": rows, "count": len(rows)}


def _strip_json_block(content):
    content = str(content or "").strip()
    if content.startswith("```"):
        lines = content.split("\n")
        content = "\n".join(lines[1:-1]).strip()
    return content


def _extract_json_object(content):
    decoder = json.JSONDecoder()
    text = _strip_json_block(content)
    try:
        return decoder.raw_decode(text)[0]
    except (json.JSONDecodeError, TypeError, ValueError):
        pass

    start = text.find("{")
    while start != -1:
        try:
            return decoder.raw_decode(text[start:])[0]
        except (json.JSONDecodeError, TypeError, ValueError):
            start = text.find("{", start + 1)
    return None


def _extract_labeled_value(content, labels, stop_labels):
    label_pattern = "|".join(re.escape(label) for label in labels)
    stop_pattern = "|".join(re.escape(label) for label in stop_labels)
    pattern = rf"(?:^|[\n\r\-*•\s])(?:{label_pattern})\s*[：:]\s*(.*?)(?=\n\s*(?:[-*•]\s*)?(?:{stop_pattern})\s*[：:]|$)"
    match = re.search(pattern, str(content or ""), re.S)
    return match.group(1).strip().strip('"“”') if match else ""


def _normalize_relation_value(value):
    if isinstance(value, list):
        return "；".join(str(item).strip() for item in value if str(item).strip())
    return str(value or "").strip()


def _merge_relation_result(base, incoming):
    if not base.get("result_name") and incoming.get("result_name"):
        base["result_name"] = incoming.get("result_name")
    return base


def _extract_relation_result_fields(data):
    result_keys = ["result_name", "resultName", "name", "title", "achievement", "achievement_name", "achievementName", "product_name", "productName", "成果名称", "科技成果名称", "科技成果", "成果", "名称", "专业名词"]
    technology_keys = ["technology", "technologies", "tech", "technical", "core_technology", "coreTechnology", "technical_details", "technicalDetails", "detail", "details", "description", "技术", "科技成果技术详情", "技术详情", "核心技术", "技术内容", "关键技术", "技术名称"]

    parsed = {"result_name": ""}

    if isinstance(data, str):
        parsed_json = _extract_json_object(data)
        if parsed_json is not None:
            _merge_relation_result(parsed, _extract_relation_result_fields(parsed_json))
        _merge_relation_result(parsed, {
            "result_name": _extract_labeled_value(data, result_keys, technology_keys),
            "technology": _extract_labeled_value(data, technology_keys, result_keys),
        })
        return parsed

    if isinstance(data, list):
        text_parts = []
        for item in data:
            if isinstance(item, str):
                text_parts.append(item)
            elif isinstance(item, dict):
                for text_key in ["text", "content", "output_text"]:
                    if isinstance(item.get(text_key), str):
                        text_parts.append(item.get(text_key))
            _merge_relation_result(parsed, _extract_relation_result_fields(item))
            if parsed.get("result_name"):
                return parsed
        if text_parts:
            _merge_relation_result(parsed, _extract_relation_result_fields("\n".join(text_parts)))
        return parsed

    if not isinstance(data, dict):
        return parsed

    for key in result_keys:
        if data.get(key):
            parsed["result_name"] = _normalize_relation_value(data.get(key))
            break
    for key in technology_keys:
        if data.get(key):
            parsed["technology"] = _normalize_relation_value(data.get(key))
            break
    if parsed.get("result_name"):
        return parsed

    for key in ["data", "result", "results", "items", "output", "outputs", "output_text", "content", "text", "answer", "message", "choices", "arguments", "function", "function_call", "tool_calls", "delta"]:
        if key in data:
            _merge_relation_result(parsed, _extract_relation_result_fields(data.get(key)))
            if parsed.get("result_name"):
                return parsed
    return parsed


def _fallback_relation_result(row):
    ip_name = str(row.get("ip_name") or "").strip()
    contract_content = str(row.get("sales_contract_keywords") or row.get("sales_contract_summary") or "").strip()
    rd_activity = str(row.get("rd_activity") or "").strip()
    ps_name = str(row.get("ps_name") or "").strip()
    result_name = re.sub(r"^(一种|一种用于|一种基于|一种新型|一种实用型)", "", ip_name).strip(" ，,。") or contract_content
    technology_source = result_name or rd_activity or ps_name
    technology = f"{technology_source}关键技术" if technology_source else ""
    return {"result_name": result_name, "technology": technology}


def _parse_relation_result_content(content):
    parsed = _extract_relation_result_fields(content)
    if parsed.get("result_name"):
        return parsed

    data = _extract_json_object(content)
    _merge_relation_result(parsed, _extract_relation_result_fields(data))
    return parsed


RD_ACTIVITY_SUFFIX = "的研发"


def _ensure_rd_activity_suffix(name):
    name = str(name or "").strip().strip('"“” ，,。；;')
    if not name:
        return ""
    if name.endswith("的研发"):
        return name
    if name.endswith("研发"):
        name = name[:-2].rstrip("的")
    elif name.endswith("的"):
        return f"{name}研发"
    return f"{name}的研发"


def _normalize_rd_activity_name(name):
    base = str(name or "").strip().strip('"“” ，,。；;')
    if base.endswith(RD_ACTIVITY_SUFFIX):
        base = base[:-len(RD_ACTIVITY_SUFFIX)]
    elif base.endswith("研发"):
        base = base[:-2].rstrip("的")
    base = _clean_rd_activity_base(base) or "科技成果转化项目"
    return _ensure_rd_activity_suffix(base)


def _parse_rd_activity_content(content):
    data = _extract_json_object(content)
    for source in [data, content]:
        if isinstance(source, dict):
            for key in ["rd_activity", "activity_name", "name", "研发活动", "研发活动名称", "项目名称"]:
                if source.get(key):
                    return _normalize_rd_activity_name(source.get(key))
        if isinstance(source, str):
            value = _extract_labeled_value(source, ["rd_activity", "activity_name", "研发活动", "研发活动名称", "项目名称"], ["说明", "理由", "依据"])
            if value:
                return _normalize_rd_activity_name(value)
            cleaned = re.sub(r"^```(?:json)?|```$", "", source.strip(), flags=re.I).strip()
            if cleaned and "{" not in cleaned and len(cleaned) <= 80:
                return _normalize_rd_activity_name(cleaned)
    return ""


def _clean_relation_topic(value):
    text = str(value or "").strip()
    text = re.sub(r"^(一种|一种用于|一种基于|一种新型|一种实用型)", "", text)
    text = re.sub(r"(的研发|研发|关键技术|技术|方法|系统|装置|设备|工艺)$", "", text)
    text = text.strip(" ，,。；;")
    if not text or re.fullmatch(r"[\d\W_]+", text):
        return ""
    if len(text) <= 2 and not re.search(r"[一-鿿]{2,}", text):
        return ""
    if re.fullmatch(r"(RD|IP|PS)?\d+", text, re.I):
        return ""
    return text


def _clean_rd_activity_base(value):
    parts = [part for part in re.split(r"[与及、]+", str(value or "")) if _clean_relation_topic(part)]
    return "与".join(parts).strip("与及、 ，,。；;")

def _fallback_rd_activity(rows):
    texts = []
    for row in rows:
        for key in ["result_name", "ip_name", "sales_contract_keywords", "technology", "ps_name"]:
            value = _clean_relation_topic(row.get(key))
            if value and value not in texts:
                texts.append(value)
    if not texts:
        return _normalize_rd_activity_name("科技成果转化项目")

    joined = "；".join(texts)
    domain_terms = [term for term in ["物联网", "卫星物联网", "光电传感", "数据安全", "网络安全", "信息安全", "数据分析", "预警", "漏洞扫描", "漏洞修复", "备份恢复", "安全审计", "智能检测", "防护", "监测", "控制", "管理", "识别"] if term in joined]
    if "卫星物联网" in domain_terms and "物联网" in domain_terms:
        domain_terms.remove("物联网")
    if len(domain_terms) >= 2:
        base = "".join(domain_terms[:3]) + "技术"
    else:
        first = texts[0]
        base = re.sub(r"(平台|系统|软件|装置|设备|产品|成果)$", "", first)
        base = f"{base}技术" if not base.endswith("技术") else base
    return _normalize_rd_activity_name(base)


def _generate_rd_activity(rows):
    rows = [row for row in rows or [] if isinstance(row, dict)]
    source_rows = [row for row in rows if any(_relation_source_details(row))]
    if not source_rows:
        return {"success": False, "error": "请先填写该项目下的专利名称或销售合同内容"}

    lines = []
    seen_lines = set()
    for row in source_rows:
        parts = []
        for label, key in [("专利名称", "ip_name"), ("成果名称", "result_name"), ("合同关键词", "sales_contract_keywords"), ("合同摘要", "sales_contract_summary"), ("历史技术", "technology"), ("PS", "ps_name")]:
            value = str(row.get(key) or "").strip()
            if value:
                parts.append(f"{label}：{value}")
        line = "；".join(parts)
        if line and line not in seen_lines:
            seen_lines.add(line)
            lines.append(f"{len(lines) + 1}. {line}")

    prompt = f"""请根据同一个 RD 项目下的专利名称和销售合同内容，结合已有成果名称，提炼一个上位的研发活动名称。

项目资料：
{chr(10).join(lines)}

要求：
1. 仅以提供的专利名称和销售合同内容作为事实依据；同时提供时必须综合两类信息，只有一类时仅依该类信息归纳。
2. 把专利名称、合同所涉产品/服务和已有成果名称作为研发产出或应用证据，反推共同研发方向。
3. 输出应是精简的上位技术项目名称，例如“物联网安全防护技术的研发”，而不是把多个名称直接拼接。
4. 不要直接复制单个专利、合同产品或成果名称，也不要用长串“与”连接多个平台/系统/软件名称。
5. 名称应体现核心技术方向，适合政府项目申报材料。
6. 名称最后必须以“的研发”结尾。
7. 只输出 JSON，不要 markdown，不要解释。

JSON 格式：
{{"rd_activity": "上位技术方向的研发"}}"""
    result = call_llm([
        {"role": "system", "content": "你是高新技术企业认定申报顾问，擅长把同一研发项目下的多个知识产权名称和成果名称综合归纳为一个研发活动名称。输出必须是可解析 JSON，且必须以“的研发”结尾。"},
        {"role": "user", "content": prompt},
    ], temperature=0.2, max_tokens=300, timeout=35, max_attempts=1)
    fallback = _fallback_rd_activity(source_rows)
    if not result.get("success"):
        return {"success": True, "rd_activity": fallback}

    rd_activity = _parse_rd_activity_content(result.get("content")) or fallback
    single_names = {
        _ensure_rd_activity_suffix(_clean_relation_topic(row.get(key)))
        for row in source_rows
        for key in ["ip_name", "result_name"]
        if row.get(key)
    }
    if rd_activity in single_names and len(single_names) > 1:
        rd_activity = fallback
    if rd_activity.count("与") >= 2 or len(rd_activity) > 28:
        rd_activity = fallback
    return {"success": True, "rd_activity": _normalize_rd_activity_name(rd_activity)}


def _parse_ps_name_content(content):
    data = _extract_json_object(content)
    for source in [data, content]:
        if isinstance(source, dict):
            for key in ["ps_name", "product_service_name", "name", "PS名称", "产品服务名称", "产品名称", "服务名称"]:
                if source.get(key):
                    return str(source.get(key)).strip().strip('"“” ，,。；;')
        if isinstance(source, str):
            value = _extract_labeled_value(source, ["ps_name", "product_service_name", "PS名称", "产品服务名称", "产品名称", "服务名称"], ["说明", "理由", "依据"])
            if value:
                return value.strip().strip('"“” ，,。；;')
            cleaned = re.sub(r"^```(?:json)?|```$", "", source.strip(), flags=re.I).strip()
            if cleaned and "{" not in cleaned and len(cleaned) <= 80:
                return cleaned.strip('"“” ，,。；;')
    return ""


def _infer_business_type(company, company_data):
    text_parts = [company.name]
    for key in ["industry_code", "sub_industry_code", "main_product_name", "main_product_category", "company_overview", "business_scope", "industry_chain"]:
        value = company_data.get(key)
        if value:
            text_parts.append(str(value))
    text = " ".join(text_parts)
    if any(word in text for word in ["服务", "平台", "软件", "信息技术", "咨询", "运营", "设计", "广告"]):
        return "service"
    if any(word in text for word in ["制造", "生产", "加工", "设备", "材料", "产品", "纸箱", "部件", "装置"]):
        return "manufacturing"
    return "unknown"


def _fallback_ps_name(rows, business_type):
    names = []
    for row in rows:
        for key in ["result_name", "rd_activity", "sales_contract_keywords", "technology", "ip_name"]:
            value = str(row.get(key) or "").strip()
            if value and value not in names:
                names.append(value)
    base = names[0] if names else "高新技术"
    base = re.sub(r"(的研发|研发)$", "", base).strip(" ，,。") or base
    return f"{base}服务" if business_type == "service" and not base.endswith("服务") else (base if business_type == "service" else base)


def _generate_ps_name(company, rows):
    rows = [row for row in rows or [] if isinstance(row, dict)]
    source_rows = [row for row in rows if str(row.get("rd_activity") or "").strip()]
    if not source_rows:
        return {"success": False, "error": "请先生成或填写 RD 名称"}

    company_data = _load_company_data(company)
    business_type = _infer_business_type(company, company_data)
    type_hint = "服务业，PS名称应为“……服务”" if business_type == "service" else "制造业，PS名称应为一个产品名称" if business_type == "manufacturing" else "根据资料判断是服务还是产品，服务业以“服务”结尾，制造业使用产品名称"
    lines = []
    seen = set()
    for row in source_rows:
        line = "；".join(
            f"{label}：{value}" for label, key in [("RD名称", "rd_activity"), ("成果名称", "result_name"), ("合同关键词", "sales_contract_keywords"), ("合同摘要", "sales_contract_summary"), ("历史技术", "technology"), ("知识产权", "ip_name")]
            for value in [str(row.get(key) or "").strip()] if value
        )
        if line and line not in seen:
            seen.add(line)
            lines.append(f"{len(lines) + 1}. {line}")

    prompt = f"""请根据企业所有 RD 名称，总结一个统一的 PS 名称；成果名称、销售合同关键词和知识产权名称作为辅助理解材料。

企业名称：{company.name}
行业判断：{type_hint}
资料：
{chr(10).join(lines)}

要求：
1. PS名称必须是一个总名称，优先概括所有 RD 名称体现的研发方向，不要逐条罗列。
2. 如果是服务业，名称应为“……服务”；如果是制造业，名称应为一个产品名称。
3. 名称要专业、简洁，适合高新技术企业申报材料。
4. 只输出 JSON，不要 markdown，不要解释。

JSON 格式：
{{"ps_name": "PS名称"}}"""
    result = call_llm([
        {"role": "system", "content": "你是高新技术企业认定申报顾问，擅长把多个研发活动名称归纳为一个产品或服务名称。输出必须是可解析 JSON。"},
        {"role": "user", "content": prompt},
    ], temperature=0.2, max_tokens=300, timeout=35, max_attempts=1)
    fallback = _fallback_ps_name(source_rows, business_type)
    if not result.get("success"):
        return {"success": True, "ps_name": fallback}
    ps_name = _parse_ps_name_content(result.get("content")) or fallback
    if business_type == "service" and ps_name and not ps_name.endswith("服务"):
        ps_name = f"{ps_name}服务"
    return {"success": True, "ps_name": ps_name}


def _generate_relation_result(row):
    ip_name, contract_content = _relation_source_details(row)
    if not ip_name and not contract_content:
        return {"success": False, "error": "请先填写专利名称或销售合同内容（合同关键词/摘要）"}

    context = {
        "研发活动": str(row.get("rd_activity") or "").strip(),
        "研发周期": str(row.get("rd_period") or "").strip(),
        "相关知识产权名称": ip_name,
        "相关知识产权授权号": str(row.get("ip_auth_no") or "").strip(),
        "对应PS编号": str(row.get("ps_code") or "").strip(),
        "对应PS名称": str(row.get("ps_name") or "").strip(),
        "销售合同摘要": str(row.get("sales_contract_summary") or "").strip(),
        "销售合同关键词": str(row.get("sales_contract_keywords") or "").strip(),
    }
    context_text = "\n".join(f"{key}：{value}" for key, value in context.items() if value)
    source_instruction = (
        "专利名称和销售合同内容均已提供，成果名称必须同时体现技术来源与合同所涉产品/服务或应用场景。"
        if ip_name and contract_content else
        "仅提供了专利名称，成果名称只能依托该专利体现的技术内容生成，不得补充未提供的合同或市场信息。"
        if ip_name else
        "仅提供了销售合同内容，成果名称只能依托合同中明确的产品、服务、应用场景或交付内容生成，不得虚构专利或授权信息。"
    )
    prompt = f"""请根据以下 RD-IP-PS 关系信息，生成高新技术企业申报材料中的科技成果名称。

{context_text}

生成要求：
1. {source_instruction}
2. 成果名称要能体现从研发活动到成果转化和 PS 名称之间的关联性，通常是产品/工艺/系统/材料/服务类专业名词。
3. 不要写成完整句子，不要虚构具体专利号、性能参数、检测数据、未给出的企业事实。
4. 只输出 JSON，不要 markdown，不要解释。

JSON 格式：
{{"result_name": "成果名称"}}"""
    result = call_llm([
        {"role": "system", "content": "你是高新技术企业认定申报顾问，擅长基于专利内容和销售合同内容生成科技成果名称。只能使用已提供的事实依据；两类依据同时存在时需要综合使用。输出必须是可解析 JSON。"},
        {"role": "user", "content": prompt},
    ], temperature=0.2, max_tokens=400, timeout=40, max_attempts=2)
    if not result.get("success"):
        return {"success": False, "error": result.get("error") or "AI 生成失败"}

    data = _parse_relation_result_content(result.get("content"))
    fallback = _fallback_relation_result(row)
    result_name = str(data.get("result_name") or fallback.get("result_name") or "").strip()
    if not result_name:
        return {"success": False, "error": "AI 返回内容不完整，请重试"}
    return {"success": True, "result_name": result_name}


def _merge_relation_fields(data):
    relation_table = data.get("gaoxin_relation_table") or {}
    relation = relation_table.get("rows") or []
    tech_field_path = str(relation_table.get("tech_field_path") or "").strip()
    if not relation:
        if tech_field_path:
            merged = dict(data)
            merged["tech_field"] = tech_field_path
            return merged
        return data

    merged = dict(data)
    if tech_field_path:
        merged["tech_field"] = tech_field_path
    rds = []
    ips = []
    pss = []
    seen_rd = set()
    seen_ip = set()
    seen_ps = set()

    for row in relation:
        rd_key = row.get("rd_code") or row.get("rd_activity")
        if rd_key and rd_key not in seen_rd:
            seen_rd.add(rd_key)
            rds.append(row)

        ip_key = row.get("ip_code") or row.get("ip_auth_no") or row.get("ip_name")
        if ip_key and ip_key not in seen_ip:
            seen_ip.add(ip_key)
            ips.append(row)

        ps_key = row.get("ps_code") or row.get("ps_name")
        if ps_key and ps_key not in seen_ps:
            seen_ps.add(ps_key)
            pss.append(row)

    relation_fields = {
        "_relation_rd_count": len(rds),
        "_relation_ip_count": len(ips),
        "_relation_ps_count": len(pss),
        "_relation_cv_count": len(relation),
    }

    for i, row in enumerate(rds):
        rd_code = row.get("rd_code") or f"RD{str(i + 1).zfill(2)}"
        relation_fields[f"rd_{i}_no"] = rd_code
        relation_fields[f"rd_{i}_name"] = row.get("rd_activity", "")
        relation_fields[f"rd_{i}_period"] = row.get("rd_period", "")
        ip_labels = []
        ps_labels = []
        tech_values = []
        for rel in relation:
            if (rel.get("rd_code") or rel.get("rd_activity")) == (row.get("rd_code") or row.get("rd_activity")):
                ip_label = _relation_label(rel.get("ip_code", ""), rel.get("ip_name", ""), 15)
                if ip_label and ip_label not in ip_labels:
                    ip_labels.append(ip_label)
                ps_label = _relation_label(rel.get("ps_code", ""), rel.get("ps_name", ""))
                if ps_label and ps_label not in ps_labels:
                    ps_labels.append(ps_label)
                if rel.get("sales_contract_keywords") and rel.get("sales_contract_keywords") not in tech_values:
                    tech_values.append(rel.get("sales_contract_keywords"))
                elif rel.get("technology") and rel.get("technology") not in tech_values:
                    tech_values.append(rel.get("technology"))
        relation_fields[f"rd_{i}_ip_no"] = ip_labels
        relation_fields[f"rd_{i}_ps_no"] = ps_labels
        relation_fields[f"rd_{i}_result"] = _build_rd_stage_result(row, ip_labels, ps_labels, tech_values)
        if tech_field_path:
            relation_fields[f"rd_{i}_field"] = tech_field_path
        elif tech_values:
            relation_fields[f"rd_{i}_field"] = "；".join(tech_values)

    for i, row in enumerate(ips):
        relation_fields[f"ip_{i}_seq"] = row.get("ip_code") or f"IP{str(i + 1).zfill(2)}"
        relation_fields[f"ip_{i}_name"] = row.get("ip_name", "")
        relation_fields[f"ip_{i}_patent_no"] = row.get("ip_auth_no", "")
        if row.get("ip_auth_no", "").upper().startswith("ZL"):
            relation_fields[f"ip_{i}_status"] = "授权"

    for i, row in enumerate(pss):
        ps_code = row.get("ps_code") or f"PS{str(i + 1).zfill(2)}"
        relation_fields[f"ps_{i}_no"] = ps_code
        relation_fields[f"ps_{i}_name"] = row.get("ps_name", "")
        if tech_field_path:
            relation_fields[f"ps_{i}_field"] = tech_field_path
        relation_fields[f"ps_{i}_is_main"] = "yes"
        rd_labels = []
        ip_labels = []
        for rel in relation:
            if (rel.get("ps_code") or rel.get("ps_name")) == (row.get("ps_code") or row.get("ps_name")):
                rd_label = _relation_label(rel.get("rd_code", ""), rel.get("rd_activity", ""))
                if rd_label and rd_label not in rd_labels:
                    rd_labels.append(rd_label)
                ip_label = _relation_label(rel.get("ip_code", ""), rel.get("ip_name", ""), 15)
                if ip_label and ip_label not in ip_labels:
                    ip_labels.append(ip_label)
        relation_fields[f"ps_{i}_rds"] = rd_labels
        relation_fields[f"ps_{i}_ip_no"] = ip_labels

    for i, row in enumerate(relation):
        relation_fields[f"cv_{i}_rd"] = _relation_label(row.get("rd_code", ""), row.get("rd_activity", ""))
        relation_fields[f"cv_{i}_ip"] = _relation_label(row.get("ip_code", ""), row.get("ip_name", ""), 15)
        relation_fields[f"cv_{i}_ps"] = _relation_label(row.get("ps_code", ""), row.get("ps_name", ""))
        result_name = row.get("result_name", "")
        if result_name:
            relation_fields[f"cv_{i}_result_name"] = result_name
            patent_name, contract_text = _relation_source_details(row)
            source_parts = []
            if patent_name:
                source_parts.append(f"以{row.get('ip_code') or patent_name}的技术内容为依据")
            if contract_text:
                source_parts.append(f"结合销售合同中的{contract_text}等内容")
            source_clause = "，".join(source_parts) or "依据已填写的关联资料"
            temporal = project_temporal_context(row.get("rd_period") or row.get("year"))
            relation_fields[f"cv_{i}_desc"] = (
                f"登记成果名称：{result_name}。"
                f"该成果与{row.get('rd_code') or '相关RD项目'}、"
                f"{row.get('ip_code') or row.get('ip_name') or '相关知识产权'}及"
                f"{row.get('ps_code') or row.get('ps_name') or '相关PS'}建立资料对应关系，"
                f"{source_clause}。当前时间状态为“{temporal['status_display']}”；"
                "实际成果形成、转化实施和应用效果应依据研发、验收、合同、发票、测试及归档材料核实。"
            )

    for key, value in relation_fields.items():
        if value in (None, "", []):
            continue
        if re.match(r"^(cv_\d+_desc|rd_\d+_result)$", key) and merged.get(key):
            continue
        merged[key] = value
    merged = _expand_relation_references(merged, relation)
    ps_kinds = {
        infer_ps_kind(row.get("ps_name"))
        for row in pss
        if str(row.get("ps_name") or "").strip()
    }
    if len(ps_kinds) == 1:
        ps_kind = next(iter(ps_kinds))
        ps_name = next(
            (str(row.get("ps_name") or "").strip() for row in pss if row.get("ps_name")),
            "",
        )
        generated_prefixes = (
            "rd_",
            "ps_",
            "cv_",
            "innovation_",
            "achievement_evidence_",
            "attachment_ps_statement_",
            "attachment_hitech_product_summary",
        )
        for key, value in list(merged.items()):
            if isinstance(value, str) and key.startswith(generated_prefixes):
                merged[key] = normalize_ps_reference_text(value, ps_name, ps_kind)
    else:
        for index, row in enumerate(pss):
            ps_name = str(row.get("ps_name") or "").strip()
            for key in [
                f"ps_{index}_source",
                f"ps_{index}_tech",
                f"ps_{index}_advantage",
                f"ps_{index}_ip_support",
                f"attachment_ps_statement_{index}",
            ]:
                if isinstance(merged.get(key), str):
                    merged[key] = normalize_ps_reference_text(merged[key], ps_name)
    return merged


def _build_relation_table_summary(data):
    """Extract the RD-IP-PS links used to ground innovation narratives."""
    relation_table = data.get("gaoxin_relation_table") if isinstance(data, dict) else {}
    rows = relation_table.get("rows") if isinstance(relation_table, dict) else []
    rows = [row for row in rows if isinstance(row, dict)]

    def collect(key, fallback_key=""):
        values = []
        seen = set()
        for row in rows:
            value = str(row.get(key) or row.get(fallback_key) or "").strip()
            if value and value not in seen:
                seen.add(value)
                values.append(value)
        return values

    links = []
    for row in rows[:30]:
        link = {
            key: str(row.get(key) or "").strip()
            for key in (
                "year", "rd_code", "rd_activity", "rd_period", "ip_code", "ip_name",
                "ip_auth_no", "ps_code", "ps_name", "result_no", "result_name",
                "sales_contract_keywords",
            )
        }
        if any(link.values()):
            links.append(link)

    return {
        "rd_count": len(collect("rd_code", "rd_activity")),
        "ip_count": len(collect("ip_code", "ip_name")),
        "ps_count": len(collect("ps_code", "ps_name")),
        "result_count": len(collect("result_no", "result_name")),
        "rd_names": collect("rd_activity", "rd_code"),
        "ip_names": collect("ip_name", "ip_code"),
        "ps_names": collect("ps_name", "ps_code"),
        "result_names": collect("result_name", "result_no"),
        "links": links,
        "tech_field_path": str(relation_table.get("tech_field_path") or data.get("tech_field") or "").strip(),
    }


def _build_rd_stage_result(row, ip_labels, ps_labels, tech_values):
    rd_name = str(row.get("rd_activity") or row.get("rd_code") or "该研发活动").strip()
    period = str(row.get("rd_period") or "").strip()
    temporal = project_temporal_context(period)
    start_date = temporal["start_display"]
    end_date = temporal["end_display"]
    technology = "；".join(tech_values) if tech_values else "相关核心技术"
    ip_text = "；".join(ip_labels) if ip_labels else "相关知识产权"
    ps_text = "；".join(ps_labels) if ps_labels else "相关PS"

    if temporal["status"] == "已完成":
        return (
            f"1、项目计划周期为{start_date}至{end_date}，研发内容围绕{rd_name}开展；\n"
            f"2、项目已到计划结束时间，{technology}相关研发任务的实际完成情况应依据研发、测试和归档记录核对；\n"
            f"3、项目关联知识产权：{ip_text}。\n"
            f"项目与{ps_text}建立对应关系，成果是否已经形成并投入应用以实际证明材料为准。"
        )
    if temporal["status"] == "研发中":
        return (
            f"1、项目计划周期自{start_date}开始，围绕{rd_name}安排研发工作；\n"
            f"2、截至{temporal['as_of_display']}，项目正在推进{technology}相关研发任务，并阶段形成研发资料；\n"
            f"3、项目计划形成或支撑知识产权：{ip_text}。\n"
            f"后续拟将相关技术成果应用于{ps_text}，并按实际进展完善RD-IP-PS对应关系。"
        )
    if temporal["status"] == "计划中":
        return (
            f"1、项目计划于{start_date}启动，围绕{rd_name}组建研发团队；\n"
            f"2、项目拟开展{technology}相关研发任务，并按阶段形成研发资料；\n"
            f"3、项目拟形成或支撑知识产权：{ip_text}。\n"
            f"项目成果拟应用于{ps_text}，实际对应关系以实施和归档资料为准。"
        )
    return (
        f"1、项目周期待补充，围绕{rd_name}的启动时间和实施节点暂不作推断；\n"
        f"2、项目拟研究{technology}，具体进度和成果状态待结合实际记录补充；\n"
        f"3、知识产权及PS对应关系分别为：{ip_text}；{ps_text}。"
    )


@docgen_bp.route("/", methods=["GET"])
@login_required
def index():
    all_records = (
        ScoreRecord.query
        .join(Company)
        .filter(Company.user_id == current_user.id)
        .order_by(ScoreRecord.created_at.desc(), ScoreRecord.id.desc())
        .all()
    )
    companies = []
    seen = set()
    for record in all_records:
        key = record.company.name
        if key in seen:
            continue
        seen.add(key)
        record.company.latest_score = record
        companies.append(record.company)
        if len(companies) >= 20:
            break
    return render_template("application_index.html", companies=companies)


@docgen_bp.route("/fill/<int:company_id>", methods=["GET", "POST"])
@login_required
def fill(company_id):
    company = Company.query.filter_by(id=company_id, user_id=current_user.id).first_or_404()
    latest_score = (
        ScoreRecord.query
        .filter_by(company_id=company.id)
        .order_by(ScoreRecord.created_at.desc(), ScoreRecord.id.desc())
        .first()
    )
    is_gaoxin = company.app_type == "高新技术" or (latest_score and latest_score.score_type == "高新技术")
    if request.method == "GET" and is_gaoxin:
        return redirect(url_for("docgen.gaoxin_relation_table", company_id=company.id))

    score_data = {}
    if company.scores:
        score_data = json.loads(company.data_json or "{}")

    if request.method == "POST":
        form = request.form
        if company.app_type == "专精特新":
            app_data = _build_zhuanjing_data(company, form)
        else:
            app_data = _build_app_data(company, form)
        app_data.update(score_data)

        try:
            if company.app_type == "专精特新":
                output_path = generate_zhuanjing(app_data)
            else:
                output_path = generate(app_data)
        except Exception as e:
            flash(f"生成失败: {e}", "error")
            tpl = "application_zhuanjing_form.html" if company.app_type == "专精特新" else "application_form.html"
            return render_template(tpl, company=company, score_data=score_data)

        draft = _upsert_application_draft(company, app_data, output_path)
        db.session.commit()
        flash("申报书生成成功！", "success")
        return redirect(url_for("docgen.download", draft_id=draft.id))

    tpl = "application_zhuanjing_form.html" if company.app_type == "专精特新" else "application_form.html"
    return render_template(tpl, company=company, score_data=score_data)


@docgen_bp.route("/download/<int:draft_id>")
@login_required
def download(draft_id):
    draft = ApplicationDraft.query.get_or_404(draft_id)
    company = Company.query.get(draft.company_id)
    if company.user_id != current_user.id:
        flash("无权访问", "error")
        return redirect(url_for("docgen.index"))
    if not draft.docx_path or not os.path.exists(draft.docx_path):
        flash("文件不存在", "error")
        return redirect(url_for("docgen.index"))
    suffix = "专精特新中小企业申请书" if draft.app_type == "专精特新" else "高新技术企业认定申请书"
    return send_file(draft.docx_path, as_attachment=True, download_name=f"{company.name}_{suffix}.docx")


@docgen_bp.route("/history", methods=["GET"])
@login_required
def history():
    drafts = (
        ApplicationDraft.query
        .join(Company)
        .filter(Company.user_id == current_user.id)
        .order_by(ApplicationDraft.created_at.desc(), ApplicationDraft.id.desc())
        .all()
    )
    return render_template("application_history.html", drafts=drafts)


@docgen_bp.route("/assistant", methods=["GET"])
@login_required
def assistant():
    """兼容旧入口，助手和体检统一进入申报评估。"""
    return redirect(url_for("docgen.assessment", **request.args))


@docgen_bp.route("/assistant/brief", methods=["GET"])
@login_required
def assistant_brief():
    """兼容旧入口，助手和体检统一进入申报评估。"""
    return redirect(url_for("docgen.assessment", **request.args))


def _indexed_form_count(form, prefix: str, suffix: str, minimum: int = 0) -> int:
    max_index = -1
    marker = f"_{suffix}"
    for key in form.keys():
        if not key.startswith(prefix) or not key.endswith(marker):
            continue
        raw = key[len(prefix):-len(marker)]
        if raw.isdigit():
            max_index = max(max_index, int(raw))
    return max(minimum, max_index + 1)


def _build_app_data(company, form):
    data = {"company_name": company.name,
            "company_english_name": _company_english_name(company, _load_company_data(company)),
            "province": form.get("province", ""),
            "city": form.get("city", ""), "tech_field": form.get("tech_field", "")}
    ip_list = []
    for i in range(_indexed_form_count(form, "ip_", "name", 10)):
        name = (form.get(f"ip_{i}_name") or form.get(f"ip_name_{i}") or "").strip()
        if name:
            patent_no = (
                form.get(f"ip_{i}_patent_no", "")
                or form.get(f"ip_{i}_no", "")
                or form.get(f"ip_{i}_auth_no", "")
                or form.get(f"ip_auth_no_{i}", "")
            )
            ip_list.append({
                "seq": form.get(f"ip_{i}_seq", "") or str(i + 1),
                "name": name,
                "type": form.get(f"ip_{i}_type", "") or form.get(f"ip_type_{i}", ""),
                "status": form.get(f"ip_{i}_status", ""),
                "patent_no": patent_no,
                "app_date": form.get(f"ip_{i}_app_date", ""),
                "applicant": form.get(f"ip_{i}_applicant", ""),
                "grant_date": form.get(f"ip_{i}_date", "") or form.get(f"ip_date_{i}", ""),
                "auth_no": patent_no,
                "date": form.get(f"ip_{i}_date", "") or form.get(f"ip_date_{i}", ""),
                "method": form.get(f"ip_{i}_method", "") or form.get(f"ip_method_{i}", "") or "自主研发",
            })
    data["ip_list"] = ip_list
    def form_int(*keys):
        for key in keys:
            value = form.get(key)
            if value not in (None, ""):
                try:
                    return int(float(value))
                except (TypeError, ValueError):
                    return 0
        return 0

    data["ip_class1_count"] = form_int("ip_class1_count")
    data["ip_class2_count"] = form_int("ip_class2_count")
    data["staff_total"] = form_int("staff_total")
    data["tech_staff"] = form_int("tech_staff")
    data["hr_detail"] = {
        "onjob": form_int("hr_fulltime", "hr_onjob"), "parttime": form_int("hr_parttime"),
        "temp": form_int("hr_temp"), "foreign": form_int("hr_foreign"),
        "returnee": form_int("hr_returnee"), "talent_plan": form_int("hr_talent_plan"),
    }
    data["hr_edu"] = {"博士": form_int("edu_phd", "hr_phd"), "硕士": form_int("edu_master", "hr_master"),
                      "本科": form_int("edu_bachelor", "hr_bachelor"), "大专及以下": form_int("edu_diploma", "hr_college")}
    data["hr_title"] = {"高级职称": form_int("title_senior", "hr_title_senior"), "中级职称": form_int("title_mid", "hr_title_mid"),
                        "初级职称": form_int("title_junior", "hr_title_junior"), "高级技工": form_int("title_tech", "hr_title_tech")}
    data["hr_age"] = {"30及以下": form_int("age_30", "hr_age_30"), "31-40": form_int("age_31_40", "hr_age_40"),
                      "41-50": form_int("age_41_50", "hr_age_50"), "51及以上": form_int("age_51", "hr_age_51")}
    data["year1"] = "2023"; data["year2"] = "2024"; data["year3"] = "2025"
    data["year1_label"] = "2023"; data["year2_label"] = "2024"; data["year3_label"] = "2025"
    data = _sync_gaoxin_finance_years({**data, **form})
    for field in ["net_assets", "sales", "profit"]:
        for y in ["year1", "year2", "year3"]:
            form_value = form.get(f"{y}_{field}") or form.get(f"fin_{y}_{field}")
            if form_value not in (None, ""):
                data[f"{y}_{field}"] = float(form_value or 0)
    data["rd_total_3y"] = float(form.get("fin_rd_total", 0) or 0)
    data["rd_domestic"] = float(form.get("fin_rd_domestic", 0) or 0)
    data["rd_basic"] = float(form.get("fin_rd_basic", 0) or 0)
    data["revenue_1y"] = float(form.get("fin_revenue", 0) or 0)
    data["hitech_revenue_1y"] = float(form.get("fin_hitech_revenue", 0) or 0)
    data["no_violation"] = form.get("no_violation", "□否")
    for f in ["innovation_ip_role", "innovation_transform", "innovation_rd_mgmt", "innovation_talent"]:
        data[f] = form.get(f, "")
    standards = []
    for i in range(3):
        name = (form.get(f"std_{i}_name") or form.get(f"std_name_{i}") or "").strip()
        if name:
            standards.append({
                "name": name,
                "no": form.get(f"std_{i}_no", "") or form.get(f"std_no_{i}", ""),
                "level": form.get(f"std_{i}_level", "") or form.get(f"std_level_{i}", ""),
                "role": form.get(f"std_{i}_role", "") or form.get(f"std_role_{i}", ""),
            })
    data["standards"] = standards
    return data


def _build_zhuanjing_data(company, form):
    data = {"company_name": company.name,
            "company_english_name": _company_english_name(company, _load_company_data(company)),
            "province": form.get("province", ""),
            "city": form.get("city", ""), "address": form.get("address", ""),
            "zipcode": form.get("zipcode", ""), "legal_rep": form.get("legal_rep", ""),
            "shareholder": form.get("shareholder", ""), "actual_controller": form.get("actual_controller", ""),
            "controller_nationality": form.get("controller_nationality", ""),
            "contact": form.get("contact", ""), "phone": form.get("phone", ""),
            "mobile": form.get("mobile", ""), "fax": form.get("fax", ""),
            "email": form.get("email", ""), "register_date": form.get("register_date", ""),
            "register_capital": form.get("register_capital", ""), "credit_code": form.get("credit_code", ""),
            "industry_code": form.get("industry_code", ""), "sub_industry_code": form.get("sub_industry_code", "")}
    fin_fields = ["employees", "rd_staff", "revenue", "main_revenue", "revenue_growth",
                  "cost", "main_cost", "product_cost", "sales_expense", "admin_expense",
                  "profit", "net_profit", "net_profit_growth", "assets", "net_assets",
                  "liabilities", "debt_ratio", "tax", "equity_financing", "valuation",
                  "bank_loan", "domestic_bond", "foreign_bond"]
    for field in fin_fields:
        for yr in ["2023", "2024", "2025"]:
            data[f"fin_{yr}_{field}"] = form.get(f"fin_{yr}_{field}", "")
    for yr in ["2023", "2024", "2025"]:
        data[f"fin_{yr}_rd_expense"] = form.get(f"fin_{yr}_rd_expense", "")
    data["audit_report_code"] = form.get("audit_report_code", "")
    data["market_years"] = form.get("market_years", "")
    data["main_revenue_ratio"] = form.get("main_revenue_ratio", "")
    data["revenue_cagr"] = form.get("revenue_cagr", "")
    for i in range(3):
        name = form.get(f"product_{i}_name", "").strip()
        rev = form.get(f"product_{i}_revenue", "").strip()
        if name or rev:
            data[f"product_{i}"] = {"name": name, "revenue": rev}
    data["std_international"] = form.get("std_international", "")
    data["std_national"] = form.get("std_national", "")
    data["std_industry"] = form.get("std_industry", "")
    data["std_names"] = form.get("std_names", "")
    rd = {}
    for inst in ["tech_academy", "tech_center", "eng_center", "design_center", "key_lab"]:
        rd[inst] = {"national": form.get(f"rd_inst_{inst}_national", ""),
                    "province": form.get(f"rd_inst_{inst}_province", ""),
                    "self_built": form.get(f"rd_inst_{inst}_self", "")}
    data["rd_institutions"] = rd
    data["has_academician_station"] = bool(form.get("has_academician_station"))
    data["has_postdoc_station"] = bool(form.get("has_postdoc_station"))
    data["partner_schools"] = form.get("partner_schools", "")
    for yr in ["2023", "2024", "2025"]:
        data[f"rd_ratio_{yr}"] = form.get(f"rd_ratio_{yr}", "")
        data[f"rd_staff_ratio_{yr}"] = form.get(f"rd_staff_ratio_{yr}", "")
    for f in ["market_position_desc", "export_amount", "brand_count", "brand_revenue",
              "industry_chain", "main_product_name", "main_product_category", "company_overview"]:
        data[f] = form.get(f, "")
    return data


@docgen_bp.route("/gaoxin_relation_table/<int:company_id>", methods=["GET", "POST"])
@login_required
def gaoxin_relation_table(company_id):
    company = Company.query.filter_by(id=company_id, user_id=current_user.id).first_or_404()

    if request.method == "POST":
        payload = _relation_payload()
        rows = _normalize_relation_rows(payload.get("rows", []))
        errors = _validate_relation_rows(rows)
        if errors:
            return jsonify({"ok": False, "errors": errors}), 400
        _save_relation_table(company, rows, payload.get("tech_field_path", ""))
        return jsonify({"ok": True})

    data = _load_company_data(company)
    relation_table = data.get("gaoxin_relation_table") or {}
    rows = relation_table.get("rows") or []
    tech_field_path = relation_table.get("tech_field_path", "") or next((str(row.get("tech_field_path") or "").strip() for row in rows if str(row.get("tech_field_path") or "").strip()), "")
    saved_attachments = json.dumps(
        data.get("gaoxin_attachments") or {},
        ensure_ascii=False,
        sort_keys=True,
    )
    attachments, contract_files = _relation_sales_contracts(data, company)
    contract_options = _relation_sales_contract_options(contract_files)
    normalized_attachments = json.dumps(attachments, ensure_ascii=False, sort_keys=True)
    if saved_attachments != normalized_attachments:
        data["gaoxin_attachments"] = attachments
        company.data_json = json.dumps(data, ensure_ascii=False)
        db.session.commit()

    ip_details = _load_ip_details(company)
    ip_options = _build_relation_ip_options(ip_details)
    tech_field_options = _load_high_tech_field_options()

    return render_template(
        "application_gaoxin_relation_table.html",
        company=company,
        rows=rows,
        ip_details=ip_details,
        ip_options=ip_options,
        contract_options=contract_options,
        tech_field_options=tech_field_options,
        tech_field_path=tech_field_path,
    )


@docgen_bp.route("/health/<int:company_id>")
@login_required
def health_check(company_id):
    """旧体检入口兼容跳转到合并后的申报评估页。"""
    Company.query.filter_by(id=company_id, user_id=current_user.id).first_or_404()
    return redirect(url_for("docgen.assessment", company_id=company_id))


@docgen_bp.route("/assessment")
@login_required
def assessment():
    """评分和申报书资料齐备后，统一展示评分分析与申报体检。"""
    companies = (
        Company.query
        .filter_by(user_id=current_user.id)
        .order_by(Company.created_at.desc(), Company.id.desc())
        .all()
    )
    company_id = request.args.get("company_id", type=int)
    company = (
        Company.query.filter_by(id=company_id, user_id=current_user.id).first()
        if company_id
        else (companies[0] if companies else None)
    )
    if not company:
        return render_template(
            "assessment.html",
            companies=[],
            company=None,
            input_state={"ready": False, "missing": []},
            score=None,
            draft=None,
            analysis=None,
            health=None,
            data={},
            application_url=url_for("docgen.index"),
        )

    data = _load_company_data(company)
    input_state = _assessment_input_state(company, data)
    score = input_state["score"]
    draft = input_state["draft"]
    analysis = _score_analysis_from_record(score, data) if input_state["ready"] else None
    health = None
    if input_state["ready"] and (
        company.app_type == "高新技术" or (score and score.score_type == "高新技术")
    ):
        health = _company_health_check(company, data)

    company_cards = []
    for item in companies:
        item_state = _assessment_input_state(item)
        company_cards.append({
            "company": item,
            "score": item_state["score"],
            "ready": item_state["ready"],
            "application_ready": item_state["application_ready"],
        })

    if company.app_type == "高新技术":
        application_url = url_for("docgen.gaoxin_relation_table", company_id=company.id)
    else:
        application_url = url_for("docgen.fill", company_id=company.id)

    return render_template(
        "assessment.html",
        companies=company_cards,
        company=company,
        input_state=input_state,
        score=score,
        draft=draft,
        analysis=analysis,
        health=health,
        data=data,
        application_url=application_url,
    )


@docgen_bp.route("/health")
@login_required
def health_index():
    """旧体检入口兼容跳转到合并后的申报评估页。"""
    company = (
        Company.query
        .filter_by(user_id=current_user.id, app_type="高新技术")
        .order_by(Company.created_at.desc(), Company.id.desc())
        .first()
    )
    if not company:
        flash("暂无高新技术企业资料，请先完成评分。", "error")
        return redirect(url_for("scoring.gaoxin"))
    return redirect(url_for("docgen.assessment", company_id=company.id))


@docgen_bp.route("/health/<int:company_id>/json")
@login_required
def health_check_json(company_id):
    company = Company.query.filter_by(id=company_id, user_id=current_user.id).first_or_404()
    data = _load_company_data(company)
    input_state = _assessment_input_state(company, data)
    if not input_state["ready"]:
        return jsonify({
            "ok": True,
            "ready": False,
            "missing": input_state["missing"],
            "health": None,
        })
    return jsonify({
        "ok": True,
        "ready": True,
        "missing": [],
        "health": _company_health_check(company, data),
    })


@docgen_bp.route("/gaoxin_relation_table/<int:company_id>/sales_contract_keywords", methods=["POST"])
@login_required
def gaoxin_relation_table_sales_contract_keywords(company_id):
    company = Company.query.filter_by(id=company_id, user_id=current_user.id).first_or_404()
    payload = request.get_json(silent=True) or {}
    file_id = str(payload.get("file_id") or "").strip()
    if not file_id:
        return jsonify({"ok": False, "errors": ["请先选择销售合同"]}), 400
    row = payload.get("row") if isinstance(payload.get("row"), dict) else {}
    force = bool(payload.get("force"))
    data = _load_company_data(company)
    attachments, contracts = _relation_sales_contracts(data, company)
    file_meta = next(
        (
            item for item in contracts
            if str(item.get("id") or "").strip() == file_id
        ),
        None,
    )
    if not file_meta:
        return jsonify({"ok": False, "errors": ["所选销售合同不存在或已失效"]}), 404

    summary = str(file_meta.get("summary") or "").strip()
    keywords = str(file_meta.get("keywords") or "").strip()
    if force or not keywords:
        relative_path = str(file_meta.get("relative_path") or "").strip()
        path = _safe_attachment_path(relative_path) if relative_path else ""
        if not path or not os.path.exists(path):
            return jsonify({"ok": False, "errors": ["销售合同文件不存在，无法解析关键词"]}), 404
        text = _extract_pdf_text(path)
        extracted = _extract_sales_contract_info(
            text,
            row,
            file_meta.get("original_filename", ""),
        )
        summary = str(extracted.get("summary") or "").strip()
        keywords = str(extracted.get("keywords") or "").strip()
        file_meta["summary"] = summary
        file_meta["keywords"] = keywords

        required_contracts = ((data.get("required_materials") or {}).get("sales_contracts") or [])
        for item in required_contracts:
            if isinstance(item, dict) and str(item.get("id") or "").strip() == file_id:
                item["summary"] = summary
                item["keywords"] = keywords

        relation_rows = ((data.get("gaoxin_relation_table") or {}).get("rows") or [])
        for relation_row in relation_rows:
            if (
                isinstance(relation_row, dict)
                and str(relation_row.get("sales_contract_file_id") or "").strip() == file_id
            ):
                relation_row["sales_contract_summary"] = summary
                relation_row["sales_contract_keywords"] = keywords

    data["gaoxin_attachments"] = attachments
    company.data_json = json.dumps(data, ensure_ascii=False)
    db.session.commit()

    return jsonify({
        "ok": True,
        "file": {
            "id": file_id,
            "code": file_meta.get("contract_code", ""),
            "year": file_meta.get("year", ""),
            "original_filename": file_meta.get("original_filename", ""),
        },
        "summary": summary,
        "keywords": keywords,
    })


@docgen_bp.route("/gaoxin_relation_table/<int:company_id>/generate_result", methods=["POST"])
@login_required
def gaoxin_relation_table_generate_result(company_id):
    Company.query.filter_by(id=company_id, user_id=current_user.id).first_or_404()
    payload = request.get_json(silent=True) or {}
    generated = _generate_relation_result(payload.get("row") or {})
    if not generated.get("success"):
        return jsonify({"ok": False, "errors": [generated.get("error") or "AI 生成失败"]}), 400
    return jsonify({"ok": True, "result_name": generated["result_name"]})


@docgen_bp.route("/gaoxin_relation_table/<int:company_id>/generate_rd_activity", methods=["POST"])
@login_required
def gaoxin_relation_table_generate_rd_activity(company_id):
    Company.query.filter_by(id=company_id, user_id=current_user.id).first_or_404()
    payload = request.get_json(silent=True) or {}
    generated = _generate_rd_activity(payload.get("rows") or [])
    if not generated.get("success"):
        return jsonify({"ok": False, "errors": [generated.get("error") or "AI 生成失败"]}), 400
    return jsonify({"ok": True, "rd_activity": generated["rd_activity"]})


@docgen_bp.route("/gaoxin_relation_table/<int:company_id>/generate_ps_name", methods=["POST"])
@login_required
def gaoxin_relation_table_generate_ps_name(company_id):
    company = Company.query.filter_by(id=company_id, user_id=current_user.id).first_or_404()
    payload = request.get_json(silent=True) or {}
    generated = _generate_ps_name(company, payload.get("rows") or [])
    if not generated.get("success"):
        return jsonify({"ok": False, "errors": [generated.get("error") or "AI 生成失败"]}), 400
    return jsonify({"ok": True, "ps_name": generated["ps_name"]})


@docgen_bp.route("/gaoxin_relation_table/<int:company_id>/import", methods=["POST"])
@login_required
def gaoxin_relation_table_import(company_id):
    Company.query.filter_by(id=company_id, user_id=current_user.id).first_or_404()
    upload = request.files.get("file")
    if not upload or not upload.filename:
        return jsonify({"ok": False, "errors": ["请先选择要上传的 Excel 文件"]}), 400
    if not upload.filename.lower().endswith((".xlsx", ".xlsm")):
        return jsonify({"ok": False, "errors": ["请上传 .xlsx 或 .xlsm 格式的 Excel 文件"]}), 400

    try:
        imported_rows = import_relation_table(upload)
        tech_field_path = next((str(row.get("tech_field_path") or "").strip() for row in imported_rows if str(row.get("tech_field_path") or "").strip()), "")
        rows = _normalize_relation_rows(imported_rows)
    except Exception as exc:
        return jsonify({"ok": False, "errors": [f"Excel 识别失败：{exc}"]}), 400

    if not rows:
        return jsonify({"ok": False, "errors": ["未从 Excel 中识别到有效数据行"]}), 400

    return jsonify({"ok": True, "rows": rows, "tech_field_path": tech_field_path, "validation_errors": _validate_relation_rows(rows)})


@docgen_bp.route("/gaoxin_relation_table/<int:company_id>/export", methods=["POST"])
@login_required
def gaoxin_relation_table_export(company_id):
    company = Company.query.filter_by(id=company_id, user_id=current_user.id).first_or_404()
    payload = _relation_payload()
    rows = _normalize_relation_rows(payload.get("rows", []))
    errors = _validate_relation_rows(rows)
    if errors:
        return jsonify({"ok": False, "errors": errors}), 400

    _save_relation_table(company, rows, payload.get("tech_field_path", ""))
    stream = export_relation_table(rows, payload.get("tech_field_path", ""))
    return send_file(
        stream,
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        as_attachment=True,
        download_name=f"{company.name}_RD-IP-PS-成果关联详情表.xlsx",
    )


@docgen_bp.route("/gaoxin_book/<int:company_id>/hr_staff_template", methods=["GET"])
@login_required
def gaoxin_book_hr_staff_template(company_id):
    company = Company.query.filter_by(id=company_id, user_id=current_user.id).first_or_404()
    stream = _create_hr_staff_template()
    return send_file(
        stream,
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        as_attachment=True,
        download_name=f"{company.name}_企业人员情况模板.xlsx",
    )


@docgen_bp.route("/gaoxin_book/<int:company_id>/hr_staff_import", methods=["POST"])
@login_required
def gaoxin_book_hr_staff_import(company_id):
    company = Company.query.filter_by(id=company_id, user_id=current_user.id).first_or_404()
    upload = request.files.get("file")
    if not upload or not upload.filename:
        return jsonify({"ok": False, "error": "请先选择要上传的 Excel 文件"}), 400
    if not upload.filename.lower().endswith((".xlsx", ".xlsm")):
        return jsonify({"ok": False, "error": "请上传 .xlsx 或 .xlsm 格式的 Excel 文件"}), 400

    try:
        rows = _import_hr_staff_excel(upload)
    except Exception as exc:
        return jsonify({"ok": False, "error": f"Excel 识别失败：{exc}"}), 400
    if not rows:
        return jsonify({"ok": False, "error": "未识别到有效人员数据"}), 400

    data = _load_company_data(company)
    summary = _summarize_hr_staff_rows(rows)
    data["hr_staff_rows"] = rows
    data.update(summary)
    data["staff_total"] = summary["hr_total"]
    company.data_json = json.dumps(data, ensure_ascii=False)
    db.session.commit()
    return jsonify({"ok": True, "rows": rows, "summary": summary})


@docgen_bp.route("/gaoxin_book/<int:company_id>", methods=["GET", "POST"])
@login_required
def gaoxin_book(company_id):
    """高新技术企业认定申请书 — 官方格式网页版"""
    company = Company.query.filter_by(id=company_id, user_id=current_user.id).first_or_404()

    if request.method == "POST":
        form_data = request.form.to_dict()
        _save_gaoxin_book_data(company, form_data, session.get("last_finance_data"))
        flash("申请书已保存", "success")
        return redirect(url_for("docgen.gaoxin_book", company_id=company.id))

    auto_data = {}
    if company.data_json:
        try:
            auto_data = json.loads(company.data_json)
        except (json.JSONDecodeError, TypeError):
            pass

    auto_data = _merge_relation_fields(auto_data)

    if "last_finance_data" in session:
        auto_data.update(session["last_finance_data"])
    auto_data = _sync_gaoxin_finance_years(auto_data)

    # 获取 IP 明细 — 优先从 DB 加载
    ip_details = session.get("ip_certificates", [])
    if not ip_details and company.ip_certs_json:
        try:
            ip_details = json.loads(company.ip_certs_json)
            session["ip_certificates"] = ip_details
        except (json.JSONDecodeError, TypeError):
            pass

    return render_template(
        "application_gaoxin_book.html",
        company=company,
        auto_data=auto_data,
        ip_details=ip_details,
        relation_summary=_build_relation_table_summary(auto_data),
        system_as_of_date=system_today().isoformat(),
    )


@docgen_bp.route("/gaoxin_attachments/<int:company_id>", methods=["GET", "POST"])
@login_required
def gaoxin_attachments(company_id):
    """高新申报附件制作页。"""
    company = Company.query.filter_by(id=company_id, user_id=current_user.id).first_or_404()
    if request.method == "POST":
        _save_gaoxin_attachment_form_data(company, request.form.to_dict())
        flash("附件信息已保存", "success")
        return redirect(url_for("docgen.gaoxin_attachments", company_id=company.id))

    data = _load_company_data(company)
    health_warning = _gaoxin_export_health_warning(
        company,
        _company_health_check(company, data),
    )
    auto_data = _sync_gaoxin_finance_years(_merge_relation_fields(data))
    auto_data["attachment_rd_staff_rows"] = _collect_december_2025_rd_staff_rows(auto_data)
    _sync_staff_month_counts(auto_data, auto_data["attachment_rd_staff_rows"])
    auto_data["attachment_staff_statement_generated"] = _build_staff_statement_from_rd_list(
        auto_data,
        company.name,
    )
    auto_data.pop("attachment_staff_difference_note", None)
    achievement_rows = _collect_achievement_transform_rows(auto_data)
    rd_project_rows = _collect_rd_project_rows(auto_data)
    for project in rd_project_rows:
        project_index = project["index"]
        project["application"] = str(auto_data.get(f"attachment_rd_project_{project_index}_application") or "").strip()
    for achievement in achievement_rows:
        achievement_index = achievement["index"]
        achievement["materials"] = {
            "description": str(auto_data.get(f"cv_{achievement_index}_desc") or "").strip(),
            "test_report": str(auto_data.get(_achievement_evidence_field_name(achievement_index, "test_report")) or "").strip(),
            "user_report": str(auto_data.get(_achievement_evidence_field_name(achievement_index, "user_report")) or "").strip(),
        }
    export_fingerprint = _gaoxin_attachment_export_fingerprint(company, data)
    latest_export_job = (
        ExportJob.query
        .filter_by(
            company_id=company.id,
            user_id=current_user.id,
            job_type=GAOXIN_ATTACHMENT_EXPORT_JOB_TYPE,
            fingerprint=export_fingerprint,
        )
        .filter(ExportJob.status.in_(("queued", "running", "ready")))
        .order_by(ExportJob.created_at.desc())
        .first()
    )
    return render_template(
        "application_gaoxin_attachments.html",
        company=company,
        auto_data=auto_data,
        attachment_sections=GAOXIN_ATTACHMENT_SECTIONS,
        attachments=_load_gaoxin_attachments_from_data(data),
        ip_attachment=_ip_attachment_context(company, data),
        achievement_rows=achievement_rows,
        rd_project_rows=rd_project_rows,
        rd_project_ai_template=RD_PROJECT_APPLICATION_TEMPLATE,
        latest_export_job=(
            _export_job_payload(latest_export_job)["job"]
            if latest_export_job
            else None
        ),
        export_health_warning=health_warning,
    )


@docgen_bp.route("/gaoxin_attachments/<int:company_id>/achievement_material/<int:achievement_index>", methods=["POST"])
@login_required
def gaoxin_achievement_material_save(company_id, achievement_index):
    """保存批量生成的一项成果材料。"""
    company = Company.query.filter_by(id=company_id, user_id=current_user.id).first_or_404()
    data = _load_company_data(company)
    auto_data = _sync_gaoxin_finance_years(_merge_relation_fields(data))
    achievements = _collect_achievement_transform_rows(auto_data)
    if achievement_index < 0 or achievement_index >= len(achievements):
        return jsonify({"ok": False, "error": "成果不存在或已被删除"}), 404

    payload = request.get_json(silent=True) or {}
    material_key = str(payload.get("material") or "").strip()
    content = str(payload.get("content") or "").strip()
    field_names = {
        "description": f"cv_{achievement_index}_desc",
        "test_report": _achievement_evidence_field_name(achievement_index, "test_report"),
        "user_report": _achievement_evidence_field_name(achievement_index, "user_report"),
    }
    if material_key not in field_names:
        return jsonify({"ok": False, "error": "不支持的成果材料类型"}), 400
    if not content:
        return jsonify({"ok": False, "error": "生成内容为空，请重试"}), 400

    achievement = achievements[achievement_index]
    data[field_names[material_key]] = normalize_ps_reference_text(
        content,
        achievement.get("ps_name") or achievement.get("ps"),
        achievement.get("ps_kind"),
    )
    if material_key == "description":
        data[f"cv_{achievement_index}_result_name"] = achievement.get("result_name", "")
        data[f"cv_{achievement_index}_rd"] = _relation_label(achievement.get("rd_code", ""), achievement.get("rd_name", ""))
        data[f"cv_{achievement_index}_ip"] = achievement.get("ip", "")
        data[f"cv_{achievement_index}_ps"] = achievement.get("ps", "")
    else:
        data[f"achievement_evidence_{achievement_index}_result_name"] = achievement.get("result_name", "")
        data[f"achievement_evidence_{achievement_index}_result_no"] = achievement.get("result_no", "")
    company.data_json = json.dumps(data, ensure_ascii=False)
    db.session.commit()
    return jsonify({"ok": True})


@docgen_bp.route("/gaoxin_attachments/<int:company_id>/achievement_summary")
@login_required
def gaoxin_achievement_summary(company_id):
    """成果转化汇总表。"""
    company = Company.query.filter_by(id=company_id, user_id=current_user.id).first_or_404()
    data = _load_company_data(company)
    auto_data = _sync_gaoxin_finance_years(_merge_relation_fields(data))
    achievements = _collect_achievement_transform_rows(auto_data)
    return render_template(
        "application_gaoxin_achievement_summary.html",
        company=company,
        auto_data=auto_data,
        achievements=achievements,
    )


@docgen_bp.route("/gaoxin_attachments/<int:company_id>/achievement_summary/<int:achievement_index>", methods=["GET", "POST"])
@login_required
def gaoxin_achievement_description(company_id, achievement_index):
    """单条成果转化描述撰写页。"""
    company = Company.query.filter_by(id=company_id, user_id=current_user.id).first_or_404()
    data = _load_company_data(company)
    auto_data = _sync_gaoxin_finance_years(_merge_relation_fields(data))
    achievements = _collect_achievement_transform_rows(auto_data)
    if achievement_index < 0 or achievement_index >= len(achievements):
        abort(404)
    achievement = achievements[achievement_index]
    field_name = f"cv_{achievement_index}_desc"

    if request.method == "POST":
        data[field_name] = normalize_ps_reference_text(
            request.form.get(field_name, ""),
            achievement.get("ps_name") or achievement.get("ps"),
            achievement.get("ps_kind"),
        )
        data[f"cv_{achievement_index}_result_name"] = achievement["result_name"]
        data[f"cv_{achievement_index}_rd"] = _relation_label(achievement.get("rd_code", ""), achievement.get("rd_name", ""))
        data[f"cv_{achievement_index}_ip"] = achievement.get("ip", "")
        data[f"cv_{achievement_index}_ps"] = achievement.get("ps", "")
        company.data_json = json.dumps(data, ensure_ascii=False)
        db.session.commit()
        flash("成果转化描述已保存", "success")
        return redirect(url_for("docgen.gaoxin_achievement_description", company_id=company.id, achievement_index=achievement_index))

    return render_template(
        "application_gaoxin_achievement_description.html",
        company=company,
        auto_data=auto_data,
        achievement=achievement,
        achievement_index=achievement_index,
        field_name=field_name,
    )


@docgen_bp.route("/gaoxin_attachments/<int:company_id>/achievement_summary/<int:achievement_index>/pdf")
@login_required
def gaoxin_achievement_description_pdf(company_id, achievement_index):
    """预览转化结果描述 PDF。"""
    company = Company.query.filter_by(id=company_id, user_id=current_user.id).first_or_404()
    data = _load_company_data(company)
    auto_data = _sync_gaoxin_finance_years(_merge_relation_fields(data))
    achievements = _collect_achievement_transform_rows(auto_data)
    if achievement_index < 0 or achievement_index >= len(achievements):
        abort(404)
    achievement = achievements[achievement_index]
    html = render_template(
        "application_gaoxin_achievement_evidence_print.html",
        company=company,
        company_english_name=_company_english_name(company, data),
        achievement=achievement,
        evidence_type={"title": "转化结果描述"},
        content_text=achievement.get("desc", ""),
    )
    return _render_html_pdf(
        html,
        f"{company.name}_{achievement.get('result_no') or '成果'}_转化结果描述.pdf",
        "docgen.gaoxin_achievement_description",
        _header_company_name=company.name,
        _header_company_english_name=_company_english_name(company, data),
        company_id=company.id,
        achievement_index=achievement_index,
    )


@docgen_bp.route("/gaoxin_attachments/<int:company_id>/achievement_evidence/<int:achievement_index>", methods=["GET", "POST"])
@login_required
def gaoxin_achievement_evidence(company_id, achievement_index):
    """单条成果转化附件制作页。"""
    company = Company.query.filter_by(id=company_id, user_id=current_user.id).first_or_404()
    data = _load_company_data(company)
    auto_data = _sync_gaoxin_finance_years(_merge_relation_fields(data))
    achievements = _collect_achievement_transform_rows(auto_data)
    if achievement_index < 0 or achievement_index >= len(achievements):
        abort(404)
    achievement = achievements[achievement_index]

    if request.method == "POST":
        for item in GAOXIN_ACHIEVEMENT_EVIDENCE_EDIT_TYPES:
            field_name = _achievement_evidence_field_name(achievement_index, item["key"])
            data[field_name] = normalize_ps_reference_text(
                request.form.get(field_name, ""),
                achievement.get("ps_name") or achievement.get("ps"),
                achievement.get("ps_kind"),
            )
        data[f"achievement_evidence_{achievement_index}_result_name"] = achievement.get("result_name", "")
        data[f"achievement_evidence_{achievement_index}_result_no"] = achievement.get("result_no", "")
        company.data_json = json.dumps(data, ensure_ascii=False)
        db.session.commit()
        flash("成果附件内容已保存", "success")
        return redirect(url_for("docgen.gaoxin_achievement_evidence", company_id=company.id, achievement_index=achievement_index))

    attachments = _sync_achievement_patent_cert_files(company, data, achievement_index, achievement)
    upload_evidence_types = []
    for item in GAOXIN_ACHIEVEMENT_EVIDENCE_UPLOAD_TYPES:
        upload_evidence_types.append({
            **item,
            "files": _achievement_evidence_files(attachments, achievement_index, item["key"]),
        })
    edit_evidence_types = []
    for item in GAOXIN_ACHIEVEMENT_EVIDENCE_EDIT_TYPES:
        edit_evidence_types.append({
            **item,
            "field_name": _achievement_evidence_field_name(achievement_index, item["key"]),
        })

    return render_template(
        "application_gaoxin_achievement_evidence.html",
        company=company,
        auto_data=auto_data,
        achievement=achievement,
        achievement_index=achievement_index,
        upload_evidence_types=upload_evidence_types,
        edit_evidence_types=edit_evidence_types,
    )


@docgen_bp.route("/gaoxin_attachments/<int:company_id>/achievement_evidence/<int:achievement_index>/<evidence_key>/pdf")
@login_required
def gaoxin_achievement_evidence_pdf(company_id, achievement_index, evidence_key):
    """预览检测报告或用户使用报告 PDF。"""
    evidence_type = _achievement_evidence_edit_map().get(evidence_key)
    if not evidence_type:
        abort(404)
    company = Company.query.filter_by(id=company_id, user_id=current_user.id).first_or_404()
    data = _load_company_data(company)
    auto_data = _sync_gaoxin_finance_years(_merge_relation_fields(data))
    achievements = _collect_achievement_transform_rows(auto_data)
    if achievement_index < 0 or achievement_index >= len(achievements):
        abort(404)
    achievement = achievements[achievement_index]
    field_name = _achievement_evidence_field_name(achievement_index, evidence_key)
    content_text = auto_data.get(field_name, "")
    html = render_template(
        "application_gaoxin_achievement_evidence_print.html",
        company=company,
        company_english_name=_company_english_name(company, data),
        achievement=achievement,
        evidence_type=evidence_type,
        content_text=content_text,
        test_report_fields=(
            _achievement_test_report_fields(content_text, company, achievement)
            if evidence_key == "test_report" else None
        ),
        user_report_fields=(
            _achievement_user_report_fields(content_text, company, achievement)
            if evidence_key == "user_report" else None
        ),
    )
    return _render_html_pdf(
        html,
        f"{company.name}_{achievement.get('result_no') or '成果'}_{evidence_type['title']}.pdf",
        "docgen.gaoxin_achievement_evidence",
        _header_company_name=company.name,
        _header_company_english_name=_company_english_name(company, data),
        company_id=company.id,
        achievement_index=achievement_index,
    )


@docgen_bp.route("/gaoxin_attachments/<int:company_id>/hitech_products/summary", methods=["GET", "POST"])
@login_required
def gaoxin_hitech_product_summary(company_id):
    """高新技术产品（服务）汇总表撰写页。"""
    company = Company.query.filter_by(id=company_id, user_id=current_user.id).first_or_404()
    data = _load_company_data(company)
    auto_data = _sync_gaoxin_finance_years(_merge_relation_fields(data))
    products = _collect_hitech_product_rows(auto_data)
    summary_labels = _hitech_product_summary_labels(products)
    field_name = "attachment_hitech_product_summary"

    if request.method == "POST":
        data[field_name] = _normalize_generated_ps_text(request.form.get(field_name, ""), products)
        company.data_json = json.dumps(data, ensure_ascii=False)
        db.session.commit()
        flash(f"{summary_labels['title']}已保存", "success")
        return redirect(url_for("docgen.gaoxin_hitech_product_summary", company_id=company.id))

    display_data = dict(auto_data)
    display_data[field_name] = _normalize_generated_ps_text(auto_data.get(field_name, ""), products)
    return render_template(
        "application_gaoxin_hitech_product_summary.html",
        company=company,
        auto_data=display_data,
        products=products,
        products_context=_hitech_product_context_text(products),
        summary_labels=summary_labels,
        field_name=field_name,
    )


@docgen_bp.route("/gaoxin_attachments/<int:company_id>/hitech_products/summary/pdf")
@login_required
def gaoxin_hitech_product_summary_pdf(company_id):
    """预览高新技术产品（服务）汇总表 PDF。"""
    company = Company.query.filter_by(id=company_id, user_id=current_user.id).first_or_404()
    data = _load_company_data(company)
    auto_data = _sync_gaoxin_finance_years(_merge_relation_fields(data))
    products = _collect_hitech_product_rows(auto_data)
    summary_labels = _hitech_product_summary_labels(products)
    summary_text = _normalize_generated_ps_text(
        auto_data.get("attachment_hitech_product_summary", ""),
        products,
    )
    html = render_template(
        "application_gaoxin_hitech_product_summary_print.html",
        company=company,
        company_english_name=_company_english_name(company, data),
        products=products,
        summary_labels=summary_labels,
        summary_text=summary_text,
        summary_html=_rd_project_application_html(summary_text),
    )

    return _render_html_pdf(
        html,
        f"{company.name}_高新技术产品服务汇总表.pdf",
        "docgen.gaoxin_hitech_product_summary",
        _header_company_name=company.name,
        _header_company_english_name=_company_english_name(company, data),
        company_id=company.id,
    )


@docgen_bp.route("/gaoxin_attachments/<int:company_id>/hitech_products/<int:product_index>/evidence")
@login_required
def gaoxin_hitech_product_evidence(company_id, product_index):
    """Upload supporting materials for one high-tech product."""
    company = Company.query.filter_by(id=company_id, user_id=current_user.id).first_or_404()
    data = _load_company_data(company)
    auto_data = _sync_gaoxin_finance_years(_merge_relation_fields(data))
    products = _collect_hitech_product_rows(auto_data)
    if product_index < 0 or product_index >= len(products):
        abort(404)

    attachments = _load_gaoxin_attachments_from_data(data)
    evidence_types = [
        {**item, "files": _hitech_product_evidence_files(attachments, product_index, item["key"])}
        for item in GAOXIN_HITECH_PRODUCT_EVIDENCE_TYPES
    ]
    return render_template(
        "application_gaoxin_hitech_product_evidence.html",
        company=company,
        product=products[product_index],
        product_index=product_index,
        evidence_types=evidence_types,
    )


@docgen_bp.route("/gaoxin_attachments/<int:company_id>/hitech_products/ps/<int:product_index>", methods=["GET", "POST"])
@login_required
def gaoxin_ps_statement(company_id, product_index):
    """单个 PS 情况说明撰写页。"""
    company = Company.query.filter_by(id=company_id, user_id=current_user.id).first_or_404()
    data = _load_company_data(company)
    auto_data = _sync_gaoxin_finance_years(_merge_relation_fields(data))
    products = _collect_hitech_product_rows(auto_data)
    if product_index < 0 or product_index >= len(products):
        abort(404)
    product = products[product_index]
    field_name = f"attachment_ps_statement_{product_index}"
    statement_template = ps_statement_case_template(product["type_label"])

    if request.method == "POST":
        data[field_name] = request.form.get(field_name, "")
        data[f"attachment_ps_statement_{product_index}_ps_code"] = product["ps_code"]
        data[f"attachment_ps_statement_{product_index}_ps_name"] = product["ps_name"]
        company.data_json = json.dumps(data, ensure_ascii=False)
        db.session.commit()
        flash("PS情况说明已保存", "success")
        return redirect(url_for("docgen.gaoxin_ps_statement", company_id=company.id, product_index=product_index))

    return render_template(
        "application_gaoxin_ps_statement.html",
        company=company,
        auto_data=auto_data,
        product=product,
        product_index=product_index,
        field_name=field_name,
        ps_statement_template=statement_template,
        statement_text=normalize_ps_reference_text(
            auto_data.get(field_name, ""),
            product.get("ps_name"),
            product.get("ps_kind"),
        ),
    )


@docgen_bp.route("/gaoxin_attachments/<int:company_id>/hitech_products/ps/<int:product_index>/pdf")
@login_required
def gaoxin_ps_statement_pdf(company_id, product_index):
    """预览单个 PS 情况说明 PDF。"""
    company = Company.query.filter_by(id=company_id, user_id=current_user.id).first_or_404()
    data = _load_company_data(company)
    auto_data = _sync_gaoxin_finance_years(_merge_relation_fields(data))
    products = _collect_hitech_product_rows(auto_data)
    if product_index < 0 or product_index >= len(products):
        abort(404)
    product = products[product_index]
    field_name = f"attachment_ps_statement_{product_index}"
    statement_text = normalize_ps_reference_text(
        auto_data.get(field_name, ""),
        product.get("ps_name"),
        product.get("ps_kind"),
    )
    html = render_template(
        "application_gaoxin_ps_statement_print.html",
        company=company,
        company_english_name=_company_english_name(company, data),
        product=product,
        statement_text=statement_text,
        statement_html=_rd_project_application_html(statement_text),
    )

    return _render_html_pdf(
        html,
        f"{company.name}_{product.get('ps_code') or 'PS'}_情况说明.pdf",
        "docgen.gaoxin_ps_statement",
        _header_company_name=company.name,
        _header_company_english_name=_company_english_name(company, data),
        company_id=company.id,
        product_index=product_index,
    )


@docgen_bp.route("/gaoxin_attachments/<int:company_id>/rd_project/<int:project_index>", methods=["GET", "POST"])
@login_required
def gaoxin_rd_project_application(company_id, project_index):
    """单个 RD 科研项目书撰写页。"""
    company = Company.query.filter_by(id=company_id, user_id=current_user.id).first_or_404()
    data = _load_company_data(company)
    auto_data = _sync_gaoxin_finance_years(_merge_relation_fields(data))
    projects = _collect_rd_project_rows(auto_data)
    if project_index < 0 or project_index >= len(projects):
        abort(404)
    project = projects[project_index]
    field_name = f"attachment_rd_project_{project_index}_application"

    if request.method == "POST":
        data[field_name] = request.form.get(field_name, "")
        data[f"attachment_rd_project_{project_index}_year"] = project["year"]
        data[f"attachment_rd_project_{project_index}_rd_code"] = project["rd_code"]
        data[f"attachment_rd_project_{project_index}_rd_activity"] = project["rd_activity"]
        data[f"attachment_rd_project_{project_index}_rd_period"] = project["rd_period"]
        company.data_json = json.dumps(data, ensure_ascii=False)
        db.session.commit()
        flash("科研项目书已保存", "success")
        return redirect(url_for("docgen.gaoxin_rd_project_application", company_id=company.id, project_index=project_index))

    return render_template(
        "application_gaoxin_rd_project.html",
        company=company,
        auto_data=auto_data,
        project=project,
        project_index=project_index,
        field_name=field_name,
        rd_project_ai_template=RD_PROJECT_APPLICATION_TEMPLATE,
    )


@docgen_bp.route("/gaoxin_attachments/<int:company_id>/rd_project/<int:project_index>/book", methods=["POST"])
@login_required
def gaoxin_rd_project_book_save(company_id, project_index):
    """保存一键生成的单份 RD 科研项目书。"""
    company = Company.query.filter_by(id=company_id, user_id=current_user.id).first_or_404()
    data = _load_company_data(company)
    auto_data = _sync_gaoxin_finance_years(_merge_relation_fields(data))
    projects = _collect_rd_project_rows(auto_data)
    if project_index < 0 or project_index >= len(projects):
        return jsonify({"ok": False, "error": "科研项目不存在或已被删除"}), 404

    payload = request.get_json(silent=True) or {}
    content = str(payload.get("content") or "").strip()
    if not content:
        return jsonify({"ok": False, "error": "生成内容为空，请重试"}), 400

    project = projects[project_index]
    data[f"attachment_rd_project_{project_index}_application"] = content
    data[f"attachment_rd_project_{project_index}_year"] = project["year"]
    data[f"attachment_rd_project_{project_index}_rd_code"] = project["rd_code"]
    data[f"attachment_rd_project_{project_index}_rd_activity"] = project["rd_activity"]
    data[f"attachment_rd_project_{project_index}_rd_period"] = project["rd_period"]
    company.data_json = json.dumps(data, ensure_ascii=False)
    db.session.commit()
    return jsonify({"ok": True})


def _rd_project_application_sections(text):
    content = str(text or "").strip()
    marker_aliases = [
        ("background", ("立项背景与必要性", "项目背景与必要性", "立项目的")),
        ("technical_problems", ("拟解决的技术问题", "关键技术问题")),
        ("objectives", ("研发目标与考核指标", "研发内容及目标")),
        ("research_content", ("研发内容",)),
        ("technical_route", ("技术路线",)),
        ("innovation", ("创新点", "核心技术及创新点")),
        ("team", ("项目组织与任务分工", "团队与任务分工")),
        ("budget", ("经费预算", "研发预算")),
        ("progress", ("计划进度", "阶段计划与里程碑")),
        ("records", ("过程记录与质量控制", "研发过程记录")),
        ("outputs", ("项目成果", "预期成果与阶段成果", "预期成果")),
        ("relations", ("RD-IP-PS关联", "研发成果关联")),
        ("acceptance_comparison", ("验收指标对照", "指标完成情况")),
        ("acceptance_opinion", ("验收意见",)),
        ("acceptance_result", ("验收结论",)),
    ]
    result = {key: "" for key, _ in marker_aliases}
    alias_lookup = [
        (alias, key)
        for key, aliases in marker_aliases
        for alias in sorted(aliases, key=len, reverse=True)
    ]
    current_key = None
    buffers = {key: [] for key in result}
    for raw_line in content.splitlines():
        stripped = raw_line.strip()
        normalized = re.sub(r"^[一二三四五六七八九十\d]+[、.．）)]\s*", "", stripped)
        matched = False
        for alias, key in alias_lookup:
            if normalized == alias or normalized.startswith(alias + "：") or normalized.startswith(alias + ":"):
                current_key = key
                remainder = normalized[len(alias):].lstrip(" ：:")
                if remainder:
                    buffers[key].append(remainder)
                matched = True
                break
        if not matched and current_key and stripped:
            buffers[current_key].append(stripped)
    for key, lines in buffers.items():
        result[key] = "\n".join(lines).strip()
    result["purpose"] = result["background"]
    result["content_goal"] = result["objectives"]
    return result


def _rd_project_application_html(text):
    from html import escape

    blocks = []
    lines = str(text or "").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if line.startswith("|") and line.endswith("|") and i + 1 < len(lines):
            sep = lines[i + 1].strip()
            if sep.startswith("|") and set(sep.replace("|", "").replace("-", "").replace(":", "").strip()) == set():
                headers = [cell.strip() for cell in line.strip("|").split("|")]
                table_rows = []
                i += 2
                while i < len(lines) and lines[i].strip().startswith("|") and lines[i].strip().endswith("|"):
                    table_rows.append([cell.strip() for cell in lines[i].strip().strip("|").split("|")])
                    i += 1
                head_html = "".join(f"<th>{escape(cell)}</th>" for cell in headers)
                rows_html = "".join("<tr>" + "".join(f"<td>{escape(cell)}</td>" for cell in row) + "</tr>" for row in table_rows)
                blocks.append(f"<table class=\"content-table\"><thead><tr>{head_html}</tr></thead><tbody>{rows_html}</tbody></table>")
                continue
        if line.startswith("### "):
            blocks.append(f"<h3>{escape(line[4:].strip())}</h3>")
        elif line.startswith("## "):
            blocks.append(f"<h2>{escape(line[3:].strip())}</h2>")
        elif line:
            blocks.append(f"<p>{escape(line)}</p>")
        i += 1
    return "\n".join(blocks)


def _export_rd_project_application_text(project, saved_text):
    """Keep every RD project represented in the consolidated attachment export.

    Project books drafted in the attachment page take precedence.  When a book
    has not been drafted yet, the print template supplies the standard
    project-specific document sections from the RD project metadata.
    """
    temporal = project.get("temporal") or project_temporal_context(project.get("rd_period"))
    content = enforce_temporal_wording(saved_text, temporal).strip()
    if content:
        return content

    project_name = str(project.get("rd_activity") or "该研发项目").strip()
    purpose = project.get("purpose") or f"围绕“{project_name}”开展自主研发，解决实际业务和技术应用中的关键问题。"
    innovation = project.get("innovation") or "围绕技术方案、实现方法、系统适配和稳定运行开展改进，具体创新内容以项目研发记录为准。"
    if project.get("result"):
        result = project["result"]
    elif temporal["status"] == "已完成":
        result = "项目已到计划结束时间，实际成果及归档情况应依据研发、测试和验收资料核对。"
    elif temporal["status"] == "研发中":
        result = "项目正在按计划推进，阶段成果及实际形成情况应依据当前研发和测试记录填写。"
    elif temporal["status"] == "计划中":
        result = "项目预期成果按立项目标安排，实际成果待项目启动并形成研发记录后填写。"
    else:
        result = "项目周期和成果状态待补充，不推断已经启动、完成或形成成果。"
    return f"""立项背景与必要性
{purpose}

拟解决的技术问题
结合项目应用场景，重点解决方案设计、关键功能实现、系统适配、稳定运行和验证评价中的技术问题。

研发目标与考核指标
建立与项目研发内容对应的功能、性能、稳定性、适用性和成果归档指标；未提供的量化指标待补充，不作编造。

研发内容
开展需求分析、技术方案设计、关键模块研发、联调测试、问题整改、方案优化和成果归档。

技术路线
按照“需求分析—方案设计—研发实现—测试验证—迭代优化—成果归档”的路线推进，并保留各阶段过程记录。

创新点
{innovation}

项目组织与任务分工
由企业组织项目实施，项目负责人、研发人员和具体任务分工以实际立项及人员记录为准，缺失信息待补充。

经费预算
项目预算为{project.get('budget_display') or '待补充'}，具体费用按研发辅助账和企业财务制度据实归集。

计划进度
按{project.get('rd_period') or '待补充的项目周期'}依次推进项目准备、设计开发、测试优化和总结验收工作。

过程记录与质量控制
保存立项审批、任务分工、设计评审、测试记录、问题整改、阶段评审、变更记录和成果归档资料。

预期成果与阶段成果
{result}

RD-IP-PS关联
关联知识产权：{'；'.join(project.get('ip_labels') or []) or '待补充'}。
关联PS：{'；'.join(project.get('ps_labels') or []) or '待补充'}。

验收指标对照
按研发目标逐项核对计划指标、实际完成情况和证明材料；未提供的量化结果留空待补充。

验收意见
{temporal['acceptance_opinion']}

验收结论
{temporal['acceptance_result']}"""


@docgen_bp.route("/gaoxin_attachments/<int:company_id>/rd_project/<int:project_index>/pdf")
@login_required
def gaoxin_rd_project_application_pdf(company_id, project_index):
    """预览单个 RD 科研项目书 PDF。"""
    company = Company.query.filter_by(id=company_id, user_id=current_user.id).first_or_404()
    data = _load_company_data(company)
    auto_data = _sync_gaoxin_finance_years(_merge_relation_fields(data))
    projects = _collect_rd_project_rows(auto_data)
    if project_index < 0 or project_index >= len(projects):
        abort(404)
    project = projects[project_index]
    field_name = f"attachment_rd_project_{project_index}_application"
    application_text = _export_rd_project_application_text(project, auto_data.get(field_name, ""))
    html = render_template(
        "application_gaoxin_rd_project_print.html",
        company=company,
        company_english_name=_company_english_name(company, data),
        project=project,
        application_text=application_text,
        application_html=_rd_project_application_html(application_text),
        application_sections=_rd_project_application_sections(application_text),
    )

    return _render_html_pdf(
        html,
        f"{company.name}_{project.get('project_no') or 'RD'}_科研项目书.pdf",
        "docgen.gaoxin_rd_project_application",
        _header_company_name=company.name,
        _header_company_english_name=_company_english_name(company, data),
        _header_skip_first_page=True,
        company_id=company.id,
        project_index=project_index,
    )


@docgen_bp.route("/gaoxin_attachments/<int:company_id>/pdf/jobs", methods=["POST"])
@login_required
def gaoxin_attachments_pdf_job_create(company_id):
    company = Company.query.filter_by(
        id=company_id,
        user_id=current_user.id,
    ).first_or_404()
    data = _load_company_data(company)
    health = _company_health_check(company, data)
    health_warning = _gaoxin_export_health_warning(company, health)

    fingerprint = _gaoxin_attachment_export_fingerprint(company, data)
    existing = (
        ExportJob.query
        .filter_by(
            company_id=company.id,
            user_id=current_user.id,
            job_type=GAOXIN_ATTACHMENT_EXPORT_JOB_TYPE,
            fingerprint=fingerprint,
        )
        .filter(ExportJob.status.in_(("queued", "running", "ready")))
        .order_by(ExportJob.created_at.desc())
        .first()
    )
    if existing:
        if not _expire_stale_export_job(existing):
            payload = _export_job_payload(existing)
            payload["export_warning"] = health_warning
            payload["cache_hit"] = existing.status == "ready"
            return jsonify(payload), 200 if existing.status == "ready" else 202

    job = ExportJob(
        id=str(uuid.uuid4()),
        company_id=company.id,
        user_id=current_user.id,
        job_type=GAOXIN_ATTACHMENT_EXPORT_JOB_TYPE,
        fingerprint=fingerprint,
        status="queued",
        stage="等待生成",
        progress=2,
        download_name=f"高新技术企业认定附件制作_{company.name}.pdf",
    )
    db.session.add(job)
    db.session.commit()
    payload = _export_job_payload(job)
    payload["export_warning"] = health_warning
    return jsonify(payload), 202


@docgen_bp.route("/gaoxin_attachments/<int:company_id>/pdf/jobs/<job_id>")
@login_required
def gaoxin_attachments_pdf_job_status(company_id, job_id):
    job = _owned_export_job(company_id, job_id)
    _expire_stale_export_job(job)
    return jsonify(_export_job_payload(job))


@docgen_bp.route("/gaoxin_attachments/<int:company_id>/pdf/jobs/<job_id>/download")
@login_required
def gaoxin_attachments_pdf_job_download(company_id, job_id):
    job = _owned_export_job(company_id, job_id)
    if job.status != "ready" or not job.result_path:
        return jsonify({"ok": False, "error": "PDF 尚未生成完成"}), 409

    local_path = _gaoxin_attachment_export_local_path(job)
    ensure_local_file(local_path, job.result_path)
    if not os.path.isfile(local_path):
        _update_export_job(
            job,
            status="failed",
            stage="导出文件已失效",
            error="已生成的 PDF 文件无法读取，请重新发起导出。",
        )
        return jsonify({"ok": False, "error": job.error_message}), 410

    return send_file(
        local_path,
        mimetype="application/pdf",
        as_attachment=True,
        download_name=job.download_name
        or f"高新技术企业认定附件制作_{company_id}.pdf",
        conditional=True,
    )


@docgen_bp.route("/gaoxin_attachments/<int:company_id>/pdf", methods=["GET", "POST"])
@login_required
def gaoxin_attachments_pdf(company_id):
    """导出高新技术企业认定附件制作页面 PDF。"""
    export_started = time.perf_counter()
    app = current_app._get_current_object()
    company = Company.query.filter_by(id=company_id, user_id=current_user.id).first_or_404()
    data = _load_company_data(company)
    job = None
    job_id = str(request.args.get("job") or "").strip()
    if job_id:
        job = _owned_export_job(company.id, job_id)
        if job.status == "ready":
            return jsonify(_export_job_payload(job))
        if job.status == "running":
            return jsonify(_export_job_payload(job)), 202
        if job.status == "failed":
            return jsonify(_export_job_payload(job)), 409
        current_fingerprint = _gaoxin_attachment_export_fingerprint(company, data)
        if current_fingerprint != job.fingerprint:
            _update_export_job(
                job,
                status="failed",
                stage="材料已发生变化",
                error="任务创建后附件内容发生变化，请重新发起导出。",
            )
            return jsonify(_export_job_payload(job)), 409
        started_at = datetime.utcnow()
        claimed = (
            ExportJob.query
            .filter_by(
                id=job.id,
                company_id=company.id,
                user_id=current_user.id,
                job_type=GAOXIN_ATTACHMENT_EXPORT_JOB_TYPE,
                status="queued",
            )
            .update(
                {
                    ExportJob.status: "running",
                    ExportJob.stage: "正在整理附件材料",
                    ExportJob.progress: 6,
                    ExportJob.error_message: "",
                    ExportJob.started_at: started_at,
                    ExportJob.updated_at: started_at,
                },
                synchronize_session=False,
            )
        )
        db.session.commit()
        job = db.session.get(ExportJob, job.id)
        if not claimed:
            status_code = 409 if job.status == "failed" else 202
            return jsonify(_export_job_payload(job)), status_code
    auto_data = _sync_gaoxin_finance_years(_merge_relation_fields(data))
    auto_data["attachment_rd_staff_rows"] = _collect_december_2025_rd_staff_rows(auto_data)
    _sync_staff_month_counts(auto_data, auto_data["attachment_rd_staff_rows"])
    auto_data["attachment_staff_statement_generated"] = _build_staff_statement_from_rd_list(
        auto_data,
        company.name,
    )
    if job:
        _update_export_job(
            job,
            stage="正在生成附件目录与内部材料",
            progress=12,
        )
    export_sections = [section for section in GAOXIN_ATTACHMENT_SECTIONS if section["key"] != "commitment"]
    attachments = _load_gaoxin_attachments_from_data(data)
    ip_attachment = _ip_attachment_context(company, data)
    rd_projects = _collect_rd_project_rows(auto_data)
    products = _collect_hitech_product_rows(auto_data)
    achievements = _collect_achievement_transform_rows(auto_data)
    system_docs = _load_gaoxin_system_docs(company)
    html = render_template(
        "application_gaoxin_attachments_print.html",
        company=company,
        company_english_name=_company_english_name(company, data),
        auto_data=auto_data,
        attachment_sections=export_sections,
        attachments=attachments,
        ip_attachment=ip_attachment,
        exported_at=datetime.now().strftime("%Y-%m-%d %H:%M"),
    )

    try:
        with tempfile.TemporaryDirectory(prefix="gaoxin-pdf-export-") as export_dir:
            section_paths = {str(number): [] for number in range(2, 14)}
            portrait_documents = []
            standalone_documents = []
            attachment_references = []
            company_english_name = _company_english_name(company, data)

            def add_attachment_files(section_no, files, label):
                for file_meta in files or []:
                    reference = {
                        "relative_path": file_meta.get("relative_path", ""),
                        "label": label,
                        "uploaded_attachment": True,
                    }
                    attachment_references.append(reference)
                    section_paths[str(section_no)].append((reference, label))

            def add_standalone_document(section_no, document_html, label):
                document = {"html": document_html, "label": label}
                standalone_documents.append(document)
                section_paths[str(section_no)].append((document, label))

            def add_portrait_document(
                section_no,
                document_html,
                label,
                *,
                skip_header_first_page=False,
            ):
                if _generated_document_needs_landscape(document_html):
                    add_standalone_document(
                        section_no,
                        _ensure_landscape_page_rule(document_html),
                        label,
                    )
                    return
                document = {
                    "html": document_html,
                    "label": label,
                    "skip_header_first_page": skip_header_first_page,
                }
                portrait_documents.append(document)
                section_paths[str(section_no)].append((document, label))

            add_attachment_files("2", attachments.get("application_pdf", {}).get("files", []), "申请书")
            add_attachment_files("3", attachments.get("business_license", {}).get("files", []), "营业执照")
            ip_detail_html = render_template(
                "application_gaoxin_ip_detail_print.html",
                company=company,
                company_english_name=company_english_name,
                rows=ip_attachment.get("rows", []),
            )
            add_portrait_document("4", ip_detail_html, "知识产权明细表")
            add_attachment_files("4", _ordered_ip_certificate_files(attachments, ip_attachment.get("rows", [])), "知识产权证书")
            staff_tables_html = render_template(
                "application_gaoxin_staff_tables_print.html",
                company=company,
                company_english_name=company_english_name,
                auto_data=auto_data,
                rows=auto_data.get("attachment_rd_staff_rows", []),
            )
            add_portrait_document("5", staff_tables_html, "研发人员及科技人员统计表")
            staff_files = attachments.get("staff_statement", {}).get("files", [])
            for file_type, label in [
                ("social_2025_12", "2025年12月份社保单"),
                ("social_latest", "最新月份社保单"),
                ("education_certificate", "毕业证书"),
                ("title_certificate", "职称证书"),
            ]:
                add_attachment_files("5", [item for item in staff_files if item.get("staff_file_type") == file_type], label)

            for project in rd_projects:
                project_index = project["index"]
                application_text = _export_rd_project_application_text(
                    project,
                    auto_data.get(f"attachment_rd_project_{project_index}_application"),
                )
                project_html = render_template(
                    "application_gaoxin_rd_project_print.html",
                    company=company,
                    company_english_name=company_english_name,
                    project=project,
                    application_text=application_text,
                    application_html=_rd_project_application_html(application_text),
                    application_sections=_rd_project_application_sections(application_text),
                )
                add_portrait_document(
                    "6",
                    project_html,
                    f"科研项目书 {project.get('project_no') or project_index}",
                    skip_header_first_page=True,
                )

            for year in GAOXIN_FINANCE_YEARS:
                add_attachment_files("7", [
                    item for item in attachments.get("rd_expense_audit", {}).get("files", [])
                    if str(item.get("attachment_year") or "") == year
                ], f"{year}年研发费用专项审计报告")

            product_summary_html = render_template(
                "application_gaoxin_hitech_product_summary_print.html",
                company=company,
                company_english_name=company_english_name,
                products=products,
                summary_labels=_hitech_product_summary_labels(products),
                summary_text=_normalize_generated_ps_text(
                    auto_data.get("attachment_hitech_product_summary", ""),
                    products,
                ),
                summary_html=_rd_project_application_html(
                    _normalize_generated_ps_text(
                        auto_data.get("attachment_hitech_product_summary", ""),
                        products,
                    )
                ),
            )
            add_standalone_document("8", product_summary_html, "高新技术产品说明汇总表")
            for product in products:
                product_index = product["index"]
                statement_text = normalize_ps_reference_text(
                    auto_data.get(f"attachment_ps_statement_{product_index}", ""),
                    product.get("ps_name"),
                    product.get("ps_kind"),
                )
                ps_html = render_template(
                    "application_gaoxin_ps_statement_print.html",
                    company=company,
                    company_english_name=company_english_name,
                    product=product,
                    statement_text=statement_text,
                    statement_html=_rd_project_application_html(statement_text),
                )
                add_portrait_document("8", ps_html, f"{product.get('ps_code') or product_index}情况说明")
                for evidence in GAOXIN_HITECH_PRODUCT_EVIDENCE_TYPES:
                    add_attachment_files("8", _hitech_product_evidence_files(attachments, product_index, evidence["key"]), evidence["title"])

            add_attachment_files("9", attachments.get("hitech_income_audit", {}).get("files", []), "高新产品收入专项审计报告")

            achievement_summary_html = render_template(
                "application_gaoxin_achievement_summary_print.html",
                company=company,
                company_english_name=company_english_name,
                achievements=achievements,
            )
            add_standalone_document("10", achievement_summary_html, "成果转化汇总表")
            for achievement in achievements:
                achievement_index = achievement["index"]
                description_html = render_template(
                    "application_gaoxin_achievement_evidence_print.html",
                    company=company,
                    company_english_name=company_english_name,
                    achievement=achievement,
                    evidence_type={"key": "description", "title": "转化结果描述"},
                    content_text=achievement.get("desc", ""),
                )
                add_portrait_document("10", description_html, f"{achievement.get('result_no')}转化结果描述")
                for evidence in GAOXIN_ACHIEVEMENT_EVIDENCE_UPLOAD_TYPES:
                    add_attachment_files("10", _achievement_evidence_files(attachments, achievement_index, evidence["key"]), evidence["title"])
                for evidence in GAOXIN_ACHIEVEMENT_EVIDENCE_EDIT_TYPES:
                    content_text = str(auto_data.get(_achievement_evidence_field_name(achievement_index, evidence["key"])) or "").strip()
                    if not content_text:
                        continue
                    evidence_html = render_template(
                        "application_gaoxin_achievement_evidence_print.html",
                        company=company,
                        company_english_name=company_english_name,
                        achievement=achievement,
                        evidence_type=evidence,
                        content_text=content_text,
                        test_report_fields=(
                            _achievement_test_report_fields(content_text, company, achievement)
                            if evidence["key"] == "test_report" else None
                        ),
                        user_report_fields=(
                            _achievement_user_report_fields(content_text, company, achievement)
                            if evidence["key"] == "user_report" else None
                        ),
                    )
                    add_portrait_document("10", evidence_html, f"{achievement.get('result_no')}{evidence['title']}")

            system_summary_data = {
                key: _normalize_generated_ps_text(value, products)
                if isinstance(value, str)
                else value
                for key, value in auto_data.items()
            }
            system_summary_html = render_template(
                "application_gaoxin_system_summary_print.html",
                company=company,
                company_english_name=company_english_name,
                auto_data=system_summary_data,
            )
            add_standalone_document(
                "11",
                _ensure_landscape_page_rule(system_summary_html),
                "制度汇总表",
            )

            system_base = system_docs.get("base") or {}
            system_company_name = system_base.get("company_name") or company.name
            for doc_type in GAOXIN_SYSTEM_DOC_TYPES:
                doc_key = doc_type["key"]
                content = _append_system_doc_signature(
                    _normalize_system_generated_text(
                        system_docs.get("docs", {}).get(doc_key) or "",
                        products,
                    ),
                    system_company_name,
                    system_base.get("effective_date", ""),
                )
                evidence = _normalize_system_doc_text(
                    _normalize_system_generated_text(
                        system_docs.get("evidence", {}).get(doc_key) or "",
                        products,
                    )
                )
                if not content and not evidence:
                    continue
                if content:
                    document_body, signature_unit, signature_date = _split_system_doc_signature(content)
                    document_html = render_template(
                        "application_gaoxin_system_doc_print.html",
                        company=company,
                        company_english_name=company_english_name,
                        title=doc_type["title"],
                        subtitle="制度正文",
                        content_text=document_body,
                        signature_unit=signature_unit,
                        signature_date=signature_date,
                    )
                    add_portrait_document("11", document_html, doc_type["title"])
                evidence_files = _evidence_file_templates(doc_key, evidence)
                if evidence_files:
                    attachment_notice_html = render_template(
                        "application_gaoxin_system_attachment_notice_print.html",
                        company=company,
                        company_english_name=company_english_name,
                        doc_type=doc_type,
                    )
                    add_portrait_document(
                        "11",
                        attachment_notice_html,
                        f"{doc_type['title']} - 佐证附件说明",
                    )
                for file_index, (file_title, purpose) in enumerate(evidence_files, start=1):
                    evidence_html = render_template(
                        "application_gaoxin_system_evidence_print.html",
                        company=company,
                        company_english_name=company_english_name,
                        doc_type=doc_type,
                        file_title=file_title,
                        sequence_label=f"{GAOXIN_SYSTEM_DOC_TYPES.index(doc_type) + 1}.{file_index}",
                        evidence_context=_normalize_generated_ps_structure(
                            _system_evidence_pdf_context(
                                system_base,
                                doc_type,
                                file_title,
                                purpose,
                                evidence,
                            ),
                            products,
                        ),
                    )
                    add_portrait_document("11", evidence_html, f"{doc_type['title']} - {file_title}")

            for year in GAOXIN_FINANCE_YEARS:
                add_attachment_files("12", [
                    item for item in attachments.get("annual_audit", {}).get("files", [])
                    if str(item.get("attachment_year") or "") == year
                ], f"{year}年年度审计报告")
                add_attachment_files("13", [
                    item for item in attachments.get("tax_settlement", {}).get("files", [])
                    if str(item.get("attachment_year") or "") == year
                ], f"{year}年汇算清缴")

            if job:
                _update_export_job(
                    job,
                    stage="正在准备已上传的附件文件",
                    progress=28,
                )
            prepared_attachment_count = _prepare_export_attachment_files(
                app,
                attachment_references,
            )
            chrome_available = bool(_chrome_executable(app))
            portrait_batch_size = (
                len(portrait_documents)
                if chrome_available and portrait_documents
                else max(
                    1,
                    int(app.config.get("PDF_PYMUPDF_EXPORT_BATCH_SIZE", 4)),
                )
            )
            portrait_batches = _portrait_export_document_batches(
                portrait_documents,
                portrait_batch_size,
            )
            landscape_batch_size = (
                len(standalone_documents)
                if chrome_available and standalone_documents
                else max(
                    1,
                    int(app.config.get("PDF_PYMUPDF_EXPORT_BATCH_SIZE", 4)),
                )
            )
            landscape_batches = _portrait_export_document_batches(
                standalone_documents,
                landscape_batch_size,
            )
            render_task_count = 1 + len(portrait_batches) + len(landscape_batches)
            configured_workers = max(1, int(app.config.get("PDF_RENDER_WORKERS", 3)))
            render_workers = min(configured_workers, render_task_count) if chrome_available else 1
            if job:
                _update_export_job(
                    job,
                    stage=f"正在渲染 PDF，共 {render_task_count} 个批次",
                    progress=40,
                )
            with ThreadPoolExecutor(max_workers=render_workers, thread_name_prefix="gaoxin-pdf") as executor:
                attachments_future = executor.submit(
                    _render_export_pdf_file, app, html, export_dir, "附件目录"
                )
                portrait_futures = [
                    executor.submit(
                        _render_export_document_batch,
                        app,
                        batch,
                        export_dir,
                        company.name,
                        company_english_name,
                        orientation="portrait",
                        label=f"附件内部材料（第{batch_index}/{len(portrait_batches)}批）",
                    )
                    for batch_index, batch in enumerate(portrait_batches, start=1)
                ]
                landscape_futures = [
                    executor.submit(
                        _render_export_document_batch,
                        app,
                        batch,
                        export_dir,
                        company.name,
                        company_english_name,
                        orientation="landscape",
                        label=f"附件横向材料（第{batch_index}/{len(landscape_batches)}批）",
                    )
                    for batch_index, batch in enumerate(landscape_batches, start=1)
                ]

                attachments_pdf_path = attachments_future.result()
                portrait_pdf_paths = [future.result() for future in portrait_futures]
                landscape_pdf_paths = [future.result() for future in landscape_futures]

            if job:
                _update_export_job(
                    job,
                    stage="PDF 渲染完成，正在检查页面",
                    progress=72,
                )
            if not attachments_pdf_path:
                raise RuntimeError("附件目录 PDF 生成失败")
            if portrait_documents:
                if len(portrait_pdf_paths) != len(portrait_batches):
                    raise RuntimeError("附件内部材料 PDF 生成失败")
            if standalone_documents:
                if len(landscape_pdf_paths) != len(landscape_batches):
                    raise RuntimeError("附件横向材料 PDF 生成失败")
            for document in portrait_documents + standalone_documents:
                if not document.get("pdf_path"):
                    raise RuntimeError(f"{document['label']} PDF 生成失败")

            try:
                import fitz
            except ImportError:
                import pymupdf as fitz

            commitment_path = Path(app.static_folder) / "attachments" / "1.企业科研诚信承诺书.pdf"
            if not commitment_path.is_file():
                raise FileNotFoundError("未找到企业科研诚信承诺书原版 PDF")

            merged_pdf_path = os.path.join(export_dir, "gaoxin-attachments-merged.pdf")
            if job:
                _update_export_job(
                    job,
                    stage="正在按 13 个板块顺序合并",
                    progress=82,
                )
            commitment_pdf = fitz.open(str(commitment_path))
            attachments_pdf = fitz.open(attachments_pdf_path)
            merged_pdf = fitz.open()
            generated_page_indexes = []
            try:
                merged_pdf.insert_pdf(commitment_pdf)

                def insert_pdf_file(source, label):
                    path = source.get("pdf_path") if isinstance(source, dict) else source
                    if not path or not os.path.isfile(path):
                        app.logger.warning("%s文件不存在，已跳过：%s", label, path)
                        return
                    try:
                        opened_document = fitz.open(path)
                        source_pdf = opened_document
                        converted_image_pdf = None
                        try:
                            if not opened_document.is_pdf:
                                converted_image_pdf = fitz.open()
                                for image_page in opened_document:
                                    image_rect = image_page.rect
                                    landscape = image_rect.width > image_rect.height
                                    page_width, page_height = (
                                        (841.89, 595.28) if landscape else (595.28, 841.89)
                                    )
                                    margin = 24
                                    available_width = page_width - margin * 2
                                    available_height = page_height - margin * 2
                                    scale = min(
                                        available_width / image_rect.width,
                                        available_height / image_rect.height,
                                    )
                                    target_width = image_rect.width * scale
                                    target_height = image_rect.height * scale
                                    left = (page_width - target_width) / 2
                                    top = (page_height - target_height) / 2
                                    target_rect = fitz.Rect(
                                        left,
                                        top,
                                        left + target_width,
                                        top + target_height,
                                    )
                                    pdf_page = converted_image_pdf.new_page(
                                        width=page_width,
                                        height=page_height,
                                    )
                                    pixmap = image_page.get_pixmap(alpha=False)
                                    pdf_page.insert_image(
                                        target_rect,
                                        stream=pixmap.tobytes("png"),
                                        keep_proportion=True,
                                    )
                                source_pdf = converted_image_pdf
                            if source_pdf.page_count:
                                from_page = int(source.get("from_page", 0)) if isinstance(source, dict) else 0
                                to_page = int(source.get("to_page", source_pdf.page_count - 1)) if isinstance(source, dict) else source_pdf.page_count - 1
                                insert_start = merged_pdf.page_count
                                merged_pdf.insert_pdf(source_pdf, from_page=from_page, to_page=to_page)
                                if (
                                    isinstance(source, dict)
                                    and not source.get("uploaded_attachment")
                                ):
                                    generated_page_indexes.extend(
                                        _generated_insert_header_page_indexes(
                                            insert_start,
                                            merged_pdf.page_count,
                                            skip_first_page=bool(
                                                source.get("skip_header_first_page")
                                            ),
                                        )
                                    )
                        finally:
                            if converted_image_pdf is not None:
                                converted_image_pdf.close()
                            opened_document.close()
                    except Exception:
                        app.logger.exception("%s文件无法合并，已跳过：%s", label, path)

                section_start_pages = {}
                for page_index, page in enumerate(attachments_pdf):
                    page_text = page.get_text() or ""
                    for section in export_sections:
                        section_no = str(section["no"])
                        marker = f"GAOXINSECTION{section_no}"
                        if marker in page_text and section_no not in section_start_pages:
                            section_start_pages[section_no] = page_index

                ordered_numbers = [str(section["no"]) for section in export_sections]
                section_ranges = _ordered_attachment_section_ranges(
                    section_start_pages,
                    attachments_pdf.page_count,
                    ordered_numbers,
                )
                for section_no, start_page, end_page in section_ranges:
                    insert_start = merged_pdf.page_count
                    merged_pdf.insert_pdf(
                        attachments_pdf,
                        from_page=start_page,
                        to_page=end_page,
                    )
                    generated_page_indexes.extend(
                        range(insert_start, merged_pdf.page_count)
                    )
                    for source, label in section_paths.get(section_no, []):
                        insert_pdf_file(source, label)
                _stamp_generated_pdf_pages(
                    merged_pdf,
                    generated_page_indexes,
                    company.name,
                    company_english_name,
                )
                merged_pdf.save(merged_pdf_path)
            finally:
                merged_pdf.close()
                attachments_pdf.close()
                commitment_pdf.close()

            pdf_bytes = Path(merged_pdf_path).read_bytes()
            duration = time.perf_counter() - export_started
            app.logger.info(
                "高新附件 PDF 导出完成 company_id=%s duration=%.2fs renderer=%s pdf_renders=%s generated_documents=%s bytes=%s",
                company.id,
                duration,
                "chrome" if chrome_available else "pymupdf",
                render_task_count,
                len(portrait_documents) + len(standalone_documents),
                len(pdf_bytes),
            )
            app.logger.info(
                "高新附件源文件准备完成 company_id=%s resolved=%s requested=%s",
                company.id,
                prepared_attachment_count,
                len(attachment_references),
            )
            if job:
                _update_export_job(
                    job,
                    stage="正在保存导出结果",
                    progress=94,
                )
                relative_path = _gaoxin_attachment_export_relative_path(job)
                job.result_path = relative_path
                output_path = _gaoxin_attachment_export_local_path(job)
                Path(output_path).parent.mkdir(parents=True, exist_ok=True)
                Path(output_path).write_bytes(pdf_bytes)
                blob_result = persist_file(output_path, relative_path)
                if blob_enabled() and not blob_result:
                    raise RuntimeError("PDF 已生成，但云端文件保存失败")
                job.result_size = len(pdf_bytes)
                job.duration_seconds = duration
                job.completed_at = datetime.utcnow()
                _update_export_job(
                    job,
                    status="ready",
                    stage="PDF 已生成，可下载",
                    progress=100,
                    error="",
                )
                return jsonify(_export_job_payload(job))
            response = send_file(
                BytesIO(pdf_bytes),
                mimetype="application/pdf",
                as_attachment=True,
                download_name=f"高新技术企业认定附件制作_{company.name}.pdf",
            )
            response.headers["X-Export-Duration"] = f"{duration:.2f}"
            response.headers["X-PDF-Render-Count"] = str(render_task_count)
            return response
    except Exception as e:
        app.logger.exception("高新附件 PDF 导出失败 company_id=%s", company.id)
        if job:
            db.session.rollback()
            current_job = db.session.get(ExportJob, job.id)
            if current_job:
                _update_export_job(
                    current_job,
                    status="failed",
                    stage="PDF 生成失败",
                    error=f"PDF生成失败：{str(e)}",
                )
        return jsonify({"ok": False, "error": f"PDF生成失败：{str(e)}"}), 500


def _staff_rows_for_certificate_match(data, raw_rows=""):
    rows = []
    try:
        submitted_rows = json.loads(raw_rows) if raw_rows else []
    except (json.JSONDecodeError, TypeError):
        submitted_rows = []
    if isinstance(submitted_rows, list):
        for fallback_index, row in enumerate(submitted_rows):
            if not isinstance(row, dict):
                continue
            try:
                index = int(row.get("index", fallback_index))
            except (TypeError, ValueError):
                index = fallback_index
            name = str(row.get("name") or "").strip()
            if name and 0 <= index < 100:
                rows.append({"index": index, "name": name})
    if rows:
        return rows

    for index in range(100):
        name = str(data.get(f"attachment_rd_staff_{index}_name") or "").strip()
        if name:
            rows.append({"index": index, "name": name})
    if rows:
        return rows

    hr_rows = data.get("hr_staff_rows")
    if isinstance(hr_rows, list):
        for index, row in enumerate(hr_rows[:100]):
            if not isinstance(row, dict):
                continue
            name = str(row.get("姓名") or "").strip()
            if name:
                rows.append({"index": index, "name": name})
    return rows


def _upload_staff_certificates(company, certificate_type):
    uploads = [item for item in request.files.getlist("files") if item and item.filename]
    if not uploads:
        upload = request.files.get("file")
        if upload and upload.filename:
            uploads = [upload]
    if not uploads:
        return jsonify({"ok": False, "error": "请先选择 PDF 或图片文件"}), 400
    if any(Path(upload.filename).suffix.lower() not in STAFF_CERTIFICATE_UPLOAD_EXTENSIONS for upload in uploads):
        return jsonify({
            "ok": False,
            "error": "毕业证书和职称证书支持 PDF、JPG、JPEG、PNG、WEBP、BMP、TIF、TIFF 文件",
        }), 400

    data = _load_company_data(company)
    attachments = _load_gaoxin_attachments_from_data(data)
    staff_rows = _staff_rows_for_certificate_match(data, request.form.get("staff_rows", ""))
    uploaded_files = []
    matches = []

    for upload in uploads:
        original_filename = upload.filename
        extension = Path(original_filename).suffix.lower()
        safe_name = secure_filename(original_filename)
        if not safe_name or "." not in safe_name:
            safe_name = f"staff_certificate_{uuid.uuid4().hex}{extension}"

        file_id = uuid.uuid4().hex
        stored_filename = f"{file_id}_{safe_name}"
        relative_path = _attachment_relative_path(
            current_user.id,
            company.id,
            "staff_statement",
            stored_filename,
        )
        target_path = _safe_attachment_path(relative_path)
        os.makedirs(os.path.dirname(target_path), exist_ok=True)
        upload.save(target_path)
        persist_file(target_path, relative_path)

        try:
            analysis = analyze_staff_certificate(
                _extract_staff_certificate_text(target_path),
                original_filename,
                certificate_type,
                staff_rows,
            )
        except Exception:
            current_app.logger.exception("人员证书识别失败：%s", original_filename)
            analysis = {
                "certificate_type": certificate_type,
                "field": STAFF_CERTIFICATE_FIELDS[certificate_type],
                "value": "",
                "matched_index": None,
                "matched_name": "",
                "status": "parse_error",
            }

        if analysis.get("status") == "matched":
            index = analysis.get("matched_index")
            field = analysis.get("field")
            value = str(analysis.get("value") or "").strip()
            if isinstance(index, int) and field and value:
                data[f"attachment_rd_staff_{index}_{field}"] = value
                matched_name = str(analysis.get("matched_name") or "").strip()
                if matched_name and not data.get(f"attachment_rd_staff_{index}_name"):
                    data[f"attachment_rd_staff_{index}_name"] = matched_name

        file_meta = {
            "id": file_id,
            "original_filename": original_filename,
            "stored_filename": stored_filename,
            "relative_path": relative_path,
            "uploaded_at": datetime.utcnow().isoformat(timespec="seconds"),
            "staff_file_type": certificate_type,
            "content_type": mimetypes.guess_type(original_filename)[0] or "application/octet-stream",
            "certificate_match": analysis,
        }
        attachments["staff_statement"]["files"].append(file_meta)
        uploaded_files.append({
            **file_meta,
            "view_url": url_for(
                "docgen.gaoxin_attachment_file",
                company_id=company.id,
                section_key="staff_statement",
                file_id=file_id,
            ),
            "delete_url": url_for(
                "docgen.gaoxin_attachment_delete",
                company_id=company.id,
                section_key="staff_statement",
                file_id=file_id,
            ),
        })
        matches.append(analysis)

    data["gaoxin_attachments"] = attachments
    company.data_json = json.dumps(data, ensure_ascii=False)
    db.session.commit()
    return jsonify({
        "ok": True,
        "file": uploaded_files[0],
        "files": uploaded_files,
        "matches": matches,
        "matched_count": sum(1 for item in matches if item.get("status") == "matched"),
    })


@docgen_bp.route("/gaoxin_attachments/<int:company_id>/upload/<section_key>", methods=["POST"])
@login_required
def gaoxin_attachment_upload(company_id, section_key):
    company = Company.query.filter_by(id=company_id, user_id=current_user.id).first_or_404()
    if section_key not in GAOXIN_ATTACHMENT_PDF_KEYS:
        return jsonify({"ok": False, "error": "该附件板块不支持 PDF 上传"}), 400

    staff_file_type = str(request.form.get("staff_file_type") or "").strip()
    if section_key == "staff_statement" and staff_file_type in STAFF_CERTIFICATE_FIELDS:
        return _upload_staff_certificates(company, staff_file_type)

    upload = request.files.get("file")
    if not upload or not upload.filename:
        return jsonify({"ok": False, "error": "请先选择 PDF 文件"}), 400

    original_filename = upload.filename
    if not original_filename.lower().endswith(".pdf"):
        return jsonify({"ok": False, "error": "仅支持上传 PDF 文件"}), 400
    safe_name = secure_filename(original_filename)
    if not safe_name or "." not in safe_name:
        safe_name = f"attachment_{uuid.uuid4().hex}.pdf"
    if not safe_name.lower().endswith(".pdf"):
        safe_name = f"{os.path.splitext(safe_name)[0]}.pdf"

    file_id = uuid.uuid4().hex
    stored_filename = f"{file_id}_{safe_name}"
    relative_path = _attachment_relative_path(current_user.id, company.id, section_key, stored_filename)
    target_path = _safe_attachment_path(relative_path)
    os.makedirs(os.path.dirname(target_path), exist_ok=True)
    upload.save(target_path)
    persist_file(target_path, relative_path)

    data = _load_company_data(company)
    attachments = _load_gaoxin_attachments_from_data(data)
    file_meta = {
        "id": file_id,
        "original_filename": original_filename,
        "stored_filename": stored_filename,
        "relative_path": relative_path,
        "uploaded_at": datetime.utcnow().isoformat(timespec="seconds"),
    }
    if section_key == "ip":
        file_meta["ip_seq"] = str(request.form.get("ip_seq") or "").strip()
    if section_key == "staff_statement":
        file_meta["staff_file_type"] = str(request.form.get("staff_file_type") or "").strip()
    if section_key in {"rd_expense_audit", "annual_audit", "tax_settlement"}:
        file_meta["attachment_year"] = str(request.form.get("attachment_year") or "").strip()
    if section_key == "achievement_transform":
        achievement_index = str(request.form.get("achievement_index") or "").strip()
        evidence_type = str(request.form.get("achievement_evidence_type") or "").strip()
        if not achievement_index or evidence_type not in _achievement_evidence_upload_map():
            return jsonify({"ok": False, "error": "成果附件参数不完整"}), 400
        file_meta["achievement_index"] = achievement_index
        file_meta["achievement_evidence_type"] = evidence_type
    if section_key == "hitech_products":
        product_index = str(request.form.get("product_index") or "").strip()
        evidence_type = str(request.form.get("product_evidence_type") or "").strip()
        if not product_index or evidence_type not in _hitech_product_evidence_map():
            return jsonify({"ok": False, "error": "产品附件参数不完整"}), 400
        file_meta["product_index"] = product_index
        file_meta["product_evidence_type"] = evidence_type
    attachments[section_key]["files"].append(file_meta)
    data["gaoxin_attachments"] = attachments
    company.data_json = json.dumps(data, ensure_ascii=False)
    db.session.commit()

    return jsonify({
        "ok": True,
        "file": {
            **file_meta,
            "view_url": url_for("docgen.gaoxin_attachment_file", company_id=company.id, section_key=section_key, file_id=file_id),
            "delete_url": url_for("docgen.gaoxin_attachment_delete", company_id=company.id, section_key=section_key, file_id=file_id),
        },
    })


@docgen_bp.route("/gaoxin_attachments/<int:company_id>/file/<section_key>/<file_id>")
@login_required
def gaoxin_attachment_file(company_id, section_key, file_id):
    company = Company.query.filter_by(id=company_id, user_id=current_user.id).first_or_404()
    files = _load_gaoxin_attachments(company).get(section_key, {}).get("files", [])
    file_meta = next((item for item in files if item.get("id") == file_id), None)
    if not file_meta:
        abort(404)
    path = _safe_attachment_path(file_meta.get("relative_path", ""))
    if not os.path.exists(path):
        abort(404)
    return send_file(
        path,
        mimetype=(
            file_meta.get("content_type")
            or mimetypes.guess_type(file_meta.get("original_filename") or path)[0]
            or "application/octet-stream"
        ),
        as_attachment=request.args.get("download") == "1",
        download_name=file_meta.get("original_filename") or "附件",
    )


@docgen_bp.route("/gaoxin_attachments/<int:company_id>/file/<section_key>/<file_id>/delete", methods=["POST"])
@login_required
def gaoxin_attachment_delete(company_id, section_key, file_id):
    company = Company.query.filter_by(id=company_id, user_id=current_user.id).first_or_404()
    data = _load_company_data(company)
    attachments = _load_gaoxin_attachments_from_data(data)
    stored_section_key = section_key
    saved_attachments = data.get("gaoxin_attachments") if isinstance(data.get("gaoxin_attachments"), dict) else {}
    if section_key == "hitech_income_audit" and not isinstance(saved_attachments.get(section_key), dict):
        stored_section_key = "special_audit"
    section = attachments.get(stored_section_key)
    if not section:
        return jsonify({"ok": False, "error": "附件板块不存在"}), 404

    kept = []
    deleted = None
    for item in section.get("files", []):
        if item.get("id") == file_id:
            deleted = item
        else:
            kept.append(item)
    if not deleted:
        return jsonify({"ok": False, "error": "文件不存在"}), 404

    section["files"] = kept
    if stored_section_key != section_key:
        legacy_section = saved_attachments.get(stored_section_key) if isinstance(saved_attachments.get(stored_section_key), dict) else {"files": []}
        legacy_section["files"] = kept
        saved_attachments[stored_section_key] = legacy_section
        data["gaoxin_attachments"] = saved_attachments
    else:
        data["gaoxin_attachments"] = attachments
    company.data_json = json.dumps(data, ensure_ascii=False)
    db.session.commit()

    try:
        path = _safe_attachment_path(deleted.get("relative_path", ""))
        if os.path.exists(path):
            os.remove(path)
        delete_file(deleted.get("relative_path", ""))
    except Exception:
        pass

    return jsonify({"ok": True})


@docgen_bp.route("/gaoxin_system_docs/<int:company_id>", methods=["GET", "POST"])
@login_required
def gaoxin_system_docs(company_id):
    """高新组织管理制度文件生成页。"""
    company = Company.query.filter_by(id=company_id, user_id=current_user.id).first_or_404()
    if request.method == "POST":
        _save_gaoxin_system_docs(company, request.form)
        flash("制度文件已保存", "success")
        return redirect(url_for("docgen.gaoxin_system_docs", company_id=company.id))

    system_docs = _load_gaoxin_system_docs(company)
    completed_count = sum(1 for content in system_docs["docs"].values() if content.strip())
    return render_template(
        "application_gaoxin_system_docs.html",
        company=company,
        doc_types=GAOXIN_SYSTEM_DOC_TYPES,
        base_fields=GAOXIN_SYSTEM_BASE_FIELDS,
        system_docs=system_docs,
        completed_count=completed_count,
    )


@docgen_bp.route("/gaoxin_system_docs/<int:company_id>/ai_generate", methods=["POST"])
@login_required
def gaoxin_system_docs_ai_generate(company_id):
    """AI 撰写单份高新组织管理制度正文。"""
    company = Company.query.filter_by(id=company_id, user_id=current_user.id).first_or_404()
    payload = request.get_json(silent=True) or {}
    doc_key = str(payload.get("doc_key") or "").strip()
    doc_type = _system_doc_type_map().get(doc_key)
    if not doc_type:
        return jsonify({"success": False, "error": "制度类型不存在"}), 400

    base = payload.get("base") if isinstance(payload.get("base"), dict) else {}
    existing = str(payload.get("existing") or "").strip()
    try:
        target_words = int(payload.get("target_words") or 1000)
    except (TypeError, ValueError):
        target_words = 1000
    target_words = max(600, min(1600, target_words))

    context_text = _format_system_doc_context(base)
    time_context = _system_doc_time_context(base)
    existing_instruction = ""
    if existing:
        existing_instruction = f"""

当前已有草稿：
{existing}

请在保留可用企业信息和制度逻辑的基础上重写完善，不要简单续写。"""

    prompt = f"""请为高新技术企业认定申报附件撰写《{doc_type['title']}》。

企业基础信息：
{context_text}

制度撰写重点：{doc_type['guide']}{existing_instruction}
{time_context['prompt']}

要求：
1. 正文约 {target_words} 字，允许上下浮动 10%。
2. 必须自然穿插企业名称、负责人、研发负责人、财务负责人、研发机构、主营业务、技术领域、研发项目等已提供信息；未提供的信息不得编造具体姓名或具体数据，可写为“由公司指定负责人”“按年度计划执行”等稳妥表述。
3. 结构要像正式企业制度文件，建议包含：总则、职责分工、流程要求、记录归档、监督检查、附则等部分；不同制度可根据主题调整。
4. 语言正式、可直接作为申报附件初稿使用，避免空洞口号。
5. 不要使用 Markdown 格式，不要输出 **、#、``` 等标记字符。
6. 段落之间最多保留一个空行，不要连续输出多个空行。
7. 正文末尾必须右下角落款，格式为“单位：{base.get('company_name') or company.name}”和“日期：{base.get('effective_date') or '    年    月    日'}”。
8. 只输出制度正文，不要输出解释或额外前缀。"""

    result = call_llm(
        [
            {"role": "system", "content": "你是高新技术企业认定申报材料和企业研发管理制度撰写专家，只输出可直接使用的制度正文。"},
            {"role": "user", "content": prompt},
        ],
        temperature=0.65,
        max_tokens=max(1600, min(5200, int(target_words * 2.4))),
        timeout=90,
        max_attempts=1,
    )
    if not result.get("success"):
        return jsonify({"success": False, "error": result.get("error", "AI 调用失败")}), 500

    data = _load_company_data(company)
    products = _collect_hitech_product_rows(_sync_gaoxin_finance_years(_merge_relation_fields(data)))
    text = _append_system_doc_signature(
        _normalize_system_generated_text(result.get("content") or "", products),
        base.get("company_name") or company.name,
        base.get("effective_date", ""),
    )
    return jsonify({"success": True, "text": text, "actual_words": len("".join(text.split()))})


@docgen_bp.route("/gaoxin_system_docs/<int:company_id>/ai_generate_evidence", methods=["POST"])
@login_required
def gaoxin_system_docs_ai_generate_evidence(company_id):
    """AI 生成单份制度对应的佐证材料清单。"""
    company = Company.query.filter_by(id=company_id, user_id=current_user.id).first_or_404()
    payload = request.get_json(silent=True) or {}
    doc_key = str(payload.get("doc_key") or "").strip()
    doc_type = _system_doc_type_map().get(doc_key)
    if not doc_type:
        return jsonify({"success": False, "error": "制度类型不存在"}), 400

    base = payload.get("base") if isinstance(payload.get("base"), dict) else {}
    doc_content = _normalize_system_doc_text(payload.get("doc_content") or "")
    existing = _normalize_system_doc_text(payload.get("existing") or "")
    context_text = _format_system_doc_context(base)
    time_context = _system_doc_time_context(base)
    existing_instruction = ""
    if existing:
        existing_instruction = f"""

当前已有佐证材料草稿：
{existing}

请在保留可用内容的基础上重新整理为更完整、更适合归档的版本。"""

    prompt = f"""请为《{doc_type['title']}》生成高新技术企业认定申报用的制度文件佐证材料内容要求，用于后续在线生成个性化 Word 佐证文件。

企业基础信息：
{context_text}

制度撰写重点：{doc_type['guide']}

制度正文摘要：
{doc_content[:1800] if doc_content else '暂未提供制度正文，请按制度名称和企业基础信息生成佐证文件内容。'}{existing_instruction}
{time_context['prompt']}

要求：
1. 佐证材料以“可直接生成的表单、台账、审批表、会议或培训记录”为主，优先使用表格字段，文字说明只保留一至两句短句，不写大段背景阐述。
2. 优先使用企业基础信息中已经填写的项目名称、编号、来源、周期、预算、研发目标、核心内容、创新点、验收指标、预期成果、人员、日期和归档位置；缺失项留空或写“待补充”，不要让用户在生成后的 Word 中再大量修改。
3. 每类佐证文件只列核心字段和默认填充值，审批/复核链条写清即可，不要扩展成可行性报告、论证报告或长篇说明。
4. 如果是研发项目立项管理制度，只围绕“研发项目立项申请表”输出核心申请表字段，不再拆成可行性论证、预算审批、验收清单等多个文件。
5. 建议按以下结构输出：一、佐证文件名称；二、表格核心字段及默认值；三、审批签字与归档信息；四、少量留空待补项。总字数尽量控制在500字以内。
6. 不要使用 Markdown，不要输出 **、#、``` 等标记字符；段落之间最多一个空行。
7. 只输出佐证文件内容要求正文，不要输出解释性前缀。"""

    result = call_llm(
        [
            {"role": "system", "content": "你是高新技术企业认定申报附件材料整理专家，只输出可直接用于生成个性化 Word 佐证文件的内容要求。"},
            {"role": "user", "content": prompt},
        ],
        temperature=0.55,
        max_tokens=2600,
        timeout=70,
        max_attempts=1,
    )
    if not result.get("success"):
        return jsonify({"success": False, "error": result.get("error", "AI 调用失败")}), 500

    data = _load_company_data(company)
    products = _collect_hitech_product_rows(_sync_gaoxin_finance_years(_merge_relation_fields(data)))
    text = _normalize_system_generated_text(enforce_temporal_wording(
        _normalize_system_doc_text(result.get("content") or ""),
        time_context["project"],
    ), products)
    return jsonify({"success": True, "text": text, "actual_words": len("".join(text.split()))})


@docgen_bp.route("/gaoxin_system_docs/<int:company_id>/word", methods=["GET", "POST"])
@login_required
def gaoxin_system_docs_word(company_id):
    """导出高新组织管理制度文件汇编 Word。"""
    company = Company.query.filter_by(id=company_id, user_id=current_user.id).first_or_404()
    if request.method == "POST":
        system_docs = _save_gaoxin_system_docs(company, request.form)
    else:
        system_docs = _load_gaoxin_system_docs(company)

    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    import tempfile

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5

    def set_run_font(run, size=12, bold=False, name="宋体"):
        run.font.size = Pt(size)
        run.bold = bold
        run.font.name = name
        run._element.rPr.rFonts.set(qn("w:eastAsia"), name)

    products = _collect_hitech_product_rows(
        _sync_gaoxin_finance_years(_merge_relation_fields(_load_company_data(company)))
    )
    base = _normalize_generated_ps_structure(system_docs.get("base") or {}, products)
    company_name = base.get("company_name") or company.name
    company_english_name = base.get("company_english_name") or _company_english_name(company, _load_company_data(company))
    _add_docx_company_header(doc, company_name, company_english_name, set_run_font)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(title.add_run(f"{company.name}高新技术企业组织管理制度文件汇编"), 18, True)

    info_lines = []
    for key, label in GAOXIN_SYSTEM_BASE_FIELDS:
        value = str(base.get(key) or "").strip()
        if value:
            info_lines.append(f"{label}：{value}")
    if info_lines:
        p = doc.add_paragraph()
        set_run_font(p.add_run("基础信息"), 13, True)
        for line in info_lines:
            para = doc.add_paragraph()
            set_run_font(para.add_run(line), 10)

    docs = system_docs.get("docs") or {}
    evidence_docs = system_docs.get("evidence") or {}
    effective_date = base.get("effective_date", "")
    for index, doc_type in enumerate(GAOXIN_SYSTEM_DOC_TYPES, start=1):
        content = _append_system_doc_signature(
            _normalize_system_generated_text(docs.get(doc_type["key"]) or "", products),
            company_name,
            effective_date,
        )
        evidence = _normalize_system_doc_text(
            _normalize_system_generated_text(evidence_docs.get(doc_type["key"]) or "", products)
        )
        if not content and not evidence:
            continue
        evidence_files = _evidence_file_templates(doc_type["key"], evidence)
        doc.add_page_break()
        heading = doc.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(heading.add_run(f"{index}. {doc_type['title']}"), 16, True)
        if content:
            for block in re.split(r"\n{2,}", content):
                block = block.strip()
                if not block:
                    continue
                para = doc.add_paragraph()
                if block.startswith("单位：") or block.startswith("落款单位：") or block.startswith("日期："):
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                set_run_font(para.add_run(block), 12, False)
        if evidence_files:
            doc.add_page_break()
            attachment_heading = doc.add_paragraph()
            attachment_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_run_font(
                attachment_heading.add_run(f"以下材料为《{doc_type['title']}》的佐证附件"),
                12,
                True,
            )
            doc.add_page_break()
            for file_index, (file_title, purpose) in enumerate(evidence_files, start=1):
                if file_index > 1:
                    doc.add_page_break()
                heading = doc.add_paragraph()
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_run_font(heading.add_run(f"{index}.{file_index} {file_title}"), 15, True)
                note = doc.add_paragraph()
                note.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_run_font(note.add_run(f"对应制度：{doc_type['title']}"), 10)
                _add_evidence_form_table(doc, set_run_font, base, doc_type, file_title, purpose, evidence)

    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    doc.save(tmp.name)
    tmp.close()
    return send_file(
        tmp.name,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=f"高新组织管理制度文件汇编_{company.name}.docx",
    )


@docgen_bp.route("/gaoxin_system_docs/<int:company_id>/evidence_zip", methods=["GET", "POST"])
@login_required
def gaoxin_system_docs_evidence_zip(company_id):
    """导出高新制度佐证材料文件包。"""
    company = Company.query.filter_by(id=company_id, user_id=current_user.id).first_or_404()
    system_docs = _save_gaoxin_system_docs(company, request.form) if request.method == "POST" else _load_gaoxin_system_docs(company)

    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    import tempfile
    import zipfile

    products = _collect_hitech_product_rows(
        _sync_gaoxin_finance_years(_merge_relation_fields(_load_company_data(company)))
    )
    base = _normalize_generated_ps_structure(system_docs.get("base") or {}, products)
    evidence_docs = system_docs.get("evidence") or {}
    company_name = base.get("company_name") or company.name
    company_english_name = base.get("company_english_name") or _company_english_name(company, _load_company_data(company))

    def set_run_font(run, size=12, bold=False, name="宋体"):
        run.font.size = Pt(size)
        run.bold = bold
        run.font.name = name
        run._element.rPr.rFonts.set(qn("w:eastAsia"), name)

    zip_tmp = tempfile.NamedTemporaryFile(suffix=".zip", delete=False)
    zip_tmp.close()
    with zipfile.ZipFile(zip_tmp.name, "w", zipfile.ZIP_DEFLATED) as zf:
        generated_count = 0
        for doc_index, doc_type in enumerate(GAOXIN_SYSTEM_DOC_TYPES, start=1):
            evidence = _normalize_system_doc_text(
                _normalize_system_generated_text(
                    evidence_docs.get(doc_type["key"]) or "",
                    products,
                )
            )
            for file_index, (file_title, purpose) in enumerate(_evidence_file_templates(doc_type["key"], evidence), start=1):
                evidence_doc = Document()
                section = evidence_doc.sections[0]
                section.top_margin = Cm(2.0)
                section.bottom_margin = Cm(2.0)
                section.left_margin = Cm(2.0)
                section.right_margin = Cm(2.0)
                style = evidence_doc.styles["Normal"]
                style.font.name = "Times New Roman"
                style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                style.font.size = Pt(11)
                _add_docx_company_header(evidence_doc, company_name, company_english_name, set_run_font)
                heading = evidence_doc.add_paragraph()
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_run_font(heading.add_run(f"{company_name}{file_title}"), 16, True)
                note = evidence_doc.add_paragraph()
                note.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_run_font(note.add_run(f"对应制度：{doc_type['title']}"), 10)
                _add_evidence_form_table(evidence_doc, set_run_font, base, doc_type, file_title, purpose, evidence)
                doc_tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
                doc_tmp.close()
                evidence_doc.save(doc_tmp.name)
                arcname = f"{doc_index:02d}_{_safe_docx_name(doc_type['short'])}/{file_index:02d}_{_safe_docx_name(file_title)}.docx"
                zf.write(doc_tmp.name, arcname)
                generated_count += 1
        if generated_count == 0:
            readme_tmp = tempfile.NamedTemporaryFile(suffix=".txt", mode="w", encoding="utf-8", delete=False)
            readme_tmp.write("未生成佐证材料文件。请先维护制度文件佐证材料后重新导出。")
            readme_tmp.close()
            zf.write(readme_tmp.name, "README.txt")

    return send_file(
        zip_tmp.name,
        mimetype="application/zip",
        as_attachment=True,
        download_name=f"高新制度佐证材料文件包_{company.name}.zip",
    )


@docgen_bp.route("/gaoxin_book/<int:company_id>/pdf")
@login_required
def gaoxin_book_pdf(company_id):
    """导出高新技术企业认定申请书 PDF"""
    company = Company.query.filter_by(id=company_id, user_id=current_user.id).first_or_404()
    
    data = {}
    if company.data_json:
        try:
            data = json.loads(company.data_json)
        except (json.JSONDecodeError, TypeError):
            pass
    if "last_finance_data" in session:
        data.update(session["last_finance_data"])
    data = _sync_gaoxin_finance_years(_merge_relation_fields(data))

    html = render_template(
        "application_gaoxin_print.html",
        company=company,
        company_english_name=_company_english_name(company, data),
        data=data,
    )
    
    try:
        with tempfile.TemporaryDirectory(prefix="gaoxin-book-pdf-") as output_dir:
            pdf_path = os.path.join(output_dir, "gaoxin-book.pdf")
            _render_pdf_file(
                current_app._get_current_object(),
                html,
                pdf_path,
                "高新技术企业认定申请书",
            )
            _stamp_pdf_file_headers(
                pdf_path,
                company.name,
                _company_english_name(company, data),
            )
            pdf_bytes = Path(pdf_path).read_bytes()
        return send_file(
            BytesIO(pdf_bytes),
            mimetype="application/pdf",
            as_attachment=True,
            download_name=f"高新技术企业认定申请书_{company.name}.pdf",
        )
    except Exception as e:
        current_app.logger.exception("高新技术企业认定申请书 PDF 生成失败")
        return f"PDF生成失败: {str(e)}", 500


@docgen_bp.route("/gaoxin_book/<int:company_id>/word", methods=["GET", "POST"])
@login_required
def gaoxin_book_word(company_id):
    """导出高新技术企业认定申请书 Word"""
    company = Company.query.filter_by(id=company_id, user_id=current_user.id).first_or_404()
    
    # 如果是 POST 请求，先保存表单数据
    if request.method == "POST":
        form_data = request.form.to_dict()
        _save_gaoxin_book_data(company, form_data, session.get("last_finance_data"))
    
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    import tempfile

    company = Company.query.filter_by(id=company_id, user_id=current_user.id).first_or_404()
    data = {}
    if company.data_json:
        try:
            data = json.loads(company.data_json)
        except:
            pass
    # 也合并 session 中的财报数据
    if "last_finance_data" in session:
        data.update(session["last_finance_data"])
    data = _sync_gaoxin_finance_years(_merge_relation_fields(data))

    doc = Document()

    # 全局默认字体
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5

    def set_run_font(run, size=12, bold=False, name='宋体'):
        run.font.size = Pt(size)
        run.bold = bold
        run.font.name = name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), name)

    _add_docx_company_header(
        doc,
        company.name,
        _company_english_name(company, data),
        set_run_font,
    )

    def add_title(text, size=22):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        set_run_font(p.add_run(text), size, True)
        return p

    def add_heading(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        set_run_font(p.add_run(text), 15, True)
        return p

    def add_body(text):
        if not text: return
        text = str(text).replace('\r\n', '\n').replace('\r', '\n').strip()
        text = re.sub(r'\n[ \t　]*\n+', '\n', text)
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.7)
        set_run_font(p.add_run(text), 11)

    def add_table(headers, rows, col_widths=None):
        tbl_obj = doc.add_table(rows=1+len(rows), cols=len(headers), style='Table Grid')
        tbl_obj.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(headers):
            c = tbl_obj.rows[0].cells[i]
            c.text = ''
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_run_font(p.add_run(h), 10, True)
        for ri, row in enumerate(rows):
            for ci, v in enumerate(row):
                c = tbl_obj.rows[ri+1].cells[ci]
                c.text = ''
                set_run_font(c.paragraphs[0].add_run(str(v or '')), 10)
        if col_widths:
            for ri in range(len(rows)+1):
                for ci, w in enumerate(col_widths):
                    if ci < len(headers):
                        tbl_obj.rows[ri].cells[ci].width = Cm(w)
        return tbl_obj

    # ====== 企业简介 ======
    intro = data.get('company_intro','')
    scope = data.get('business_scope','')
    if intro or scope:
        add_heading('企业简介')
        if intro:
            add_body(intro)
        if scope:
            add_body(f'经营范围：{scope}')

    # ====== 一、主要情况 ======
    add_heading('一、主要情况')
    add_table(
        ['项目','内容','项目','内容'],
        [
            ['企业名称', company.name, '技术领域', data.get('tech_field','')],
            ['职工总数', data.get('staff_total',''), '科技人员', data.get('tech_staff','')],
            ['Ⅰ类IP', data.get('ip_class1_count','0'), 'Ⅱ类IP', data.get('ip_class2_count','0')],
        ]
    )
    # 三年经营数据
    add_table(
        ['年度','净资产(万元)','销售收入(万元)','利润总额(万元)'],
        [
            [data.get('year1_label','2023'), data.get('year1_net_assets',''), data.get('year1_sales',''), data.get('year1_profit','')],
            [data.get('year2_label','2024'), data.get('year2_net_assets',''), data.get('year2_sales',''), data.get('year2_profit','')],
            [data.get('year3_label','2025'), data.get('year3_net_assets',''), data.get('year3_sales',''), data.get('year3_profit','')],
        ]
    )
    add_table(
        ['项目','金额(万元)'],
        [
            ['近三年研发费用总额', data.get('rd_total_3y','')],
            ['其中：境内研发费用', data.get('rd_domestic','')],
            ['近一年总收入', data.get('revenue_1y','')],
            ['近一年高新技术产品收入', data.get('hitech_revenue','')],
        ]
    )

    # ====== 二、知识产权 ======
    add_heading('二、知识产权汇总表')
    ips = []
    ip_count = _indexed_form_count(data, "ip_", "name", 10)
    for i in range(ip_count):
        nm = data.get(f'ip_{i}_name','')
        if nm:
            patent_no = data.get(f'ip_{i}_patent_no','') or data.get(f'ip_{i}_no','') or data.get(f'ip_{i}_auth_no','')
            ips.append([data.get(f'ip_{i}_seq','') or str(i + 1), nm, data.get(f'ip_{i}_type',''),
                       data.get(f'ip_{i}_status',''), patent_no,
                       data.get(f'ip_{i}_app_date',''), data.get(f'ip_{i}_applicant',''),
                       data.get(f'ip_{i}_date','')])
    if ips:
        add_table(['编号','名称','专利类型','法律状态','专利号','申请日期','专利权人','授权日期'], ips)
    else:
        add_body('（未填写）')

    # ====== 三、人力资源 ======
    add_heading('三、人力资源情况表')
    add_table(
        ['类别','人数','类别','人数'],
        [
            ['职工总数', data.get('hr_total',''), '科技人员', data.get('hr_tech','')],
            ['在职', data.get('hr_fulltime',''), '兼职', data.get('hr_parttime','')],
            ['临时聘用', data.get('hr_temp',''), '外籍', data.get('hr_foreign','')],
            ['留学归国', data.get('hr_returnee',''), '人才计划', data.get('hr_talent_plan','')],
        ]
    )
    add_table(
        ['学历','博士','硕士','本科','大专及以下'],
        [['人数', data.get('edu_phd',''), data.get('edu_master',''), data.get('edu_bachelor',''), data.get('edu_diploma','')]]
    )
    add_table(
        ['职称','高级职称','中级职称','初级职称','高级技工'],
        [['人数', data.get('title_senior',''), data.get('title_mid',''), data.get('title_junior',''), data.get('title_tech','')]]
    )
    hr_staff_rows = data.get('hr_staff_rows') if isinstance(data.get('hr_staff_rows'), list) else []
    if hr_staff_rows:
        add_table(
            ['序号','姓名','身份证号','是否签订合同','入职时间','是否缴纳社保','工作性质','学历','是否科技人员'],
            [[
                row.get('序号',''),
                row.get('姓名',''),
                row.get('身份证号','') or row.get('身份证',''),
                row.get('是否签订合同',''),
                row.get('入职时间',''),
                row.get('是否缴纳社保',''),
                row.get('工作性质',''),
                row.get('学历',''),
                row.get('是否科技人员',''),
            ] for row in hr_staff_rows]
        )

    # ====== 四、研发活动 ======
    add_heading('四、企业研究开发活动情况表（近三年）')
    rd_count = 0
    for i in range(20):
        name = data.get(f'rd_{i}_name','')
        if not name: continue
        rd_count += 1
        doc.add_paragraph()
        p = doc.add_paragraph()
        set_run_font(p.add_run(f"项目{rd_count}：{data.get(f'rd_{i}_no',f'RD{(rd_count):02d}')} {name}"), 13, True)
        
        add_table(
            ['项目','内容'],
            [
                ['起止时间', data.get(f'rd_{i}_period','')],
                ['技术领域', data.get(f'rd_{i}_field','')],
                ['技术来源', data.get(f'rd_{i}_source','')],
                ['知识产权编号', data.get(f'rd_{i}_ip_no','')],
                ['研发经费预算(万元)', data.get(f'rd_{i}_budget','')],
                ['近三年总支出(万元)', data.get(f'rd_{i}_total','')],
            ]
        )
        for lb, k in [
            ('目的及组织实施方式（限400字）','purpose'),
            ('核心技术及创新点（限400字）','innovation'),
            ('取得的阶段性成果（限400字）','result'),
        ]:
            v = data.get(f'rd_{i}_{k}','')
            if v:
                p = doc.add_paragraph()
                set_run_font(p.add_run(f'{lb}：'), 11, True)
                add_body(v)

    # ====== 五、费用明细 ======
    add_heading('五、企业年度研究开发费用结构明细表（万元）')
    for yi, ylabel in enumerate(['第一年','第二年','第三年']):
        add_body(f'【{ylabel}】')
        rows = []
        for rdi in range(min(rd_count, 8)):
            rd_no = f'RD{(rdi+1):02d}'
            rows.append([rd_no, data.get(f'fee_{yi}_labor_rd{rdi}','')])
        if rows:
            hdrs = ['科目'] + [r[0] for r in rows] + ['合计']
            vals = ['人员人工费用'] + [r[1] for r in rows] + [data.get(f'fee_{yi}_labor_total','')]
            add_table(hdrs, [vals])

    # ====== 六、PS ======
    add_heading('六、上年度高新技术产品（服务）情况表')
    for i in range(20):
        psname = data.get(f'ps_{i}_name','')
        if not psname: continue
        doc.add_paragraph()
        p = doc.add_paragraph()
        set_run_font(p.add_run(f"{data.get(f'ps_{i}_no',f'PS{(i+1):02d}')}：{psname}"), 13, True)
        add_table(
            ['项目','内容'],
            [
                ['技术领域', data.get(f'ps_{i}_field','')],
                ['技术来源', data.get(f'ps_{i}_source','')],
                ['销售收入(万元)', data.get(f'ps_{i}_revenue','')],
                ['是否主要产品', '是' if data.get(f'ps_{i}_is_main','yes') == 'yes' else '否'],
                ['关联RD', data.get(f'ps_{i}_rds','')],
                ['知识产权编号', data.get(f'ps_{i}_ip_no','')],
            ]
        )
        for lb, k in [
            ('关键技术及主要技术指标（限400字）','tech'),
            ('与同类产品的竞争优势（限400字）','advantage'),
            ('知识产权对产品的支持作用（限400字）','ip_support'),
        ]:
            v = data.get(f'ps_{i}_{k}','')
            if v:
                p = doc.add_paragraph()
                set_run_font(p.add_run(f'{lb}：'), 11, True)
                add_body(v)

    # ====== 七、创新能力 ======
    add_heading('七、企业创新能力')
    cv_rows = []
    cv_descriptions = []
    for i in range(50):
        desc = re.sub(r'\n[ \t　]*\n+', '\n', str(data.get(f'cv_{i}_desc','') or '').strip())
        rd = data.get(f'cv_{i}_rd','')
        ip = data.get(f'cv_{i}_ip','')
        ps = data.get(f'cv_{i}_ps','')
        result_name = data.get(f'cv_{i}_result_name','')
        if desc or rd or ip or ps or result_name:
            seq = str(len(cv_rows) + 1)
            cv_rows.append([seq, rd, ip, ps, result_name])
            cv_descriptions.append((seq, result_name, desc))
    if cv_rows:
        set_run_font(doc.add_paragraph().add_run('成果转化明细表'), 12, True)
        add_table(['序号','RD项目','知识产权','PS产品','成果名称'], cv_rows, [1.1, 3.1, 3.4, 3.1, 4.2])
        set_run_font(doc.add_paragraph().add_run('转化结果描述'), 12, True)
        for seq, result_name, desc in cv_descriptions:
            p = doc.add_paragraph()
            set_run_font(p.add_run(f'{seq}. {result_name or "成果转化"}：'), 11, True)
            add_body(desc or '（未填写）')

    for lb, k in [
        ('知识产权对企业竞争力的作用','innovation_ip'),
        ('科技成果转化情况','innovation_transform'),
        ('研究开发与技术创新组织管理情况','innovation_rd_mgmt'),
    ]:
        v = data.get(k,'')
        p = doc.add_paragraph()
        set_run_font(p.add_run(f'{lb}（限400字）：'), 12, True)
        add_body(v or '（未填写）')

    add_heading('组织管理制度文件框架')
    add_table(
        ['制度模块','建议制度文件','框架要点及佐证材料'],
        [
            ['企业研究开发组织管理', data.get('system_rd_files',''), data.get('system_rd_points','')],
            ['研发机构与产学研合作', data.get('system_org_files',''), data.get('system_org_points','')],
            ['科技成果转化与创新平台', data.get('system_transform_files',''), data.get('system_transform_points','')],
            ['科技人员培养与绩效激励', data.get('system_talent_files',''), data.get('system_talent_points','')],
        ],
        [3.2, 5.0, 7.0]
    )

    v = data.get('innovation_staff','')
    p = doc.add_paragraph()
    set_run_font(p.add_run('管理与科技人员情况（限400字）：'), 12, True)
    add_body(v or '（未填写）')

    # ====== 八、标准 ======
    add_heading('八、企业参与国家标准或行业标准制定情况')
    std_rows = []
    for i in range(5):
        sn = data.get(f'std_{i}_name','')
        if sn:
            std_rows.append([str(i+1), sn, data.get(f'std_{i}_level',''), data.get(f'std_{i}_no',''), data.get(f'std_{i}_role','')])
    if std_rows:
        add_table(['序号','标准名称','级别','编号','参与方式'], std_rows)
    else:
        add_body('（未填写）')

    tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
    doc.save(tmp.name)
    tmp.close()
    from flask import send_file
    return send_file(tmp.name, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    as_attachment=True, download_name=f'高新技术企业认定申请书_{company.name}.docx')
