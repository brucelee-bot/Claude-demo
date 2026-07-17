"""
申报书生成器 — 基于模板填充数据生成 .docx
"""
import os
import shutil
from datetime import datetime
from copy import deepcopy
from typing import Any

from docx import Document
from docx.shared import Pt, Cm

from modules.docgen.document_headers import add_company_header


TEMPLATE_PATH = os.path.expanduser("~/Desktop/高新技术企业认定申请书模板.docx")

# 字段映射：{ table_index: { row: { col: "field_name" } } }
FIELD_MAP = {
    # 一、主要情况 (Table 0)
    0: {
        0: {0: "tech_field"},           # 技术领域
        1: {1: "ip_class1_count", 3: "ip_class2_count"},  # I类/II类数量
        2: {1: "staff_total", 3: "tech_staff"},           # 职工总数/科技人员
        3: {2: "_hdr_assets", 3: "_hdr_sales", 4: "_hdr_profit"},  # header
        4: {1: "year1", 2: "year1_net_assets", 3: "year1_sales", 5: "year1_profit"},
        5: {1: "year2", 2: "year2_net_assets", 3: "year2_sales", 5: "year2_profit"},
        6: {1: "year3", 2: "year3_net_assets", 3: "year3_sales", 5: "year3_profit"},
        7: {1: "rd_total_3y", 3: "rd_domestic", 4: "rd_basic"},
        9: {1: "revenue_1y"},
        10: {1: "hitech_revenue_1y"},
        11: {2: "no_violation"},        # 是否违法
    },
}


def _safe_str(val, default=""):
    if val is None:
        return default
    return str(val)


def _set_cell_text(table, row_idx, col_idx, text):
    """安全设置单元格文本"""
    try:
        cell = table.rows[row_idx].cells[col_idx]
        # 清除现有段落内容，保留第一个段落
        for p in cell.paragraphs:
            for run in p.runs:
                run.text = ""
        if cell.paragraphs:
            cell.paragraphs[0].text = _safe_str(text)
        else:
            cell.add_paragraph(_safe_str(text))
    except IndexError:
        pass


def generate(data: dict, output_path: str = None) -> str:
    """
    生成申报书

    data 结构:
    {
        "company_name": "XX科技",
        "province": "浙江", "city": "杭州",
        "tech_field": "电子信息",
        # 知识产权
        "ip_class1_count": 2, "ip_class2_count": 8,
        "ip_list": [{"name": "xxx", "type": "发明专利", "date": "2024-01", "no": "ZL...", "method": "自主研发"}, ...],
        # 人力资源
        "staff_total": 150, "tech_staff": 45,
        "hr_detail": {"onjob": 140, "parttime": 5, "temp": 3, "foreign": 1, "returnee": 1, "talent_plan": 2},
        "hr_edu": {"博士": 5, "硕士": 20, "本科": 60, "大专及以下": 65},
        "hr_title": {"高级职称": 8, "中级职称": 25, "初级职称": 30, "高级技工": 10},
        "hr_age": {"30及以下": 40, "31-40": 60, "41-50": 35, "51及以上": 15},
        # 经营数据
        "year1": "2023", "year2": "2024", "year3": "2025",
        "year1_net_assets": 5000, ...,
        "rd_total_3y": 1200, "rd_domestic": 1100, "rd_basic": 100,
        "revenue_1y": 8000, "hitech_revenue_1y": 6500,
        "no_violation": "□否",
        # R&D 活动
        "rd_activities": [{...}, ...],
        # 产品
        "products": [{...}, ...],
        # 创新能力 (AI 生成)
        "innovation_ip_role": "知识产权对竞争力作用...(限400字)",
        "innovation_transform": "成果转化情况...(限400字)",
        "innovation_rd_mgmt": "研发组织管理...(限400字)",
        "innovation_talent": "人才情况...(限400字)",
        # 标准
        "standards": [{...}, ...],
    }
    """
    data = dict(data or {})
    data.setdefault("year1", "2023")
    data.setdefault("year2", "2024")
    data.setdefault("year3", "2025")

    if not os.path.exists(TEMPLATE_PATH):
        raise FileNotFoundError(f"模板文件不存在: {TEMPLATE_PATH}")

    if output_path is None:
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        # __file__ = modules/docgen/generator.py → 上溯3级到项目根
        project_root = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
        output_path = os.path.join(
            project_root, "outputs",
            f"国家高新技术企业认定申请书_{data.get('company_name', 'unknown')}_{ts}.docx"
        )

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    shutil.copy(TEMPLATE_PATH, output_path)

    doc = Document(output_path)
    add_company_header(
        doc,
        data.get("company_name"),
        data.get("company_english_name"),
    )

    # === 替换段落文本中的占位符 ===
    _fill_paragraphs(doc, data)

    # === 填充一、主要情况 (Table 0) ===
    if doc.tables:
        _fill_main_table(doc.tables[0], data)

    # === 填充二、知识产权明细 (Table 2) ===
    if len(doc.tables) >= 3:
        _fill_ip_table(doc.tables[2], data.get("ip_list", []))

    # === 填充三、人力资源 (Table 3) ===
    if len(doc.tables) >= 4:
        _fill_hr_table(doc.tables[3], data)

    # === 填充七、创新能力 (Table 7) ===
    if len(doc.tables) >= 8:
        _fill_innovation_table(doc.tables[7], data)

    # === 填充八、标准 (Table 8) ===
    if len(doc.tables) >= 9:
        _fill_standard_table(doc.tables[8], data.get("standards", []))

    doc.save(output_path)
    return output_path


def _fill_paragraphs(doc, data):
    """替换文档中的占位符"""
    replace_map = {
        "企业名称：": f"企业名称：{data.get('company_name', '')}",
        "企业所在地区：      省     市(区、自治州)": f"企业所在地区：{data.get('province', '')}省{data.get('city', '')}市",
    }
    for p in doc.paragraphs:
        for old, new in replace_map.items():
            if old in p.text:
                for run in p.runs:
                    if old in run.text:
                        run.text = run.text.replace(old, new)
                        break

    # 设置日期
    now = datetime.now()
    for p in doc.paragraphs:
        if "申请日期：" in p.text and "年" in p.text:
            for run in p.runs:
                run.text = run.text.replace("年", str(now.year)).replace("月", str(now.month)).replace("日", str(now.day))
            break


def _fill_main_table(table, data):
    """填充一、主要情况表"""
    for row_idx, cols in FIELD_MAP.get(0, {}).items():
        for col_idx, field in cols.items():
            if not field.startswith("_hdr_"):
                _set_cell_text(table, row_idx, col_idx, data.get(field, ""))


def _fill_ip_table(table, ip_list):
    """填充知识产权明细表 — 从第1行开始填充"""
    headers = ["知识产权编号", "知识产权名称", "类别", "授权日期", "授权号", "获得方式"]
    # 保留表头行(0)，填充后续行
    for i, ip in enumerate(ip_list[:8]):  # 最多8行
        row_idx = i + 1
        if row_idx < len(table.rows):
            _set_cell_text(table, row_idx, 0, ip.get("no", f"IP{i+1:02d}"))
            _set_cell_text(table, row_idx, 1, ip.get("name", ""))
            _set_cell_text(table, row_idx, 2, ip.get("type", ""))
            _set_cell_text(table, row_idx, 3, ip.get("date", ""))
            _set_cell_text(table, row_idx, 4, ip.get("auth_no", ""))
            _set_cell_text(table, row_idx, 5, ip.get("method", ""))


def _fill_hr_table(table, data):
    """填充人力资源表"""
    hr = data.get("hr_detail", {})
    field_map_row = {
        2: ("总   数（人）", "total"),
        3: ("其中：在职人员", "onjob"),
        4: ("兼职人员", "parttime"),
        5: ("临时聘用人员", "temp"),
        6: ("外籍人员", "foreign"),
        7: ("留学归国人员", "returnee"),
        8: ("人才计划", "talent_plan"),
    }
    for row_idx, row in enumerate(table.rows):
        first_cell = row.cells[0].text.strip()
        if row_idx in field_map_row:
            match_text, field = field_map_row[row_idx]
            if match_text in first_cell:
                _set_cell_text(table, row_idx, 2, hr.get(field, ""))
                _set_cell_text(table, row_idx, 5, hr.get(field, ""))

    # 学历
    edu = data.get("hr_edu", {})
    _set_cell_text(table, 11, 1, edu.get("博士", ""))
    _set_cell_text(table, 11, 3, edu.get("硕士", ""))
    _set_cell_text(table, 11, 4, edu.get("本科", ""))
    _set_cell_text(table, 11, 6, edu.get("大专及以下", ""))

    # 职称
    title = data.get("hr_title", {})
    _set_cell_text(table, 13, 1, title.get("高级职称", ""))
    _set_cell_text(table, 13, 3, title.get("中级职称", ""))
    _set_cell_text(table, 13, 4, title.get("初级职称", ""))
    _set_cell_text(table, 13, 6, title.get("高级技工", ""))

    # 年龄
    age = data.get("hr_age", {})
    _set_cell_text(table, 15, 1, age.get("30及以下", ""))
    _set_cell_text(table, 15, 3, age.get("31-40", ""))
    _set_cell_text(table, 15, 4, age.get("41-50", ""))
    _set_cell_text(table, 15, 6, age.get("51及以上", ""))


def _fill_innovation_table(table, data):
    """填充七、企业创新能力（4个400字文本块）"""
    fields = [
        ("innovation_ip_role", "知识产权对企业竞争力的作用"),
        ("innovation_transform", "科技成果转化情况"),
        ("innovation_rd_mgmt", "研究开发与技术创新组织管理情况"),
        ("innovation_talent", "管理与科技人员情况"),
    ]
    for i, (field, _) in enumerate(fields):
        text = data.get(field, "")
        if text and i < len(table.rows):
            _set_cell_text(table, i, 1, text)


def _fill_standard_table(table, standards):
    """填充八、标准制定"""
    for i, std in enumerate(standards[:3]):  # 最多3行
        row_idx = i + 1
        if row_idx < len(table.rows):
            _set_cell_text(table, row_idx, 0, str(i + 1))
            _set_cell_text(table, row_idx, 1, std.get("name", ""))
            _set_cell_text(table, row_idx, 3, std.get("no", ""))
            _set_cell_text(table, row_idx, 4, std.get("role", ""))
