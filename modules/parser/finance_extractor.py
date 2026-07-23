"""
财报数据智能提取器 v3 — 支持 .xls/.xlsx/.pdf/.docx，适配中国小企业报表格式
横纵双栏资产负债表自动识别
"""
import json
import os, re
from datetime import datetime

import openpyxl

try:
    import pdfplumber
except Exception:
    pdfplumber = None

# 字段映射：行标签关键词 → 字段名
VALUE_LABELS = [
    # (正则, 字段名)
    # === 资产负债表 — 先匹配更具体的合计项，避免“流动资产合计”误命中“资产合计” ===
    (r"(?<!非)流动资产合计", "current_assets"),
    (r"非流动资产合计", "noncurrent_assets"),
    (r"资产总[计額]|(?<!流动)资产合计", "assets"),
    (r"(?<!非)流动负债合计", "current_liabilities"),
    (r"非流动负债合计", "noncurrent_liabilities"),
    (r"负债和所有者权益.*合计|负债及所有者权益.*合计", "liabilities_and_equity"),
    (r"所有者权益.*合计|股东权益.*合计", "net_assets"),
    (r"(?<!流动)(?<!非流动)负债.*合计|负债总[计額]", "liabilities"),
    (r"实收资本|股本", "paid_capital"),
    (r"未分配利润", "undistributed_profit"),
    # === 利润表 ===
    (r"主营业务收入", "main_revenue"),
    (r"营业[总]?收入|营业收[入總]", "revenue"),
    (r"营业利润", "operating_profit"),
    (r"利润总额", "profit"),
    (r"净利润(?!率)", "net_profit"),
    (r"营业成本", "cost"),
    (r"销售费用", "sales_expense"),
    (r"管理费用", "admin_expense"),
    (r"研究费用|研发费用|研发投入", "rd_expense"),
    (r"上缴税金|所得税费用", "tax"),
    # === 人员 ===
    (r"职工总数|员工总数|从业人员", "employees"),
    (r"科技人员|研发人员|技术人员", "rd_staff"),
    # === 综合 ===
    (r"高新技术产品.*收入", "hitech_revenue"),
    (r"近三年研发.*总额|三年研发.*合计", "rd_total_3y"),
    (r"近一年.*总收入|上年度.*总收入", "revenue_1y"),
    (r"审计报告编码", "audit_report_code"),
]

YEAR_RE = re.compile(r"20(\d{2})(?:年|度|年末|年初|[-/.])")
COMPANY_RE = re.compile(r"(?:编制单位|企业名称|单位名称)[：:]\s*(.+?)$")


def _normalize_rows(rows):
    normalized = []
    for row in rows or []:
        normalized.append([c if c is not None else "" for c in row])
    return normalized


def _merge_results(*chunks):
    merged = {}
    for chunk in chunks:
        if not isinstance(chunk, dict):
            continue
        for key, value in chunk.items():
            if value not in (None, ""):
                merged[key] = value
    return merged


def _pdf_tables(filepath: str) -> list:
    if pdfplumber is None:
        return []
    tables = []
    try:
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages[:5]:
                for table in page.extract_tables() or []:
                    normalized = _normalize_rows(table)
                    if normalized:
                        tables.append(normalized)
    except Exception:
        return []
    return tables


def _sheet_rows_from_pdf(filepath: str) -> list:
    rows = []
    for table in _pdf_tables(filepath):
        rows.extend(table)
    return rows


def _rows_to_text(rows: list, max_rows: int = 80) -> str:
    lines = []
    for row in (rows or [])[:max_rows]:
        values = [str(c).strip() for c in row if c not in (None, "")]
        if values:
            lines.append("\t".join(values))
    return "\n".join(lines)


def _raw_text_from_xlsx(filepath: str) -> str:
    chunks = []
    try:
        wb = openpyxl.load_workbook(filepath, data_only=True)
        for sheet_name in wb.sheetnames[:5]:
            ws = wb[sheet_name]
            rows = [[c if c is not None else "" for c in row] for row in ws.iter_rows(values_only=True)]
            text = _rows_to_text(rows)
            if text:
                chunks.append(f"## Sheet: {sheet_name}\n{text}")
        wb.close()
    except Exception:
        return ""
    return "\n\n".join(chunks)


def _raw_text_from_pdf(filepath: str) -> str:
    chunks = []
    table_rows = _sheet_rows_from_pdf(filepath)
    if table_rows:
        chunks.append("## PDF Tables\n" + _rows_to_text(table_rows, max_rows=120))
    try:
        import fitz
        doc = fitz.open(filepath)
        text = "\n".join(page.get_text() for page in doc[:5])
        doc.close()
        if text.strip():
            chunks.append("## PDF Text\n" + text[:12000])
    except Exception:
        pass
    return "\n\n".join(chunks)


def _raw_text_from_file(filepath: str) -> str:
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ('.xlsx',):
        return _raw_text_from_xlsx(filepath)
    if ext in ('.pdf',):
        return _raw_text_from_pdf(filepath)
    if ext in ('.docx',):
        try:
            from docx import Document
            doc = Document(filepath)
            chunks = []
            for table in doc.tables[:5]:
                rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
                chunks.append(_rows_to_text(rows, max_rows=80))
            return "\n\n".join(c for c in chunks if c)
        except Exception:
            return ""
    return ""


def _coerce_ai_financial_json(parsed: dict, fallback_year: str = "") -> dict:
    result = {}
    if not isinstance(parsed, dict):
        return result
    if parsed.get("company_name"):
        result["company_name"] = str(parsed.get("company_name")).strip()

    years_obj = parsed.get("years")
    if isinstance(years_obj, dict):
        for year, values in years_obj.items():
            year_match = re.search(r"20\d{2}", str(year))
            if not year_match or not isinstance(values, dict):
                continue
            y = year_match.group(0)
            for src, dst in [
                ("net_assets", "net_assets"),
                ("sales", "revenue"),
                ("revenue", "revenue"),
                ("main_revenue", "main_revenue"),
                ("rd_expense", "rd_expense"),
                ("employees", "employees"),
                ("rd_staff", "rd_staff"),
            ]:
                if values.get(src) not in (None, ""):
                    result[f"fin_{y}_{dst}"] = _extract_value(values.get(src)) or str(values.get(src)).strip()

    for key, value in parsed.items():
        if re.match(r"year[123]_(net_assets|sales)$", key) and value not in (None, ""):
            result[key] = _extract_value(value) or str(value).strip()
        if re.match(r"fin_20\d{2}_", key) and value not in (None, ""):
            result[key] = _extract_value(value) or str(value).strip()

    return _normalize_for_gaoxin_score(result, fallback_year)


def _financial_extract_prompt(filepath: str, raw_text: str) -> str:
    return f"""请从下面的财务报表原始文本中提取高新技术企业评分需要的数据。

只输出严格 JSON，不要 markdown，不要解释。金额保持原单位，不要换算；如果表中是元就输出元。如果某项缺失，留空字符串。

重点字段：企业名称、各年度营业收入、净资产/所有者权益合计、研发费用、职工总数、科技人员/研发人员。

字段口径：
1. net_assets 只取资产负债表中的“所有者权益合计”或“股东权益合计”，不要取“负债和所有者权益合计”。
2. sales 和 revenue 都按利润表中的“营业收入”填写；不要用“主营业务收入”替代销售收入。
3. 年份必须优先使用表头“税款所属期起止”或“所属期起止”的截止年份。例如税款所属期为 2024-01-01 至 2024-12-31，则所有数据归入 2024 年；不要用文件名年份或报表标题年份覆盖税款所属期年份。

输出格式：
{{
  "company_name": "",
  "years": {{
    "2023": {{"revenue": "", "sales": "", "net_assets": "", "rd_expense": "", "employees": "", "rd_staff": ""}},
    "2024": {{"revenue": "", "sales": "", "net_assets": "", "rd_expense": "", "employees": "", "rd_staff": ""}},
    "2025": {{"revenue": "", "sales": "", "net_assets": "", "rd_expense": "", "employees": "", "rd_staff": ""}}
  }}
}}

文件名：{os.path.basename(filepath)}
原始文本：
{raw_text[:18000]}"""


def _parse_financial_json_response(content: str) -> dict:
    content = str(content or "").strip()
    if content.startswith("```"):
        content = "\n".join(content.splitlines()[1:-1])
    try:
        parsed = json.loads(content)
        return parsed if isinstance(parsed, dict) else {}
    except Exception:
        return {}


def _llm_extract_financials_from_raw(filepath: str, source: str) -> tuple:
    raw_text = _raw_text_from_file(filepath)
    if not raw_text.strip():
        return {}, f"{source}: 未读取到可识别文本"

    fallback_year = _tax_period_year_from_text(raw_text) or _detect_file_year(filepath)
    prompt = _financial_extract_prompt(filepath, raw_text)

    try:
        from modules.ai.llm_client import call_llm
    except Exception as exc:
        return {}, f"LLM 客户端不可用: {str(exc)[:120]}"
    result = call_llm([
        {"role": "system", "content": "你是财务报表数据抽取助手，只输出可解析的 JSON。"},
        {"role": "user", "content": prompt},
    ], temperature=0.0, max_tokens=1800, timeout=60, max_attempts=1)

    if not result.get("success"):
        return {}, result.get("error") or f"{source} 调用失败"

    parsed = _parse_financial_json_response(result.get("content", ""))
    if not parsed:
        return {}, f"{source} 返回内容不是可解析 JSON"
    return _coerce_ai_financial_json(parsed, fallback_year), ""


def _primary_llm_extract_financials_from_raw(filepath: str) -> tuple:
    return _llm_extract_financials_from_raw(filepath, "primary_llm")


def _secondary_llm_extract_financials_from_raw(filepath: str) -> tuple:
    return _llm_extract_financials_from_raw(filepath, "validation_llm")


def _llm_normalize_financials(raw: dict) -> dict:
    try:
        from modules.ai.llm_client import call_llm
    except Exception:
        return {}

    prompt = f"""你是财务报表结构化抽取助手。请把下面的原始财务字段整理成 JSON，只输出 JSON，不要解释。

要求：
1. 尽量整理出 company_name 和 2023/2024/2025 三年的 net_assets、sales、revenue、rd_expense、employees、rd_staff。
2. 如果只有 fin_YYYY_* 字段，归一化成 year1/year2/year3，按年份从小到大排序。
3. 缺失值留空字符串。

原始字段：
{json.dumps(raw, ensure_ascii=False, indent=2)}

输出格式：
{{
  "company_name": "",
  "year1_net_assets": "",
  "year2_net_assets": "",
  "year3_net_assets": "",
  "year1_sales": "",
  "year2_sales": "",
  "year3_sales": ""
}}"""

    result = call_llm([
        {"role": "system", "content": "你是财务报表结构化抽取助手，只输出严格 JSON。"},
        {"role": "user", "content": prompt},
    ], temperature=0.0, max_tokens=1200, timeout=35, max_attempts=1)
    if not result.get("success"):
        return {}
    try:
        content = result.get("content", "").strip()
        if content.startswith("```"):
            content = "\n".join(content.splitlines()[1:-1])
        parsed = json.loads(content)
        return parsed if isinstance(parsed, dict) else {}
    except Exception:
        return {}


def _extract_value(text) -> str:
    """从文本/数字提取数值字符串"""
    if text is None:
        return ""
    if isinstance(text, (int, float)):
        if text == 0:
            return ""
        return str(float(text))
    text = str(text).replace(",", "").replace("，", "").strip()
    try:
        f = float(text)
        if f == 0:
            return ""
        return str(f)
    except ValueError:
        pass
    m = re.search(r'\(?-?[\d,]+\.?\d*\)?', text)
    if m:
        val = m.group(0).replace(",", "")
        negative = val.startswith("(") and val.endswith(")")
        val = val.strip("()")
        try:
            f = float(val)
            if negative:
                f = -f
            if f == 0:
                return ""
            return str(f)
        except ValueError:
            pass
    return ""


def _extract_last_value(text) -> str:
    """从包含科目和年份的整行文本中取最后一个数值，避免把年份当金额。"""
    if text is None:
        return ""
    if isinstance(text, (int, float)):
        return _extract_value(text)
    matches = re.findall(r'\(?-?[\d,]+\.?\d*\)?', str(text).replace("，", ","))
    for raw in reversed(matches):
        val = raw.replace(",", "")
        negative = val.startswith("(") and val.endswith(")")
        val = val.strip("()")
        try:
            f = float(val)
        except ValueError:
            continue
        if 1900 <= abs(f) <= 2100 and str(int(abs(f))).startswith("20"):
            continue
        if negative:
            f = -f
        if f != 0:
            return str(f)
    return ""


def _match_label(text: str):
    """匹配行标签，返回字段名或 None。"""
    if not text:
        return None
    text = str(text).strip().replace("　", "").replace(" ", "")
    if not text:
        return None
    for pattern, field in VALUE_LABELS:
        if re.search(pattern, text):
            return field
    return None


def _extract_last_value(text) -> str:
    """从包含科目和年份的整行文本中取最后一个数值，避免把年份当金额。"""
    if text is None:
        return ""
    if isinstance(text, (int, float)):
        return _extract_value(text)
    matches = re.findall(r'\(?-?[\d,]+\.?\d*\)?', str(text).replace("，", ","))
    for raw in reversed(matches):
        val = raw.replace(",", "")
        negative = val.startswith("(") and val.endswith(")")
        val = val.strip("()")
        try:
            f = float(val)
        except ValueError:
            continue
        if 1900 <= abs(f) <= 2100 and str(int(abs(f))).startswith("20"):
            continue
        if negative:
            f = -f
        if f != 0:
            return str(f)
    return ""


    """匹配行标签，返回 (field_name, regex_match) 或 None"""
    if not text:
        return None
    text = str(text).strip().replace("　", "").replace(" ", "")
    if not text:
        return None
    for pattern, field in VALUE_LABELS:
        m = re.search(pattern, text)
        if m:
            return field
    return None


def _years_from_text(text: str) -> list:
    years = []
    for m in re.finditer(r"20\d{2}(?=\s*(?:年|年度|度|年末|年初|[-/.]))", str(text)):
        year = m.group(0)
        if year not in years:
            years.append(year)
    return years


def _tax_period_year_from_text(text: str) -> str:
    text = re.sub(r"\s+", "", str(text or ""))
    date = r"20\d{2}(?:年\d{1,2}月(?:\d{1,2}日?)?|[-/.]\d{1,2}(?:[-/.]\d{1,2})?)"
    patterns = [
        rf"(?:税款|款税)?所属期(?:起止|起止日期|起始|起)?[：:：]?{date}(?:至|到|-|—|~|－)({date})",
        rf"(?:税款|款税)?所属期(?:止|截止|终止)?[：:：]?.*?({date})",
    ]
    for pattern in patterns:
        m = re.search(pattern, text)
        if m:
            year_match = re.search(r"20\d{2}", m.group(1))
            if year_match:
                return year_match.group(0)
    return ""


def _tax_period_year_from_lines(lines: list) -> str:
    text = "\n".join(str(line or "") for line in lines[:80])
    return _tax_period_year_from_text(text)


def _tax_period_year_from_rows(rows: list) -> str:
    lines = [" ".join(str(cell or "") for cell in row) for row in (rows or [])[:80]]
    return _tax_period_year_from_lines(lines)


def _detect_file_year(filepath: str) -> str:
    years = _years_from_text(os.path.basename(filepath))
    return years[0] if years else ""


def _recent_years(count: int = 3) -> list:
    gaoxin_years = ["2023", "2024", "2025"]
    if count <= len(gaoxin_years):
        return gaoxin_years[-count:]
    first = int(gaoxin_years[0])
    prefix = [str(year) for year in range(first - (count - len(gaoxin_years)), first)]
    return prefix + gaoxin_years


def _normalize_for_gaoxin_score(data: dict, fallback_year: str = "") -> dict:
    """补充高新评分表单需要的 year1/year2/year3 字段。"""
    result = dict(data or {})

    if fallback_year:
        for field in ("net_assets", "revenue", "sales", "rd_expense", "employees", "rd_staff"):
            if result.get(field) and not result.get(f"fin_{fallback_year}_{field}"):
                result[f"fin_{fallback_year}_{field}"] = result[field]

    years = sorted({m.group(1) for key in result for m in [re.match(r"fin_(20\d{2})_", key)] if m})
    if not years:
        years = _recent_years()
    else:
        years = years[-3:]
        if len(years) < 3:
            first = int(years[0])
            prefix = [str(y) for y in range(first - (3 - len(years)), first)]
            years = prefix + years

    for idx, year in enumerate(years[-3:], start=1):
        net_assets = result.get(f"fin_{year}_net_assets")
        sales = result.get(f"fin_{year}_revenue")
        if net_assets and not result.get(f"year{idx}_net_assets"):
            result[f"year{idx}_net_assets"] = net_assets
        if sales and not result.get(f"year{idx}_sales"):
            result[f"year{idx}_sales"] = sales

    return result



def _looks_like_report_line_no(value: str) -> bool:
    value = str(value or "").strip().replace("，", "").replace(",", "")
    if not re.fullmatch(r"\d{1,3}(?:\.0)?", value):
        return False
    try:
        number = float(value)
    except ValueError:
        return False
    return 1 <= number <= 120 and number == int(number)


def _extract_numeric_tokens(text: str) -> list:
    tokens = []
    for raw in re.findall(r"\(?-?[\d,，]+\.?\d*\)?", str(text or "")):
        normalized = raw.replace("，", "").replace(",", "").strip()
        if not normalized or _looks_like_report_line_no(normalized):
            continue
        try:
            value = float(normalized.strip("()"))
        except ValueError:
            continue
        if 1900 <= abs(value) <= 2100 and str(int(abs(value))).startswith("20"):
            continue
        if value != 0:
            tokens.append(raw)
    return tokens


def _extract_pdf_value_after_label(lines: list, start_index: int) -> str:
    candidates = []
    start_line = str(lines[start_index] if start_index < len(lines) else "")
    start_match = None
    for pattern, _ in VALUE_LABELS:
        m = re.search(pattern, start_line)
        if m and (start_match is None or m.start() < start_match.start()):
            start_match = m
    if start_match:
        candidates.extend(_extract_numeric_tokens(start_line[start_match.end():]))
    if not candidates:
        candidates.extend(_extract_numeric_tokens(start_line))

    for offset in range(1, 8):
        idx = start_index + offset
        if idx >= len(lines):
            break
        text = str(lines[idx]).strip()
        if not text or text == "\xa0":
            continue
        if _match_label(text):
            break
        candidates.extend(_extract_numeric_tokens(text))

    for raw in candidates:
        value = _extract_value(raw)
        if value:
            return value
    return ""



def _clean_company_name(value: str) -> str:
    value = str(value or "").strip().strip("：: ")
    value = re.sub(r"\s+", "", value)
    value = re.sub(r"(资产负债表|利润表|现金流量表|会小企\d+表.*)$", "", value)
    return value.strip()


def _extract_company_name_from_lines(lines: list) -> str:
    labels = ("编制单位", "企业名称", "单位名称", "公司名称")
    for i, raw in enumerate(lines[:80]):
        line = str(raw or "").strip()
        for label in labels:
            if label not in line:
                continue
            tail = line.split(label, 1)[1]
            tail = re.sub(r"^[：:\s]*", "", tail)
            if tail:
                name = _clean_company_name(tail)
                if name:
                    return name
            for nxt in lines[i + 1:i + 4]:
                name = _clean_company_name(nxt)
                if name and not any(skip in name for skip in ["资产负债表", "利润表", "行次", "金额", "余额"]):
                    return name
    return ""


def _extract_pdf_lines(lines: list, fallback_year: str = "") -> dict:
    result = {}
    company_name = _extract_company_name_from_lines(lines)
    if company_name:
        result["company_name"] = company_name
    tax_period_year = _tax_period_year_from_lines(lines)
    current_year = tax_period_year or fallback_year
    for i, line in enumerate(lines):
        line = str(line).strip()
        years = _years_from_text(line)
        if years and not (tax_period_year or fallback_year):
            current_year = years[-1]

        for pattern, field in VALUE_LABELS:
            if re.search(pattern, line):
                val = _extract_pdf_value_after_label(lines, i)
                if val:
                    if current_year:
                        result[f"fin_{current_year}_{field}"] = val
                    result[field] = val
                break
    return result


def _extract_pdf_ocr(filepath: str, fallback_year: str = "") -> dict:
    try:
        import fitz
        import pytesseract
        from PIL import Image
        import io

        doc = fitz.open(filepath)
        chunks = []
        for page in doc[: min(len(doc), 5)]:
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            image = Image.open(io.BytesIO(pix.tobytes("png")))
            chunks.append(pytesseract.image_to_string(image, lang="chi_sim+eng"))
        doc.close()
        return _extract_pdf_lines("\n".join(chunks).split("\n"), fallback_year)
    except Exception:
        return {}


# ═══════════════════════════════════════════════════════════════
#  双栏资产负债表提取
# ═══════════════════════════════════════════════════════════════

def _extract_balance_sheet(rows: list) -> dict:
    """
    处理小企业资产负债表格式：左右双栏
    左栏 (cols 0-3): 资产 | 行次 | 期末余额 | 年初余额
    右栏 (cols 4-7): 负债和所有者权益 | 行次 | 期末余额 | 年初余额
    年份从表头行获取（如 "2023-12-31"）
    """
    result = {}
    sheet_year = _tax_period_year_from_rows(rows)

    # 1. 找年份 + 企业名称
    for row in rows[:6]:
        for cell in row:
            text = str(cell)
            if not sheet_year:
                m = YEAR_RE.search(text)
                if m:
                    sheet_year = f"20{m.group(1)}"
            m2 = COMPANY_RE.search(text)
            if m2:
                result["company_name"] = m2.group(1).strip()
        if sheet_year:
            break

    if not sheet_year:
        return result

    # 2. 遍历数据行
    for row in rows:
        if not row:
            continue

        # 左栏：col 0 = 标签, col 2 = 期末余额, col 3 = 年初余额
        left_label = str(row[0]) if len(row) > 0 else ""
        left_field = _match_label(left_label)
        if left_field:
            val = _extract_value(row[2]) if len(row) > 2 else ""
            if val:
                result[f"fin_{sheet_year}_{left_field}"] = val

        # 右栏：col 4 = 标签, col 6 = 期末余额, col 7 = 年初余额
        right_label = str(row[4]) if len(row) > 4 else ""
        right_field = _match_label(right_label)
        if right_field:
            val = _extract_value(row[6]) if len(row) > 6 else ""
            if val:
                result[f"fin_{sheet_year}_{right_field}"] = val

    return result


def _looks_like_sequence_col(rows: list, ci: int) -> bool:
    vals = []
    for row in rows[3:min(12, len(rows))]:
        if ci >= len(row):
            continue
        try:
            v = float(str(row[ci]))
            if v == int(v) and 1 <= v <= 80:
                vals.append(int(v))
        except (ValueError, TypeError):
            pass
    return len(vals) >= 3 and vals == list(range(vals[0], vals[0] + len(vals)))


def _preferred_value_columns(rows: list) -> list:
    preferred = []
    for ri in range(min(8, len(rows))):
        for ci, cell in enumerate(rows[ri]):
            text = str(cell or "")
            if any(k in text for k in ["本年累计", "期末余额", "金额", "余额"]):
                if ci not in preferred and not _looks_like_sequence_col(rows, ci):
                    preferred.append(ci)
    return preferred


def _find_label_cell(row: list):
    for ci, cell in enumerate(row[: min(len(row), 4)]):
        field = _match_label(str(cell))
        if field:
            return ci, field
    return -1, None


def _extract_value_from_row_after_label(rows: list, row: list, label_ci: int, preferred_cols: list) -> str:
    candidate_cols = [ci for ci in preferred_cols if ci > label_ci]
    if not candidate_cols:
        candidate_cols = [ci for ci in range(label_ci + 1, len(row)) if not _looks_like_sequence_col(rows, ci)]
    for ci in candidate_cols:
        if ci < len(row):
            val = _extract_value(row[ci])
            if val:
                return val
    return ""


def _extract_generic(rows: list) -> dict:
    """
    通用提取：标准横向表格（表头含年份列 + 数据行）
    """
    result = {}

    # 检测列年份
    tax_period_year = _tax_period_year_from_rows(rows)
    col_years = {}
    for ri in range(min(8, len(rows))):
        for ci, cell in enumerate(rows[ri]):
            text = str(cell or "")
            if re.fullmatch(r"[\d,，]+(?:\.\d+)?", text.strip()):
                continue
            m = YEAR_RE.search(text)
            if m:
                y = f"20{m.group(1)}"
                if ci not in col_years:
                    col_years[ci] = y

    # 排除行次列（值 = 1.0, 2.0, 3.0... 的列）
    bad_cols = set()
    label_positions = [label_ci for row in rows for label_ci, field in [_find_label_cell(row)] if field]
    min_label_ci = min(label_positions) if label_positions else 0
    for ci in list(col_years.keys()):
        if label_positions and ci <= min_label_ci:
            bad_cols.add(ci)
            continue
        vals = []
        for row in rows[3:min(10, len(rows))]:
            if ci < len(row):
                try:
                    v = float(str(row[ci]))
                    if v == int(v) and 1 <= v <= 50:
                        vals.append(int(v))
                except (ValueError, TypeError):
                    pass
        if len(vals) >= 3 and vals == list(range(vals[0], vals[0] + len(vals))):
            bad_cols.add(ci)
    for ci in bad_cols:
        del col_years[ci]

    # 找全局年份（如果没检测到列级年份）
    global_year = tax_period_year
    if not col_years and not global_year:
        for row in rows[:6]:
            for cell in row:
                text = str(cell or "")
                if re.fullmatch(r"[\d,，]+(?:\.\d+)?", text.strip()):
                    continue
                m = YEAR_RE.search(text)
                if m:
                    global_year = f"20{m.group(1)}"
                    break
            if global_year:
                break

    preferred_cols = _preferred_value_columns(rows)

    # 遍历数据行
    for row in rows:
        if not row:
            continue
        label_ci, field = _find_label_cell(row)
        if not field:
            continue

        if col_years:
            for ci, year in col_years.items():
                if ci <= label_ci or ci >= len(row):
                    continue
                val = _extract_value(row[ci])
                if val:
                    result[f"fin_{year}_{field}"] = val
        elif global_year:
            val = _extract_value_from_row_after_label(rows, row, label_ci, preferred_cols)
            if val:
                result[f"fin_{global_year}_{field}"] = val
        else:
            val = _extract_value_from_row_after_label(rows, row, label_ci, preferred_cols)
            if val:
                result[field] = val

    return result


def _is_balance_sheet(rows: list) -> bool:
    """判断是否为中国小企业资产负债表（左右双栏格式）"""
    if len(rows) < 3:
        return False
    for row in rows[:5]:
        left = str(row[0]).strip() if len(row) > 0 else ""
        right = str(row[4]).strip() if len(row) > 4 else ""
        if "资产" in left and ("负债" in right or "所有者" in right):
            return True
        if "流动" in left and "流动" in right:
            return True
    return False


def _extract_sheet(rows: list) -> dict:
    """提取单个 sheet 的数据"""
    if _is_balance_sheet(rows):
        return _extract_balance_sheet(rows)
    return _extract_generic(rows)


# ═══════════════════════════════════════════════════════════════
#  文件格式读取器
# ═══════════════════════════════════════════════════════════════

def extract_from_xlsx(filepath: str) -> dict:
    """openpyxl 读取 .xlsx"""
    result = {}
    wb = openpyxl.load_workbook(filepath, data_only=True)
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = [[c if c is not None else "" for c in row] for row in ws.iter_rows(values_only=True)]
        result.update(_extract_sheet(rows))
    wb.close()
    return result


def extract_from_xls(filepath: str) -> dict:
    """
    读取 .xls（旧格式 Excel）
    优先使用 calamine（Rust，快速），失败回退到 xlrd3，最后用 LibreOffice 转换为 xlsx 再解析。
    """
    errors = []

    # 优先 calamine
    try:
        from python_calamine import CalamineWorkbook
        result = {}
        wb = CalamineWorkbook.from_path(filepath)
        for sheet_name in wb.sheet_names:
            rows = wb.get_sheet_by_name(sheet_name).to_python()
            normalized = [[c if c is not None else "" for c in row] for row in rows]
            result.update(_extract_sheet(normalized))
        return result
    except ImportError as exc:
        errors.append(f"python-calamine 未安装: {exc}")
    except Exception as exc:
        errors.append(f"python-calamine 解析失败: {exc}")

    # 回退 xlrd3（兼容部分 calamine 不支持的编码）
    try:
        import xlrd3
        result = {}
        wb = xlrd3.open_workbook(filepath)
        for sheet_name in wb.sheet_names():
            ws = wb.sheet_by_name(sheet_name)
            rows = [[ws.cell_value(r, c) for c in range(ws.ncols)] for r in range(ws.nrows)]
            result.update(_extract_sheet(rows))
        return result
    except ImportError as exc:
        errors.append(f"xlrd3 未安装: {exc}")
    except Exception as exc:
        errors.append(f"xlrd3 解析失败: {exc}")

    # 最后兜底：LibreOffice 头less 转成 xlsx，再复用 xlsx 解析逻辑
    try:
        import subprocess
        import tempfile
        import shutil

        office_bin = shutil.which("libreoffice") or shutil.which("soffice")
        if not office_bin:
            raise FileNotFoundError("未找到 libreoffice/soffice")

        with tempfile.TemporaryDirectory() as tmpdir:
            subprocess.run(
                [
                    office_bin,
                    "--headless",
                    "--nologo",
                    "--nolockcheck",
                    "--nodefault",
                    "--convert-to",
                    "xlsx",
                    "--outdir",
                    tmpdir,
                    filepath,
                ],
                check=True,
                capture_output=True,
            )

            base = os.path.splitext(os.path.basename(filepath))[0]
            converted = os.path.join(tmpdir, f"{base}.xlsx")
            if not os.path.exists(converted):
                # 兼容 LibreOffice 可能重命名输出文件的情况
                candidates = [
                    os.path.join(tmpdir, f)
                    for f in os.listdir(tmpdir)
                    if f.lower().endswith(".xlsx")
                ]
                converted = candidates[0] if candidates else ""

            if converted and os.path.exists(converted):
                return extract_from_xlsx(converted)
    except Exception as exc:
        errors.append(f"LibreOffice 转换失败: {exc}")

    missing_readers = any("未安装" in error for error in errors)
    if missing_readers:
        raise ValueError("服务器暂不支持解析旧版 .xls 文件：缺少 .xls 解析依赖。请安装 python-calamine 或 xlrd3，或临时将文件另存为 .xlsx 后上传。")
    raise ValueError("无法解析此 .xls 文件：文件可能已损坏、格式过旧或不兼容。请尝试用 Excel/LibreOffice 另存为 .xlsx 后再上传。")



def extract_from_docx(filepath: str) -> dict:
    """从 Word 表格提取"""
    from docx import Document
    result = {}
    doc = Document(filepath)
    for table in doc.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        result.update(_extract_sheet(rows))
    return result


def _detect_tax_period_year_from_file(filepath: str) -> str:
    ext = os.path.splitext(filepath)[1].lower()
    if ext == ".pdf":
        raw = _raw_text_from_pdf(filepath)
    else:
        raw = _raw_text_from_file(filepath)
    return _tax_period_year_from_text(raw)


def extract_from_pdf(filepath: str) -> dict:
    """从 PDF 提取"""
    fallback_year = _detect_tax_period_year_from_file(filepath) or _detect_file_year(filepath)

    table_result = {}
    for table_rows in _pdf_tables(filepath):
        table_result.update(_extract_sheet(table_rows))
    if table_result:
        return _normalize_for_gaoxin_score(table_result, fallback_year)

    try:
        import fitz
        doc = fitz.open(filepath)
        text = "\n".join(page.get_text() for page in doc)
        doc.close()

        result = _extract_pdf_lines(text.split('\n'), fallback_year)
        if result:
            return _normalize_for_gaoxin_score(result, fallback_year)
    except Exception:
        pass

    try:
        import subprocess, tempfile
        with tempfile.NamedTemporaryFile(suffix='.txt', delete=False) as tmp:
            txt_path = tmp.name
        subprocess.run(['pdftotext', '-layout', filepath, txt_path], check=True, capture_output=True)
        with open(txt_path, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()
        try:
            os.remove(txt_path)
        except OSError:
            pass

        result = _extract_pdf_lines(text.split('\n'), fallback_year)
        if result:
            return _normalize_for_gaoxin_score(result, fallback_year)
    except Exception:
        pass

    if os.getenv("FINANCE_PDF_OCR", "").strip().lower() in {"1", "true", "yes", "on"}:
        ocr_result = _extract_pdf_ocr(filepath, fallback_year)
        if ocr_result:
            return _normalize_for_gaoxin_score(ocr_result, fallback_year)
    return {}


# ═══════════════════════════════════════════════════════════════
#  入口
# ═══════════════════════════════════════════════════════════════

def _add_fixed_growth_year_fields(data: dict) -> dict:
    result = dict(data or {})
    for idx, year in enumerate(["2023", "2024", "2025"], start=1):
        net_assets = result.get(f"fin_{year}_net_assets")
        sales = result.get(f"fin_{year}_revenue")
        if net_assets not in (None, ""):
            result[f"year{idx}_net_assets"] = net_assets
        if sales not in (None, ""):
            result[f"year{idx}_sales"] = sales
    return result


def _strip_single_file_year_fields(data: dict) -> dict:
    result = dict(data or {})
    years = {m.group(1) for key in result for m in [re.match(r"fin_(20\d{2})_", key)] if m}
    if len(years) <= 1:
        for key in list(result.keys()):
            if re.match(r"year[123]_(net_assets|sales)$", key):
                result.pop(key, None)
    return result


def _has_growth_financial_fields(data: dict) -> bool:
    keys = data or {}
    return any(re.match(r"fin_20\d{2}_net_assets$", key) for key in keys) and any(
        re.match(r"fin_20\d{2}_(sales|revenue|main_revenue)$", key) for key in keys
    )


def _has_key_financial_fields(data: dict) -> bool:
    keys = data or {}
    return any(re.match(r"fin_20\d{2}_(revenue|sales|main_revenue|net_assets)$", key) for key in keys)


def _values_match(a, b) -> bool:
    if a in (None, "") or b in (None, ""):
        return False
    av = _extract_value(a) or str(a).strip()
    bv = _extract_value(b) or str(b).strip()
    try:
        af = float(av)
        bf = float(bv)
        tolerance = max(1.0, abs(af) * 0.01)
        return abs(af - bf) <= tolerance
    except (TypeError, ValueError):
        return av == bv


def _merge_missing_fields(primary: dict, secondary: dict) -> dict:
    merged = dict(primary or {})
    for key, value in (secondary or {}).items():
        if key.startswith("_") or value in (None, ""):
            continue
        if key not in merged or merged.get(key) in (None, ""):
            merged[key] = value
    return merged


def _build_validation(primary: dict, verifier: dict, primary_llm_error: str = "", verifier_error: str = "") -> dict:
    comparable_keys = sorted(
        key for key in set(primary or {}) | set(verifier or {})
        if key == "company_name" or re.match(r"fin_20\d{2}_", key)
    )
    matched = []
    mismatched = []
    only_in_primary = []
    only_in_verifier = []

    for key in comparable_keys:
        pv = (primary or {}).get(key)
        vv = (verifier or {}).get(key)
        if pv not in (None, "") and vv not in (None, ""):
            if _values_match(pv, vv):
                matched.append(key)
            else:
                mismatched.append({"field": key, "primary": pv, "verifier": vv})
        elif pv not in (None, ""):
            only_in_primary.append(key)
        elif vv not in (None, ""):
            only_in_verifier.append(key)

    return {
        "primary_source": "local_rules_then_configured_llm",
        "secondary_source": "configured_llm_validation",
        "primary_llm_error": primary_llm_error,
        "verifier_error": verifier_error,
        "matched": matched,
        "mismatched": mismatched,
        "only_in_primary": only_in_primary,
        "only_in_verifier": only_in_verifier,
        "status": "warning" if mismatched or verifier_error else "ok",
    }


def _extract_rule_data(filepath: str) -> dict:
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ('.xlsx',):
        return extract_from_xlsx(filepath)
    if ext in ('.xls',):
        return extract_from_xls(filepath)
    if ext in ('.pdf',):
        return extract_from_pdf(filepath)
    if ext in ('.docx',):
        return extract_from_docx(filepath)
    return {"error": f"不支持的文件格式: {ext}"}


def extract_with_validation(filepath: str) -> dict:
    """Extract financial data with rule recognition and a single configured LLM."""
    rule_data = _extract_rule_data(filepath)
    if rule_data.get("error"):
        return {"data": rule_data, "validation": {"status": "error"}, "sources": {}}

    primary_llm_data, primary_llm_error = _primary_llm_extract_financials_from_raw(filepath)

    fallback_year = _detect_tax_period_year_from_file(filepath) or _detect_file_year(filepath)
    primary = _merge_results(
        _normalize_for_gaoxin_score(rule_data, fallback_year),
        _normalize_for_gaoxin_score(primary_llm_data, fallback_year),
    )

    secondary_data, secondary_error = _secondary_llm_extract_financials_from_raw(filepath)
    normalized_secondary = _normalize_for_gaoxin_score(secondary_data, fallback_year)
    validation = _build_validation(primary, normalized_secondary, primary_llm_error, secondary_error)

    normalized = _merge_missing_fields(primary, normalized_secondary)
    return {
        "data": _add_fixed_growth_year_fields(normalized),
        "validation": validation,
        "sources": {
            "rule_fields": len(rule_data or {}),
            "primary_llm_fields": len(primary_llm_data or {}),
            "llm_fields": len(secondary_data or {}),
        },
    }


def extract(filepath: str) -> dict:
    """根据扩展名自动选择读取器，返回可直接回填表单的数据。"""
    return extract_with_validation(filepath).get("data", {})
